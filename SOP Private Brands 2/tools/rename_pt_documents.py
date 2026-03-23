"""
P&T SOP Document Renamer
Renames backup P&T documents to [XX] format and updates version fields to 1.0.

Usage:
    python tools/rename_pt_documents.py [--dry-run]
"""

import os
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

BACKUP_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Documents", "SOP_Pipes & tubes", "_backup_20260227_183825"
)

# Complete file mapping: old_filename -> {doc_num, clean_name, skip_version_update}
RENAME_MAPPING = {
    # --- Numbered files (suffix matches document ID) ---
    "SOP for incoming coil inspection for P & T_1.docx": {
        "doc_num": 1,
        "clean_name": "SOP for Incoming Coil Inspection",
    },
    "SOP for slitted coil inspection for P & T_2.docx": {
        "doc_num": 2,
        "clean_name": "SOP for Slitted Coil Inspection",
    },
    "SOP for final pipe inspection for P & T_3.docx": {
        "doc_num": 3,
        "clean_name": "SOP for Final Pipe Inspection",
    },
    "SOP for Joint Inspection for P & T_4.docx": {
        "doc_num": 4,
        "clean_name": "SOP for Joint Inspection",
    },
    "SOP_P & T Tensile Test joint inspection P & T_5.docx": {
        "doc_num": 5,
        "clean_name": "SOP for Tensile Test Joint Inspection",
    },
    "SOP_P & T_Bend Test_7.docx": {
        "doc_num": 7,
        "clean_name": "SOP for Bend Test",
    },
    "SOP_P & T_Flattening Test_8.docx": {
        "doc_num": 8,
        "clean_name": "SOP for Flattening Test",
    },
    "SOP_P & T_Dimension Test_9.docx": {
        "doc_num": 9,
        "clean_name": "SOP for Dimension Test",
    },
    "SOP for PDI of P&T_12.docx": {
        "doc_num": 12,
        "clean_name": "SOP for Pre-Dispatch Inspection",
    },
    "SOP for Test Certificate Generation P&T_13.docx": {
        "doc_num": 13,
        "clean_name": "SOP for Test Certificate Generation",
    },
    "SOP_P&T_Traceability and identification products_14.docx": {
        "doc_num": 14,
        "clean_name": "SOP for Traceability and Identification",
    },
    "SOP_P&T_Handling of NC products_15.docx": {
        "doc_num": 15,
        "clean_name": "SOP for Handling of NC Products",
    },
    # --- Unnumbered files (assigned next available sequential IDs) ---
    "SOP for incoming coil receiving for P & T.docx": {
        "doc_num": 6,
        "clean_name": "SOP for Incoming Coil Receiving",
    },
    "SOP for Rolling Plan.docx": {
        "doc_num": 10,
        "clean_name": "SOP for Rolling Plan",
    },
    "SOP for Finishing Activities.docx": {
        "doc_num": 11,
        "clean_name": "SOP for Finishing Activities",
    },
    "One_Helix_Pipes_Tubes_SOP_REVISED_FINAL 1.2.docx": {
        "doc_num": 16,
        "clean_name": "One Helix Plant Operations SOP",
        "skip_version_update": True,
    },
    "SOP for Dispatch Process- Pipes & Tubes.docx": {
        "doc_num": 17,
        "clean_name": "SOP for Dispatch Process",
    },
    "SOP for DO Generation for Dispatch process.docx": {
        "doc_num": 18,
        "clean_name": "SOP for DO Generation for Dispatch",
    },
    "SOP for Vehilce Placement.docx": {
        "doc_num": 19,
        "clean_name": "SOP for Vehicle Placement",
    },
    "SOP for Invoicing process.docx": {
        "doc_num": 20,
        "clean_name": "SOP for Invoicing Process",
    },
}

DUPLICATE_TO_DELETE = "SOP_P&T_Traceability and identification products_14[1].docx"


def build_new_filename(doc_num, clean_name):
    return f"[{doc_num:02d}] {clean_name}.docx"


def update_version_in_docx(filepath, version_str="1.0"):
    """Update the Version cell in the standard P&T header table (table[0], row 1, cell 2).

    Returns (success, message).
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        return False, f"Cannot open file: {e}"

    if not doc.tables:
        return False, "No tables found"

    table0 = doc.tables[0]
    if len(table0.rows) < 2:
        return False, f"Header table has only {len(table0.rows)} rows (need >= 2)"

    row1 = table0.rows[1]
    if len(row1.cells) < 3:
        return False, f"Row 1 has only {len(row1.cells)} cells (need >= 3)"

    version_cell = row1.cells[2]
    cell_text = version_cell.text.strip()

    if not cell_text.lower().startswith("version"):
        return False, f"Cell text '{cell_text}' does not start with 'Version'"

    # Capture original font size from first run
    original_font_size = None
    para = version_cell.paragraphs[0]
    if para.runs:
        original_font_size = para.runs[0].font.size

    # Clear paragraph content (preserve paragraph properties)
    p_element = para._element
    for child in list(p_element):
        if child.tag != qn('w:pPr'):
            p_element.remove(child)

    # Write new version text
    new_run = para.add_run(f"Version: {version_str} ")
    if original_font_size:
        new_run.font.size = original_font_size

    doc.save(filepath)
    return True, f"'{cell_text}' -> 'Version: {version_str}'"


def main():
    parser = argparse.ArgumentParser(description="Rename P&T SOP backup documents")
    parser.add_argument("--dry-run", action="store_true",
                        help="Preview changes without modifying files")
    args = parser.parse_args()

    backup_dir = BACKUP_DIR
    if not os.path.isdir(backup_dir):
        print(f"ERROR: Backup directory not found: {backup_dir}")
        sys.exit(1)

    prefix = "[DRY RUN] " if args.dry_run else ""
    print(f"{prefix}P&T SOP Document Renamer")
    print(f"Directory: {backup_dir}")
    print("=" * 60)

    # --- Phase 1: Validate source files ---
    print("\n[Phase 1] Validating source files...")
    errors = []
    for old_name in RENAME_MAPPING:
        old_path = os.path.join(backup_dir, old_name)
        if not os.path.isfile(old_path):
            errors.append(f"  MISSING: {old_name}")

    if errors:
        for e in errors:
            print(e)
        print("ABORT: Fix missing files before proceeding.")
        sys.exit(1)
    print(f"  All {len(RENAME_MAPPING)} source files found.")

    # --- Phase 2: Check for filename collisions ---
    print("\n[Phase 2] Checking for collisions...")
    new_names = {}
    for old_name, config in RENAME_MAPPING.items():
        new_name = build_new_filename(config["doc_num"], config["clean_name"])
        if new_name in new_names:
            errors.append(f"  COLLISION: {new_name} from both '{old_name}' and '{new_names[new_name]}'")
        new_names[new_name] = old_name

    if errors:
        for e in errors:
            print(e)
        print("ABORT: Resolve collisions.")
        sys.exit(1)
    print("  No collisions.")

    # --- Phase 3: Delete duplicate ---
    print("\n[Phase 3] Removing duplicate...")
    dup_path = os.path.join(backup_dir, DUPLICATE_TO_DELETE)
    if os.path.isfile(dup_path):
        if args.dry_run:
            print(f"  Would delete: {DUPLICATE_TO_DELETE}")
        else:
            os.remove(dup_path)
            print(f"  Deleted: {DUPLICATE_TO_DELETE}")
    else:
        print(f"  Already removed: {DUPLICATE_TO_DELETE}")

    # --- Phase 4: Update versions ---
    print("\n[Phase 4] Updating version numbers...")
    sorted_items = sorted(RENAME_MAPPING.items(), key=lambda x: x[1]["doc_num"])

    for old_name, config in sorted_items:
        skip_version = config.get("skip_version_update", False)
        old_path = os.path.join(backup_dir, old_name)

        if skip_version:
            print(f"  [{config['doc_num']:02d}] SKIP version (non-standard): {old_name}")
            continue

        if args.dry_run:
            print(f"  [{config['doc_num']:02d}] Would update version to 1.0: {old_name}")
        else:
            success, msg = update_version_in_docx(old_path, "1.0")
            status = "OK" if success else "WARN"
            print(f"  [{config['doc_num']:02d}] {status}: {msg}")

    # --- Phase 5: Rename files (two-pass via temp names) ---
    print("\n[Phase 5] Renaming files...")
    temp_pairs = []

    for old_name, config in sorted_items:
        new_name = build_new_filename(config["doc_num"], config["clean_name"])
        old_path = os.path.join(backup_dir, old_name)
        new_path = os.path.join(backup_dir, new_name)

        print(f"  [{config['doc_num']:02d}] {old_name}")
        print(f"       -> {new_name}")

        if old_name == new_name:
            print("       (same name, skip)")
            continue

        if args.dry_run:
            continue

        # Rename to temp first
        temp_path = old_path + ".renaming"
        try:
            os.rename(old_path, temp_path)
            temp_pairs.append((temp_path, new_path, new_name))
        except OSError as e:
            print(f"       ERROR: {e}")
            if os.path.isfile(temp_path) and not os.path.isfile(old_path):
                os.rename(temp_path, old_path)

    # Finalize temp -> new
    for temp_path, new_path, new_name in temp_pairs:
        try:
            os.rename(temp_path, new_path)
        except OSError as e:
            print(f"       ERROR finalizing {new_name}: {e}")

    # --- Phase 6: Verify ---
    print(f"\n{'=' * 60}")
    print("[Phase 6] Verification:")

    if args.dry_run:
        print("  Dry run complete. No files were modified.")
    else:
        expected = {build_new_filename(c["doc_num"], c["clean_name"])
                    for c in RENAME_MAPPING.values()}
        actual = {f for f in os.listdir(backup_dir) if f.endswith(".docx")}

        missing = expected - actual
        extra = actual - expected

        if missing:
            print(f"  MISSING: {missing}")
        if extra:
            print(f"  UNEXPECTED: {extra}")
        if not missing and not extra:
            print(f"  All {len(expected)} files present with correct names.")

    print("\nDone.")


if __name__ == "__main__":
    main()
