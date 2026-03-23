"""
JSW ONE Pipes & Tubes SOP Master Excel Generator

Reads all SOP documents from Documents/SOP_Pipes & tubes/ folder,
parses them, classifies into journey/section, and generates a formatted
master Excel file matching the Private Brands output format.

Usage:
    python tools/generate_pt_sop_excel.py
"""

import hashlib
import json
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from collections import defaultdict

from docx import Document

# Add tools directory to path
TOOLS_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, TOOLS_DIR)

from models.sop_data import Activity, Section, Journey
from excel_writer.writer import write_master_excel


# ---------------------------------------------------------------------------
# Document classification: maps filename patterns to journey/section
# ---------------------------------------------------------------------------

PT_DOCUMENT_MAPPING = {
    # Pre-Production: Raw Material Management
    "SOP for incoming coil receiving*": {
        "journey": "Pre-Production",
        "section": "A. RAW MATERIAL MANAGEMENT",
        "activity_prefix": "Incoming Coil Receiving",
    },
    "SOP for incoming coil inspection*": {
        "journey": "Pre-Production",
        "section": "A. RAW MATERIAL MANAGEMENT",
        "activity_prefix": "Incoming Coil Inspection",
    },
    "SOP for slitted coil inspection*": {
        "journey": "Pre-Production",
        "section": "A. RAW MATERIAL MANAGEMENT",
        "activity_prefix": "Slitted Coil Inspection",
    },
    # Pre-Production: Production Planning
    "SOP for Rolling Plan*": {
        "journey": "Pre-Production",
        "section": "B. PRODUCTION PLANNING",
        "activity_prefix": "Rolling Plan",
    },
    # Production: Manufacturing & Finishing
    "SOP for Finishing Activities*": {
        "journey": "Production",
        "section": "C. MANUFACTURING & FINISHING",
        "activity_prefix": "Finishing Activities",
    },
    "One_Helix*": {
        "journey": "Production",
        "section": "C. MANUFACTURING & FINISHING",
        "activity_prefix": "Pipes & Tubes Plant Operations",
        "is_comprehensive": True,
    },
    # Production: Quality Testing
    "SOP for Joint Inspection*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Joint Inspection",
    },
    "SOP*Tensile Test*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Tensile Test & Joint Inspection",
    },
    "SOP*Bend Test*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Bend Test",
    },
    "SOP*Flattening Test*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Flattening Test",
    },
    "SOP*Dimension Test*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Dimension Test",
    },
    "SOP for final pipe inspection*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Final Pipe Inspection",
    },
    "SOP for PDI*": {
        "journey": "Production",
        "section": "D. QUALITY TESTING",
        "activity_prefix": "Pre-Dispatch Inspection (PDI)",
    },
    # Production: Quality Management
    "SOP*Handling of NC*": {
        "journey": "Production",
        "section": "E. QUALITY MANAGEMENT",
        "activity_prefix": "Handling of Non-Conforming Products",
    },
    "SOP*Traceability*": {
        "journey": "Production",
        "section": "E. QUALITY MANAGEMENT",
        "activity_prefix": "Traceability & Product Identification",
    },
    # Post-Production: Dispatch & Logistics
    "SOP for Dispatch Process*": {
        "journey": "Post-Production",
        "section": "F. DISPATCH & LOGISTICS",
        "activity_prefix": "Dispatch Process",
    },
    "SOP for DO Generation*": {
        "journey": "Post-Production",
        "section": "F. DISPATCH & LOGISTICS",
        "activity_prefix": "DO Generation for Dispatch",
    },
    "SOP for Veh*Placement*": {
        "journey": "Post-Production",
        "section": "F. DISPATCH & LOGISTICS",
        "activity_prefix": "Vehicle Placement",
    },
    # Post-Production: Documentation & Invoicing
    "SOP for Invoicing*": {
        "journey": "Post-Production",
        "section": "G. DOCUMENTATION & INVOICING",
        "activity_prefix": "Invoicing Process",
    },
    "SOP for Test Certificate*": {
        "journey": "Post-Production",
        "section": "G. DOCUMENTATION & INVOICING",
        "activity_prefix": "Test Certificate Generation",
    },
}

PT_SECTION_ORDER = {
    "Pre-Production": [
        "A. RAW MATERIAL MANAGEMENT",
        "B. PRODUCTION PLANNING",
    ],
    "Production": [
        "C. MANUFACTURING & FINISHING",
        "D. QUALITY TESTING",
        "E. QUALITY MANAGEMENT",
    ],
    "Post-Production": [
        "F. DISPATCH & LOGISTICS",
        "G. DOCUMENTATION & INVOICING",
    ],
}

PT_JOURNEY_ORDER = ["Pre-Production", "Production", "Post-Production"]


# ---------------------------------------------------------------------------
# Document classifier
# ---------------------------------------------------------------------------

def classify_pt_document(filename: str) -> dict | None:
    """Classify a P&T document by filename pattern matching."""
    import fnmatch

    normalized = filename.replace('+', ' ').replace('  ', ' ')

    for pattern, config in PT_DOCUMENT_MAPPING.items():
        pattern_normalized = pattern.replace('+', ' ')
        if fnmatch.fnmatch(normalized, pattern_normalized) or fnmatch.fnmatch(filename, pattern):
            return config

    # Fallback: keyword matching
    name_lower = normalized.lower()
    for pattern, config in PT_DOCUMENT_MAPPING.items():
        keywords = [w.strip('*').lower() for w in pattern.split('*') if w.strip('*')]
        if all(kw in name_lower for kw in keywords if len(kw) > 2):
            return config

    return None


# ---------------------------------------------------------------------------
# P&T Document parser
# ---------------------------------------------------------------------------

def parse_pt_docx(filepath: str) -> dict:
    """Parse a Pipes & Tubes .docx SOP document.

    Returns a dict with:
        filename, title, doc_id, version, date,
        teams (list of {team, responsibility}),
        purpose, scope, process_steps (list of str),
        process_tables (list of table data),
        approval_info, raw_text
    """
    doc = Document(filepath)
    filename = os.path.basename(filepath)

    result = {
        'filename': filename,
        'title': '',
        'doc_id': '',
        'version': '',
        'date': '',
        'teams': [],
        'purpose': '',
        'scope': '',
        'process_steps': [],
        'process_tables': [],
        'approval_info': '',
        'raw_text': '',
    }

    # Extract all tables
    all_tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            all_tables.append(rows)

    # Extract all paragraph text
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    result['raw_text'] = '\n'.join(paragraphs)

    # --- Parse header table (table 0) ---
    # Standard JSW ONE template: 3 rows x 3 columns
    #   Row 0: "A JSW ONE PRODUCT" | "STANDARD OPERATING PROCEDURE" | "Document ID:..."
    #   Row 1: "A JSW ONE PRODUCT" | "STANDARD OPERATING PROCEDURE" | "Version: 1"
    #   Row 2: "A JSW ONE PRODUCT" | <PROCESS NAME>                 | "Date: DD/MM/YYYY"
    if all_tables:
        header_table = all_tables[0]
        for row in header_table:
            for cell in row:
                cell_lower = cell.lower().strip()
                # Extract document ID
                id_match = re.search(r'(JODL[^\s|]+)', cell)
                if id_match:
                    result['doc_id'] = id_match.group(1)
                # Extract version
                if cell_lower.startswith('version'):
                    result['version'] = re.sub(r'^version[:\s]*', '', cell, flags=re.IGNORECASE).strip()
                # Extract date
                if cell_lower.startswith('date'):
                    result['date'] = re.sub(r'^date[:\s]*', '', cell, flags=re.IGNORECASE).strip()

        # Title is in the last row (row 2), middle column - the process name
        # It's the cell that is NOT "A JSW ONE PRODUCT", NOT "Date:...", NOT empty
        if len(header_table) >= 3:
            last_row = header_table[-1]
            for cell in last_row:
                cell_upper = cell.upper().strip()
                if cell and \
                        'JSW ONE' not in cell_upper and \
                        'STANDARD OPERATING' not in cell_upper and \
                        not cell_upper.startswith('DATE') and \
                        not cell.startswith('JODL') and \
                        len(cell) > 3:
                    result['title'] = cell
                    break

    if not result['title']:
        result['title'] = filename.replace('.docx', '').replace('_', ' ')

    # --- Parse teams table (table 1 - "Sl No | Team | Responsibility") ---
    if len(all_tables) > 1:
        teams_table = all_tables[1]
        if teams_table and len(teams_table) > 1:
            header_lower = [h.lower() for h in teams_table[0]]
            is_teams = any('team' in h for h in header_lower) or any('responsibility' in h for h in header_lower)
            if is_teams:
                team_col = next((i for i, h in enumerate(header_lower) if 'team' in h), 1)
                resp_col = next((i for i, h in enumerate(header_lower) if 'resp' in h), 2)
                for row in teams_table[1:]:
                    team = row[team_col] if team_col < len(row) else ''
                    resp = row[resp_col] if resp_col < len(row) else ''
                    if team.strip():
                        result['teams'].append({'team': team.strip(), 'responsibility': resp.strip()})

    # --- Parse approval table (last table) ---
    if len(all_tables) > 2:
        approval_table = all_tables[-1]
        if approval_table:
            approval_text = ' | '.join(
                cell for row in approval_table for cell in row if cell.strip()
            )
            if any(kw in approval_text.lower() for kw in ['created', 'approved', 'reviewed']):
                result['approval_info'] = approval_text

    # --- Parse process-specific tables (tables between teams and approval) ---
    for i, table_data in enumerate(all_tables):
        if i <= 1 or i == len(all_tables) - 1:
            continue  # Skip header, teams, and approval tables
        result['process_tables'].append(table_data)

    # --- Extract purpose, scope, and process steps from paragraphs ---
    in_purpose = False
    in_scope = False
    in_process = False
    in_records = False

    def _strip_section_number(text: str) -> str:
        """Strip leading section numbers like '1.0 ', '6.0. ' from text."""
        return re.sub(r'^\d+(\.\d+)*\.?\s*', '', text)

    def _is_section(text_lower: str, keyword: str) -> bool:
        """Check if text starts with keyword, with or without section number prefix."""
        stripped = _strip_section_number(text_lower)
        return stripped.startswith(keyword)

    for para in paragraphs:
        para_lower = para.lower().strip()

        # Section detection (handles both "Purpose:" and "1.0 Purpose:" formats)
        if _is_section(para_lower, 'purpose'):
            in_purpose = True
            in_scope = False
            in_process = False
            in_records = False
            after = para.split(':', 1)
            if len(after) > 1 and after[1].strip():
                result['purpose'] = after[1].strip()
            continue
        elif _is_section(para_lower, 'scope'):
            in_purpose = False
            in_scope = True
            in_process = False
            in_records = False
            after = para.split(':', 1)
            if len(after) > 1 and after[1].strip():
                result['scope'] = after[1].strip()
            continue
        elif _is_section(para_lower, 'process') or _is_section(para_lower, 'procedure') or \
                _is_section(para_lower, 'test procedure') or \
                _is_section(para_lower, 'teams and responsibilities') or \
                _is_section(para_lower, 'equipment') or \
                _is_section(para_lower, 'applicable standard'):
            in_purpose = False
            in_scope = False
            in_process = True
            in_records = False
            continue
        elif _is_section(para_lower, 'records') or \
                re.match(r'^(\d+(\.\d+)*\.?\s*)?records?\s*[:.]?\s*$', para_lower) or \
                _is_section(para_lower, 'annexure'):
            in_purpose = False
            in_scope = False
            in_process = False
            in_records = True
            continue
        elif _is_section(para_lower, 'inspection') or _is_section(para_lower, 'acceptance') or \
                _is_section(para_lower, 're-test') or _is_section(para_lower, 'retest'):
            # Acceptance criteria and retest sections are still process-related
            in_purpose = False
            in_scope = False
            in_process = True
            in_records = False
            continue

        # Collect content
        if in_purpose and not result['purpose']:
            result['purpose'] = para
        elif in_scope and not result['scope']:
            result['scope'] = para
        elif in_process and para.strip():
            # Clean up numbered/bulleted steps
            cleaned = re.sub(r'^[\d]+[.\)]\s*', '', para.strip())
            cleaned = re.sub(r'^[a-zA-Z][.\)]\s*', '', cleaned)
            cleaned = cleaned.lstrip('- \u2022*')
            if cleaned and len(cleaned) > 3:
                result['process_steps'].append(cleaned.strip())

    return result


def parse_one_helix_docx(filepath: str) -> list:
    """Parse the comprehensive One_Helix document into multiple activities.

    This document has a unique structure with sections A.1-A.6, each containing
    an activity table with columns: Activity | Steps | Team | Interface/Utilities.
    """
    doc = Document(filepath)
    filename = os.path.basename(filepath)

    activities = []

    # Extract all tables
    all_tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            all_tables.append(rows)

    # Extract headings to identify sections
    headings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        if text and ('Heading' in style_name or text.startswith('A.')):
            headings.append(text)

    # Table 0: Responsibilities (Maker/Checker) - skip
    # Tables 1-6: Activity tables (A.1 HR Coil Purchase through A.6 Yield & Loss)
    # Table 7: Yield Metrics (formulas) - extract as data
    # Table 8: Yield Status/Action - extract as data

    section_names = [
        "HR Coil Purchase & Receipt",
        "Data Entry & Recording",
        "Slitting Process",
        "FG Conversion (Pipe Making)",
        "Sales Order & Dispatch",
        "Yield & Loss Tracking",
    ]

    # Map to journeys for the comprehensive doc
    section_journeys = {
        0: ("Pre-Production", "A. RAW MATERIAL MANAGEMENT"),
        1: ("Pre-Production", "A. RAW MATERIAL MANAGEMENT"),
        2: ("Production", "C. MANUFACTURING & FINISHING"),
        3: ("Production", "C. MANUFACTURING & FINISHING"),
        4: ("Post-Production", "F. DISPATCH & LOGISTICS"),
        5: ("Production", "E. QUALITY MANAGEMENT"),
    }

    for idx, table_data in enumerate(all_tables):
        if idx == 0:
            continue  # Skip responsibilities table
        if idx > 6:
            break  # Only process activity tables (1-6)

        activity_idx = idx - 1
        if activity_idx >= len(section_names):
            break

        header_lower = [h.lower() for h in table_data[0]] if table_data else []
        has_steps = any('step' in h for h in header_lower)
        has_activity = any('activity' in h for h in header_lower)

        if not (has_steps or has_activity):
            continue

        # Identify columns
        step_col = next((i for i, h in enumerate(header_lower) if 'step' in h), None)
        team_col = next((i for i, h in enumerate(header_lower) if 'team' in h), None)
        iface_col = next((i for i, h in enumerate(header_lower)
                          if 'interface' in h or 'utilit' in h), None)

        steps = []
        teams = set()
        interfaces = set()

        for row in table_data[1:]:
            if step_col is not None and step_col < len(row) and row[step_col].strip():
                steps.append(row[step_col].strip())
            if team_col is not None and team_col < len(row) and row[team_col].strip():
                teams.add(row[team_col].strip())
            if iface_col is not None and iface_col < len(row) and row[iface_col].strip():
                interfaces.add(row[iface_col].strip())

        journey_name, section_label = section_journeys.get(
            activity_idx, ("Production", "C. MANUFACTURING & FINISHING"))

        activity = Activity(
            activity_name=section_names[activity_idx],
            description=f"One Helix P&T Plant Operations - {section_names[activity_idx]}",
            owner='; '.join(sorted(teams)) if teams else "JOTS; Contract Manufacturer",
            interface='; '.join(sorted(interfaces)) if interfaces else "SAP/ERP",
            sign_off="JSW One Operational Head",
            steps=steps,
            flow_type="Positive",
            sop_link=None,
            remarks=None,
            source_document=filename,
        )
        activity._journey = journey_name
        activity._section = section_label
        activities.append(activity)

    # Extract yield metrics table (table 7) as a separate activity
    if len(all_tables) > 7:
        yield_table = all_tables[7]
        yield_steps = []
        for row in yield_table[1:]:
            if len(row) >= 2 and row[0].strip():
                yield_steps.append(f"{row[0].strip()}: {row[1].strip()}")

        if yield_steps:
            act = Activity(
                activity_name="Yield Metrics & Formulas",
                description="Yield calculation formulas for P&T operations",
                owner="Plant Operations Manager",
                interface="SAP/ERP; Excel",
                sign_off="JSW One Operational Head",
                steps=yield_steps,
                flow_type="Positive",
                sop_link=None,
                remarks=None,
                source_document=filename,
            )
            act._journey = "Production"
            act._section = "E. QUALITY MANAGEMENT"
            activities.append(act)

    # Extract yield status/action table (table 8)
    if len(all_tables) > 8:
        status_table = all_tables[8]
        status_steps = []
        for row in status_table[1:]:
            parts = [cell for cell in row if cell.strip()]
            if parts:
                status_steps.append(' | '.join(parts))

        if status_steps:
            act = Activity(
                activity_name="Yield Status & Corrective Actions",
                description="Yield range thresholds and required corrective actions",
                owner="Plant Operations Manager; QA Head",
                interface="SAP/ERP",
                sign_off="JSW One Operational Head",
                steps=status_steps,
                flow_type="Positive",
                sop_link=None,
                remarks=None,
                source_document=filename,
            )
            act._journey = "Production"
            act._section = "E. QUALITY MANAGEMENT"
            activities.append(act)

    return activities


# ---------------------------------------------------------------------------
# Activity extractor
# ---------------------------------------------------------------------------

def extract_pt_activities(parsed: dict, config: dict) -> list:
    """Convert parsed P&T document data into Activity objects."""
    activities = []

    prefix = config.get('activity_prefix', parsed['title'])

    # Build owner from teams
    owner_parts = []
    for t in parsed['teams']:
        owner_parts.append(t['team'])
    owner = '; '.join(owner_parts) if owner_parts else "Contract Manufacturer; JOTS"

    # Build sign-off from approval info
    sign_off = ""
    if parsed['approval_info']:
        # Extract approver name if present
        approval_lower = parsed['approval_info'].lower()
        if 'approved' in approval_lower:
            sign_off = parsed['approval_info'][:100]
    if not sign_off:
        sign_off = "JSW One Operation Manager"

    # Detect interface from text
    interface = _detect_pt_interface(parsed['raw_text'])

    # Build description from purpose
    description = parsed['purpose'] or f"SOP for {prefix} - Pipes & Tubes"

    # Collect steps from process_steps
    steps = parsed['process_steps'][:20]

    # If no steps from text, try extracting from process tables
    if not steps and parsed['process_tables']:
        for table_data in parsed['process_tables']:
            if not table_data or len(table_data) < 2:
                continue
            header_lower = [h.lower() for h in table_data[0]]

            # Look for tables with descriptive content
            for row in table_data[1:]:
                for cell in row:
                    if cell.strip() and len(cell.strip()) > 10:
                        steps.append(cell.strip())
                if len(steps) >= 15:
                    break

    # Build remarks from process tables (e.g., acceptance criteria)
    remarks = None
    remarks_details = []
    for table_data in parsed['process_tables']:
        if not table_data or len(table_data) < 2:
            continue
        header_lower = [h.lower() for h in table_data[0]]

        # Check for criteria/specification tables
        has_criteria = any(kw in h for h in header_lower
                          for kw in ['criterion', 'acceptance', 'characteristic',
                                     'nonconform', 'rejection', 'disposition'])
        if has_criteria:
            remarks = f"See {parsed['filename']} for detailed acceptance criteria"
            for row in table_data[1:]:
                row_text = ' | '.join(cell for cell in row if cell.strip())
                if row_text:
                    remarks_details.append(row_text)

    # Build teams into description_details
    desc_details = []
    if parsed['teams']:
        for t in parsed['teams']:
            if t['responsibility']:
                desc_details.append(f"{t['team']}: {t['responsibility']}")

    activity = Activity(
        activity_name=prefix,
        description=description,
        description_details=desc_details,
        owner=owner,
        interface=interface,
        sign_off=sign_off,
        steps=steps,
        flow_type="Positive",
        sop_link=None,  # No SOP links for P&T yet
        remarks=remarks,
        remarks_details=remarks_details[:10],
        source_document=parsed['filename'],
    )

    activities.append(activity)
    return activities


def _detect_pt_interface(raw_text: str) -> str:
    """Detect interfaces/systems from P&T document content."""
    text = raw_text.lower()
    interfaces = []

    if 'sap' in text or 'erp' in text:
        interfaces.append("SAP/ERP")
    if 'jots' in text or 'jsw one' in text:
        interfaces.append("JOTS")
    if 'excel' in text or 'spreadsheet' in text:
        interfaces.append("Excel")
    if 'email' in text:
        interfaces.append("Email")
    if 'utm' in text or 'universal testing' in text:
        interfaces.append("UTM")
    if 'vernier' in text or 'caliper' in text or 'tape' in text:
        interfaces.append("Measuring Instruments")

    return '; '.join(interfaces[:3]) if interfaces else "JOTS; SAP/ERP"


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def main():
    project_root = os.path.dirname(TOOLS_DIR)
    docs_dir = os.path.join(project_root, 'Documents', 'SOP_Pipes & tubes')
    output_dir = os.path.join(project_root, 'output')

    print(f"=" * 60)
    print(f"  JSW ONE Pipes & Tubes SOP Master Excel Generator")
    print(f"=" * 60)

    # Step 1: Discover documents
    print(f"\n[1/5] Discovering documents in: {docs_dir}")
    doc_files = discover_pt_documents(docs_dir)
    print(f"  Found {len(doc_files)} unique documents")

    # Step 2: Parse and classify
    print(f"\n[2/5] Parsing and classifying documents...")
    all_activities = []  # List of (journey, section, [Activity])
    seen_titles = set()

    for filepath in doc_files:
        filename = os.path.basename(filepath)
        print(f"  Processing: {filename}")

        config = classify_pt_document(filename)
        if config is None:
            print(f"    WARNING: No classification found, skipping")
            continue

        print(f"    -> {config['journey']} / {config['section']}")

        # Handle One_Helix comprehensive doc specially
        if config.get('is_comprehensive'):
            try:
                helix_activities = parse_one_helix_docx(filepath)
                for act in helix_activities:
                    journey = getattr(act, '_journey', config['journey'])
                    section = getattr(act, '_section', config['section'])
                    all_activities.append((journey, section, [act]))
                print(f"    Extracted {len(helix_activities)} activities (comprehensive doc)")
            except Exception as e:
                print(f"    ERROR: {e}")
                import traceback
                traceback.print_exc()
            continue

        # Standard P&T document parsing
        try:
            parsed = parse_pt_docx(filepath)

            # Content-based dedup
            dedup_key = parsed['title'].strip().lower()
            if dedup_key in seen_titles:
                print(f"    Skipping content duplicate")
                continue
            seen_titles.add(dedup_key)

            activities = extract_pt_activities(parsed, config)
            journey = config['journey']
            section = config['section']
            all_activities.append((journey, section, activities))
            print(f"    Extracted {len(activities)} activity(ies)")
            if parsed['teams']:
                print(f"    Teams: {', '.join(t['team'] for t in parsed['teams'][:3])}")

        except Exception as e:
            print(f"    ERROR: {e}")
            import traceback
            traceback.print_exc()

    # Step 3: Build journey structure
    print(f"\n[3/5] Building journey structure...")
    journeys = build_pt_journey_structure(all_activities)

    total_activities = sum(
        len(section.activities)
        for journey in journeys
        for section in journey.sections
    )
    print(f"  Total journeys: {len(journeys)}")
    for j in journeys:
        section_count = len(j.sections)
        act_count = sum(len(s.activities) for s in j.sections)
        print(f"    {j.name}: {section_count} sections, {act_count} activities")
    print(f"  Total activities: {total_activities}")

    # Step 4: Generate Excel
    print(f"\n[4/5] Generating Excel file...")
    output_path = os.path.join(output_dir, 'JSW_ONE_PT_SOPs_Master.xlsx')
    write_master_excel(journeys, output_path, sheet_name="Pipes & Tubes")
    print(f"  Generated: {output_path}")

    # Step 5: Save manifest
    print(f"\n[5/5] Saving manifest...")
    manifest_path = os.path.join(output_dir, 'pt_manifest.json')
    save_pt_manifest(manifest_path, doc_files)
    print(f"  Manifest saved: {manifest_path}")

    print(f"\n{'=' * 60}")
    print(f"  Done! Master Excel generated at: {output_path}")
    print(f"{'=' * 60}")
    return output_path


def discover_pt_documents(docs_dir: str) -> list:
    """Find all .docx files in the P&T folder, deduplicate by content hash."""
    files = []
    seen_hashes = set()

    if not os.path.isdir(docs_dir):
        print(f"  ERROR: Directory not found: {docs_dir}")
        return files

    for f in sorted(Path(docs_dir).glob('*.docx')):
        if not f.is_file():
            continue
        h = hashlib.md5(f.read_bytes()).hexdigest()
        if h not in seen_hashes:
            seen_hashes.add(h)
            files.append(str(f))
        else:
            print(f"  Skipping duplicate: {f.name}")

    return files


def build_pt_journey_structure(all_activities: list) -> list:
    """Build Journey > Section > Activity hierarchy for P&T."""
    structure = defaultdict(lambda: defaultdict(list))

    for journey_name, section_label, activities in all_activities:
        structure[journey_name][section_label].extend(activities)

    journeys = []
    sr_no_counter = 1

    for journey_name in PT_JOURNEY_ORDER:
        if journey_name not in structure:
            continue

        sections = []
        section_data = structure[journey_name]

        # Order sections from PT_SECTION_ORDER
        ordered_labels = PT_SECTION_ORDER.get(journey_name, [])
        seen_labels = set()

        for label in ordered_labels:
            if label in section_data:
                seen_labels.add(label)
                activities = section_data[label]
                for act in activities:
                    act.sr_no = sr_no_counter
                    sr_no_counter += 1
                sections.append(Section(label=label, activities=activities))

        # Any remaining sections not in predefined order
        for label, activities in section_data.items():
            if label not in seen_labels:
                for act in activities:
                    act.sr_no = sr_no_counter
                    sr_no_counter += 1
                sections.append(Section(label=label, activities=activities))

        if sections:
            journeys.append(Journey(name=journey_name, sections=sections))

    return journeys


def save_pt_manifest(manifest_path: str, doc_files: list):
    """Save a build manifest."""
    manifest = {
        "generated_at": datetime.now().isoformat(),
        "generator": "tools/generate_pt_sop_excel.py",
        "product_line": "Pipes & Tubes",
        "files": {},
    }

    for filepath in doc_files:
        h = hashlib.md5(Path(filepath).read_bytes()).hexdigest()
        manifest["files"][os.path.basename(filepath)] = {
            "hash": h,
            "size": os.path.getsize(filepath),
        }

    os.makedirs(os.path.dirname(manifest_path) or '.', exist_ok=True)
    with open(manifest_path, 'w') as f:
        json.dump(manifest, f, indent=2)


if __name__ == '__main__':
    main()
