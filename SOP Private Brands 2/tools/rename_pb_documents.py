"""
Private Brands SOP Document Renamer
Renames PB SOP documents to [XX] SOP_CleanName format, sequenced by Journey > Section order.
Attempts version update to 1.0 in .docx files where possible.

Usage:
    python tools/rename_pb_documents.py [--dry-run]
"""

import os
import sys
import argparse
from pathlib import Path

SOP_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Documents", "SOP's"
)

# Complete file mapping: old_filename -> {doc_num, clean_name}
# Sequenced by Journey (Pre-Order -> Order -> Post Order) then Section
RENAME_MAPPING = {
    # ── Pre - Order ──────────────────────────────────────────────────
    # A. DATA MANAGEMENT
    "Influencer,+Retailer+&+Distributor_+Data+Enrichment (1).doc": {
        "doc_num": 1,
        "clean_name": "SOP_Data Enrichment Specification",
    },
    "PB+-+Retailer+&+Influencer+Data+Enrichment+-+Demo (1).doc": {
        "doc_num": 2,
        "clean_name": "SOP_PB Retailer & Influencer Data Enrichment Demo",
    },
    "Updates+to+Add+Category+Component+on+the+PB+Retailer+&+Influencer.doc": {
        "doc_num": 3,
        "clean_name": "SOP_Category Component Updates",
    },
    # E. ORDER LOGGING (spans Pre-Order / Order / Post Order)
    "JSWOrderLogging_V18.docx": {
        "doc_num": 4,
        "clean_name": "SOP_JSW Order Logging",
        "try_version_update": True,
    },
    # G. INFLUENCER MANAGEMENT
    "3-+PB+Influencer+meets+bill+clearance+(mason_contractor_dealer)-+Ops+process.doc": {
        "doc_num": 5,
        "clean_name": "SOP_PB Influencer Meets Bill Clearance",
    },
    "4-+Influencer+data+management+SOP+-+Private+Brands (1).doc": {
        "doc_num": 6,
        "clean_name": "SOP_Influencer Data Management",
    },
    "5-+PB+Influencer+scheme+disbursal+and+compliance-+Ops+SOP.doc": {
        "doc_num": 7,
        "clean_name": "SOP_PB Influencer Scheme Disbursal & Compliance",
    },
    "New SOP for Influencer Meets.pdf": {
        "doc_num": 8,
        "clean_name": "SOP_Influencer Meets",
    },
    # K. MARKETING ACTIVITIES
    "GSB_Dealer_Sign_Board.docx": {
        "doc_num": 9,
        "clean_name": "SOP_Dealer Sign Board GSB",
        "try_version_update": True,
    },
    "Wall_Painting.docx": {
        "doc_num": 10,
        "clean_name": "SOP_Wall Painting",
        "try_version_update": True,
    },
    "SOP Marketing Activities.xlsx": {
        "doc_num": 11,
        "clean_name": "SOP_Marketing Activities",
    },
    # ── Order ────────────────────────────────────────────────────────
    # C. OPPORTUNITY MANAGEMENT
    "Documentation_+Document+the+PB+Opportunity+Workflow.doc": {
        "doc_num": 12,
        "clean_name": "SOP_PB Opportunity Workflow",
    },
    # H. APPROVAL WORKFLOWS
    "Documentation+of+Approval+Workflows+that+are+in+PB.doc": {
        "doc_num": 13,
        "clean_name": "SOP_PB Approval Workflows",
    },
    # I. DIGITAL CHANNELS
    "PB+Retailer_Influencer+in+App.doc": {
        "doc_num": 14,
        "clean_name": "SOP_PB Retailer Influencer in App",
    },
    "PB+Retailer_Influencer+in+App (1).doc": {
        "doc_num": 15,
        "clean_name": "SOP_PB Retailer Influencer in App (1)",
    },
    # J. QUALITY ASSURANCE (spans Order + Post Order)
    "JSW One TMT Quality Manual.pdf": {
        "doc_num": 16,
        "clean_name": "SOP_JSW One TMT Quality Manual",
    },
    # ── Post Order ───────────────────────────────────────────────────
    # F. FINANCE & RECONCILIATION
    "1-+PB+Rebates,+Schemes+and+Price+difference-+Ops+process+SOP.doc": {
        "doc_num": 17,
        "clean_name": "SOP_PB Rebates Schemes & Price Difference",
    },
}

# File excluded from renaming (reference file, not an SOP)
EXCLUDED_FILES = ["PB SOP Links Confluence.xlsx"]


def build_new_filename(doc_num, clean_name, original_name):
    """Build filename in [XX] SOP_Name format, preserving original extension."""
    ext = os.path.splitext(original_name)[1]
    return f"[{doc_num:02d}] {clean_name}{ext}"


def try_update_version_docx(filepath, version_str="1.0"):
    """Attempt to update version cell in a .docx file's first table.

    Returns (success, message). Non-fatal — warns and continues on failure.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return False, "python-docx not installed"

    try:
        doc = Document(filepath)
    except Exception as e:
        return False, f"Cannot open: {e}"

    if not doc.tables:
        return False, "No tables found"

    # Search first few tables for a version cell
    for table_idx, table in enumerate(doc.tables[:3]):
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text.lower().startswith("version"):
                    # Found version cell — update it
                    original_font_size = None
                    para = cell.paragraphs[0]
                    if para.runs:
                        original_font_size = para.runs[0].font.size

                    # Clear paragraph content (preserve properties)
                    p_element = para._element
                    for child in list(p_element):
                        if child.tag != qn('w:pPr'):
                            p_element.remove(child)

                    new_run = para.add_run(f"Version: {version_str} ")
                    if original_font_size:
                        new_run.font.size = original_font_size

                    doc.save(filepath)
                    return True, f"'{cell_text}' -> 'Version: {version_str}'"

    return False, "No version cell found in tables"


def main():
    parser = argparse.ArgumentParser(description="Rename PB SOP documents")
    parser.add_argument("--dry-run", action="store_true",
                        help="Preview changes without modifying files")
    args = parser.parse_args()

    sop_dir = SOP_DIR
    if not os.path.isdir(sop_dir):
        print(f"ERROR: Directory not found: {sop_dir}")
        sys.exit(1)

    prefix = "[DRY RUN] " if args.dry_run else ""
    print(f"{prefix}Private Brands SOP Document Renamer")
    print(f"Directory: {sop_dir}")
    print("=" * 70)

    # --- Phase 1: Validate source files ---
    print("\n[Phase 1] Validating source files...")
    errors = []
    for old_name in RENAME_MAPPING:
        old_path = os.path.join(sop_dir, old_name)
        if not os.path.isfile(old_path):
            errors.append(f"  MISSING: {old_name}")

    if errors:
        for e in errors:
            print(e)
        print("ABORT: Fix missing files before proceeding.")
        sys.exit(1)
    print(f"  All {len(RENAME_MAPPING)} source files found.")

    # Note excluded files
    for exc in EXCLUDED_FILES:
        exc_path = os.path.join(sop_dir, exc)
        status = "present" if os.path.isfile(exc_path) else "not found"
        print(f"  Excluded ({status}): {exc}")

    # --- Phase 2: Check for collisions ---
    print("\n[Phase 2] Checking for collisions...")
    new_names = {}
    for old_name, config in RENAME_MAPPING.items():
        new_name = build_new_filename(config["doc_num"], config["clean_name"], old_name)
        if new_name in new_names:
            errors.append(f"  COLLISION: {new_name}")
        new_names[new_name] = old_name

    if errors:
        for e in errors:
            print(e)
        print("ABORT: Resolve collisions.")
        sys.exit(1)
    print("  No collisions.")

    # --- Phase 3: Version updates (.docx only) ---
    print("\n[Phase 3] Updating versions in .docx files...")
    sorted_items = sorted(RENAME_MAPPING.items(), key=lambda x: x[1]["doc_num"])

    for old_name, config in sorted_items:
        if not config.get("try_version_update", False):
            continue

        old_path = os.path.join(sop_dir, old_name)
        if args.dry_run:
            print(f"  [{config['doc_num']:02d}] Would try version update: {old_name}")
        else:
            success, msg = try_update_version_docx(old_path, "1.0")
            status = "OK" if success else "SKIP"
            print(f"  [{config['doc_num']:02d}] {status}: {msg}")

    # --- Phase 4: Rename files (two-pass) ---
    print("\n[Phase 4] Renaming files...")
    temp_pairs = []

    for old_name, config in sorted_items:
        new_name = build_new_filename(config["doc_num"], config["clean_name"], old_name)
        old_path = os.path.join(sop_dir, old_name)
        new_path = os.path.join(sop_dir, new_name)

        print(f"  [{config['doc_num']:02d}] {old_name}")
        print(f"       -> {new_name}")

        if old_name == new_name:
            print("       (same name, skip)")
            continue

        if args.dry_run:
            continue

        # Two-pass: old -> temp -> new
        temp_path = old_path + ".renaming"
        try:
            os.rename(old_path, temp_path)
            temp_pairs.append((temp_path, new_path, new_name))
        except OSError as e:
            print(f"       ERROR: {e}")
            if os.path.isfile(temp_path) and not os.path.isfile(old_path):
                os.rename(temp_path, old_path)

    # Finalize temp -> new
    for temp_path, new_path, new_name in temp_pairs:
        try:
            os.rename(temp_path, new_path)
        except OSError as e:
            print(f"       ERROR finalizing {new_name}: {e}")

    # --- Phase 5: Verify ---
    print(f"\n{'=' * 70}")
    print("[Phase 5] Verification:")

    if args.dry_run:
        print("  Dry run complete. No files were modified.")
    else:
        expected = {build_new_filename(c["doc_num"], c["clean_name"], n)
                    for n, c in RENAME_MAPPING.items()}
        # Also include excluded files in expected set
        for exc in EXCLUDED_FILES:
            expected.add(exc)

        actual = set(os.listdir(sop_dir))

        renamed_present = expected & actual
        renamed_missing = {f for f in expected if f not in actual} - set(EXCLUDED_FILES)
        unexpected = actual - expected

        print(f"  Renamed files present: {len(renamed_present - set(EXCLUDED_FILES))}/{len(RENAME_MAPPING)}")
        if renamed_missing:
            print(f"  MISSING: {renamed_missing}")
        if unexpected:
            print(f"  Other files in folder: {unexpected}")
        if not renamed_missing:
            print("  All renamed files verified.")

    print("\nDone.")


if __name__ == "__main__":
    main()
