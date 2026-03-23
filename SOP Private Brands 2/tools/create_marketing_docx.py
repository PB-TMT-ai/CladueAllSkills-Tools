"""
One-time script to create .docx files from marketing specification images.
Creates GSB_Dealer_Sign_Board.docx and Wall_Painting.docx in Documents/SOP's/
with tables matching the OrderLogging_V18 format.
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(PROJECT_ROOT, "Documents", "SOP's")


def create_gsb_docx():
    """Create GSB_Dealer_Sign_Board.docx from GSB.png content."""
    doc = Document()

    # Title
    title = doc.add_heading("Dealer Sign Board - GSB", level=1)

    # Specifications section
    doc.add_heading("Specifications", level=2)
    spec_table = doc.add_table(rows=7, cols=2)
    spec_table.style = "Table Grid"
    # Header
    spec_table.rows[0].cells[0].text = "Item"
    spec_table.rows[0].cells[1].text = "Detail"
    # Data
    specs = [
        ("Material", "Flex / Rikasons/Star 400 GSM"),
        ("Tube Light / Module", "Surya/Anchor/Heads"),
        ("Cabinet Frame", '150 Mock 16 gauge G.I. Sheet & 20 gauge 1" ms pipe'),
        ("Box Side Sheet", "Powder Coated"),
        ("Box Thickness", "8"),
        ("Iron Angle", "35 x 35 x 5mm to be used (to mount the GSB in lesser space available at the dealer store)"),
    ]
    for i, (item, detail) in enumerate(specs, 1):
        spec_table.rows[i].cells[0].text = item
        spec_table.rows[i].cells[1].text = detail

    # Approval and Execution section
    doc.add_heading("Approval and Execution", level=2)
    exec_table = doc.add_table(rows=4, cols=5)
    exec_table.style = "Table Grid"
    # Header
    headers = ["Activity", "Steps", "Team", "Interface / Utilities used", "Sign off"]
    for j, h in enumerate(headers):
        exec_table.rows[0].cells[j].text = h
    # Data rows
    rows_data = [
        ("1.", "Channel team to provide list of dealer that are needed to be covered along with the dealer print name, address and dimensions.", "Channel Team", "WhatsApp / Field", "Marketing"),
        ("2.", "Recce to be performed by the vendor, further to which report will be shared the channel team for validation.", "Vendor", "WhatsApp / Field", "Channel Team"),
        ("3.", "Post validation, execution will be started, WhatsApp group will be created to track the placing and recording the execution.", "Marketing / Vendor", "WhatsApp", "Marketing"),
    ]
    for i, row_data in enumerate(rows_data, 1):
        for j, val in enumerate(row_data):
            exec_table.rows[i].cells[j].text = val

    filepath = os.path.join(DOCS_DIR, "GSB_Dealer_Sign_Board.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_wall_painting_docx():
    """Create Wall_Painting.docx from Wall Painting (1).png content."""
    doc = Document()

    # Title
    doc.add_heading("Wall Painting", level=1)

    # Specifications section
    doc.add_heading("Specifications", level=2)
    spec_table = doc.add_table(rows=5, cols=2)
    spec_table.style = "Table Grid"
    spec_table.rows[0].cells[0].text = "Item"
    spec_table.rows[0].cells[1].text = "Detail"
    specs = [
        ("Paint Type", "Asian Paints/Burger/Water based"),
        ("Branding Rule", "In case there is not a separate patch with dealer branding name and IVR number to be mandate at all walls (exception only if walls are less than 100 sq ft)"),
        ("Adjacent Wall Rule", "In case of dealer adjacent wall, dealer can be allowed to use retailer/his/her current information, but IVR information should not be tampered at all conditions."),
        ("Artwork", "Black strip needed to be added below the artwork for respective RM/TM contact/dealer/distributor name (dealer/distributor will not be written). No alteration on the wall painting artwork is allowed."),
    ]
    for i, (item, detail) in enumerate(specs, 1):
        spec_table.rows[i].cells[0].text = item
        spec_table.rows[i].cells[1].text = detail

    # Approval and Execution section
    doc.add_heading("Approval and Execution", level=2)
    exec_table = doc.add_table(rows=7, cols=5)
    exec_table.style = "Table Grid"
    headers = ["Activity", "Steps", "Team", "Interface / Utilities used", "Sign off"]
    for j, h in enumerate(headers):
        exec_table.rows[0].cells[j].text = h
    rows_data = [
        ("1.", "WhatsApp groups will be created by the vendor for the respective regions.", "Vendor", "WhatsApp", "Marketing"),
        ("2.", "Respective RMs and SMs are needed to be added in the group either by vendor or by marketing.", "Marketing / Vendor", "WhatsApp", "Marketing"),
        ("3.", "Painters are needed to be added in the group as most, who will report through photos for approval before starting the wall painting.", "Painter", "WhatsApp / Geo-tagged photos", "RM/SM"),
        ("4.", "Photo shared in the group should be geo tagged and each wall photo should come in set of 6 photos (including before photo).", "Painter", "WhatsApp / Geo-tagged photos", "RM/SM"),
        ("5.", "No alteration on the wall painting artwork is allowed.", "-", "-", "Marketing"),
        ("6.", "RM/TM/RM will approve the wall, no walls should be painted without the approval.", "RM/TM", "WhatsApp", "Marketing"),
    ]
    for i, row_data in enumerate(rows_data, 1):
        for j, val in enumerate(row_data):
            exec_table.rows[i].cells[j].text = val

    filepath = os.path.join(DOCS_DIR, "Wall_Painting.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


if __name__ == "__main__":
    create_gsb_docx()
    create_wall_painting_docx()
    print("Done! Both .docx files created.")
