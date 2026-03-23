"""
Parser for .docx files (specifically JSWOrderLogging_V18.docx).
Extracts structured activity data from tables organized by journey phases.
"""

import os
import sys
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from models.sop_data import SOPDocument


def parse_docx_sop(filepath: str) -> SOPDocument:
    """Parse a .docx SOP document and return a structured SOPDocument."""
    doc = Document(filepath)
    filename = os.path.basename(filepath)

    # Extract title from first heading or Title style
    title = ""
    headings = []
    paragraphs_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if 'Title' in style_name and not title:
            title = text
        elif 'Heading' in style_name:
            level = 1
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break
            headings.append((level, text))
            if not title and level == 1:
                title = text

        paragraphs_text.append(text)

    if not title:
        title = filename.replace('.docx', '').replace('_', ' ')

    # Extract tables with their content
    tables_data = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(cells)
        if table_rows:
            tables_data.append(table_rows)

    # Build raw text
    raw_text = '\n'.join(paragraphs_text)

    # Extract steps from tables that have Activity/Steps columns
    all_steps = []
    for table_data in tables_data:
        if not table_data:
            continue
        header = [h.lower() for h in table_data[0]]
        step_col = None
        for i, h in enumerate(header):
            if 'step' in h:
                step_col = i
                break
        if step_col is not None:
            for row in table_data[1:]:
                if step_col < len(row) and row[step_col].strip():
                    all_steps.append(row[step_col].strip())

    return SOPDocument(
        filename=filename,
        title=title,
        doc_type="",  # Will be set by document_classifier
        purpose=None,
        stakeholders={},
        steps=all_steps,
        tables=tables_data,
        headings=headings,
        raw_text=raw_text,
        sign_off_info=None,
        escalation_info=None,
    )


def extract_docx_activities(filepath: str) -> list:
    """Extract detailed activity records from JSWOrderLogging_V18.docx.

    Returns a list of dicts with keys:
        phase, section, activity_num, activity_name, steps, team, interface, sign_off
    """
    doc = Document(filepath)

    # Map heading context to understand which phase/section we're in
    current_phase = "Pre - Order"
    current_section = ""
    heading_context = []

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        if 'Heading' in style_name and text:
            text_upper = text.upper()
            if 'PRE-ORDER' in text_upper or 'PRE ORDER' in text_upper:
                current_phase = "Pre - Order"
            elif 'POST-ORDER' in text_upper or 'POST ORDER' in text_upper:
                current_phase = "Post Order"
            elif 'ORDER' in text_upper and 'PRE' not in text_upper and 'POST' not in text_upper:
                current_phase = "Order"
            heading_context.append((current_phase, text))

    # Process tables - each table represents activities within a section
    activities = []
    phase_for_table = _map_tables_to_phases(doc)

    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if len(rows) < 2:
            continue

        # Check if this is an activity table (has Activity/Steps/Team columns)
        header = rows[0]
        header_lower = [h.lower() for h in header]

        col_map = _identify_columns(header_lower)
        if not col_map:
            continue

        phase = phase_for_table.get(table_idx, "Pre - Order")

        for row in rows[1:]:
            if all(not cell.strip() for cell in row):
                continue

            activity = {
                'phase': phase,
                'section': '',
                'activity_num': _safe_get(row, col_map.get('activity', -1), ''),
                'steps': _safe_get(row, col_map.get('steps', -1), ''),
                'team': _safe_get(row, col_map.get('team', -1), ''),
                'interface': _safe_get(row, col_map.get('interface', -1), ''),
                'sign_off': _safe_get(row, col_map.get('signoff', -1), ''),
            }
            activities.append(activity)

    return activities


def _identify_columns(header_lower: list) -> dict | None:
    """Identify column indices from header row. Returns None if not an activity table."""
    col_map = {}

    for i, h in enumerate(header_lower):
        if 'activity' in h and 'activity' not in col_map:
            col_map['activity'] = i
        elif 'step' in h:
            col_map['steps'] = i
        elif 'team' in h:
            col_map['team'] = i
        elif 'interface' in h or 'utilit' in h:
            col_map['interface'] = i
        elif 'sign' in h:
            col_map['signoff'] = i

    # Must have at least activity or steps column to be relevant
    if 'activity' not in col_map and 'steps' not in col_map:
        return None

    return col_map


def _map_tables_to_phases(doc) -> dict:
    """Map table indices to journey phases based on surrounding headings."""
    phase_map = {}
    current_phase = "Pre - Order"
    table_idx = 0

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Check if it's a heading
            for child in element:
                child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if child_tag == 'pPr':
                    for style_elem in child:
                        style_tag = style_elem.tag.split('}')[-1] if '}' in style_elem.tag else style_elem.tag
                        if style_tag == 'pStyle':
                            style_val = style_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                            if 'Heading' in style_val:
                                text = element.text or ''
                                # Also get text from runs
                                for run in element.iter():
                                    run_tag = run.tag.split('}')[-1] if '}' in run.tag else run.tag
                                    if run_tag == 't' and run.text:
                                        text += run.text

                                text_upper = text.upper()
                                # Check for phase markers (must be specific phase headings, not title)
                                if 'PRE-ORDER' in text_upper or 'PRE ORDER' in text_upper:
                                    current_phase = "Pre - Order"
                                elif 'POST-ORDER' in text_upper or 'POST ORDER' in text_upper:
                                    current_phase = "Post Order"
                                elif ('ORDER PHASE' in text_upper or text_upper.strip() == 'ORDER') and 'PRE' not in text_upper and 'POST' not in text_upper and 'LOGGING' not in text_upper:
                                    current_phase = "Order"

        elif tag == 'tbl':
            phase_map[table_idx] = current_phase
            table_idx += 1

    return phase_map


def _safe_get(lst: list, index: int, default: str = '') -> str:
    """Safely get an element from a list."""
    if index < 0 or index >= len(lst):
        return default
    return lst[index].strip() if lst[index] else default
