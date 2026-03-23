"""
Script to update SOP document v1.1 to v1.2 based on Process Flow changes.
Applies 9 identified changes.
"""
import copy
from docx import Document
from docx.shared import Pt

INPUT_PATH = r'D:\SOP_OneHelixPipe\One_Helix_Pipes_Tubes_SOP_REVISED_FINAL 1.1.docx'
OUTPUT_PATH = r'D:\SOP_OneHelixPipe\One_Helix_Pipes_Tubes_SOP_REVISED_FINAL 1.2.docx'

doc = Document(INPUT_PATH)

# Helper: replace text in a paragraph while preserving formatting
def replace_paragraph_text(paragraph, old_text, new_text):
    """Replace text in paragraph, preserving run formatting."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False
    # For simple single-run paragraphs
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = paragraph.runs[0].text.replace(old_text, new_text)
        return True
    # For multi-run: reconstruct
    combined = ''.join(r.text for r in paragraph.runs)
    if old_text in combined:
        new_combined = combined.replace(old_text, new_text)
        # Put all text in first run, clear others
        if paragraph.runs:
            paragraph.runs[0].text = new_combined
            for r in paragraph.runs[1:]:
                r.text = ''
        return True
    return False


def replace_cell_text(cell, old_text, new_text):
    """Replace text in a table cell, preserving formatting."""
    for paragraph in cell.paragraphs:
        if old_text in paragraph.text:
            replace_paragraph_text(paragraph, old_text, new_text)
            return True
    return False


# ============================================================
# CHANGE 1: Fix typo "Receives" -> "Receipt" in Section A.1 heading (P[20])
# ============================================================
print("Change 1: Fix typo in A.1 heading...")
for p in doc.paragraphs:
    if 'HR Coil Purchase & Receives' in p.text:
        replace_paragraph_text(p, 'HR Coil Purchase & Receives', 'HR Coil Purchase & Receipt')
        print(f"  Fixed: '{p.text}'")
        break

# ============================================================
# CHANGE 2: Add "approved suppliers" reference in Section A.1
# ============================================================
print("Change 2: Add approved suppliers reference...")
for p in doc.paragraphs:
    if p.text.strip() == 'This section describes the process for receiving and recording Hot Rolled (HR) coils at the One Helix Pipes & Tubes manufacturing facility.':
        replace_paragraph_text(
            p,
            'This section describes the process for receiving and recording Hot Rolled (HR) coils at the One Helix Pipes & Tubes manufacturing facility.',
            'This section describes the process for receiving and recording Hot Rolled (HR) coils procured from approved suppliers at the One Helix Pipes & Tubes manufacturing facility. Upon receipt, all coil details are recorded in the Master Sheet (Coil sheet).'
        )
        print(f"  Updated section description")
        break

# Also update table 0 step 1 to mention approved suppliers
table0 = doc.tables[0]
steps_cell = table0.rows[1].cells[1]
for para in steps_cell.paragraphs:
    if 'Receive HR coil delivery at plant facility' in para.text:
        replace_paragraph_text(
            para,
            'Receive HR coil delivery at plant facility',
            'Receive HR coil delivery from approved supplier at plant facility'
        )
        print("  Updated table step")
        break

# ============================================================
# CHANGE 3: Add "Master Sheet" reference - already done in Change 2
# The section description now mentions "Master Sheet (Coil sheet)"
# ============================================================
print("Change 3: Master Sheet reference added in Change 2")

# ============================================================
# CHANGE 4: Revise Section A.2 to reflect dropdown & auto-population
# ============================================================
print("Change 4: Update A.2 workflow description...")
for p in doc.paragraphs:
    if p.text.startswith('This section describes the process for recording detailed information of each HR coil in the tracking system.'):
        replace_paragraph_text(
            p,
            'This section describes the process for recording detailed information of each HR coil in the tracking system. All data entered in this step forms the master record that will be referenced throughout the manufacturing process.',
            'This section describes the process for recording detailed information of each HR coil in the Master Sheet (Coil sheet). All data entered in this step forms the master record. For tracking purposes, the plant team selects the HR Coil ID from a dropdown list linked to the Master Sheet, and the coil details are automatically populated in the HR Coil Tracking Sheet to ensure traceability throughout the manufacturing process.'
        )
        print("  Updated A.2 description")
        break

# ============================================================
# CHANGE 5: Add parameter-based filtering in Section A.3 (Slitting)
# Update table 2, step about HR Coil ID selection
# ============================================================
print("Change 5: Add parameter-based filtering in A.3...")
table2 = doc.tables[2]
steps_cell_a3 = table2.rows[1].cells[1]
for para in steps_cell_a3.paragraphs:
    if 'Select HR Coil ID from dropdown' in para.text:
        replace_paragraph_text(
            para,
            'Select HR Coil ID from dropdown (auto-populates from "Coil sheet")',
            'Enter Thickness and Width to be slit. Based on the entered parameters, the relevant HR Coil ID appears in a filtered dropdown. Select the applicable HR Coil ID'
        )
        print("  Updated slitting dropdown step")
        break

# ============================================================
# CHANGE 6: Add "Date of Inward" to baby coil fields in A.3
# ============================================================
print("Change 6: Add Date of Inward to baby coil fields...")
for para in steps_cell_a3.paragraphs:
    if 'Cost Price and PO Number auto-populate from master data' in para.text:
        replace_paragraph_text(
            para,
            'Cost Price and PO Number auto-populate from master data',
            'Cost Price, PO Number, and Date of Inward auto-populate from master data'
        )
        print("  Updated baby coil fields")
        break

# ============================================================
# CHANGE 7: Clarify Baby Coil ID is auto-populated in A.4
# ============================================================
print("Change 7: Update Baby Coil ID handling in A.4...")
table3 = doc.tables[3]
steps_cell_a4 = table3.rows[1].cells[1]
for para in steps_cell_a4.paragraphs:
    if 'Select Baby Coil ID from dropdown' in para.text:
        replace_paragraph_text(
            para,
            'Select Baby Coil ID from dropdown (auto-populates from "Coil to Slit" sheet)',
            'Baby Coil ID is auto-populated from the "Coil to Slit" sheet based on the selected SKU Code'
        )
        print("  Updated Baby Coil ID handling")
        break

# ============================================================
# CHANGE 8: Rename Section A.6 to "Yield & Loss Calculation"
# ============================================================
print("Change 8: Rename A.6...")
# Update heading
for p in doc.paragraphs:
    if p.text.strip() == 'A.6 Yield Calculation' and p.style.name == 'Heading 1':
        replace_paragraph_text(p, 'A.6 Yield Calculation', 'A.6 Yield & Loss Calculation')
        print("  Updated heading")
        break

# Update table of contents bullet
for p in doc.paragraphs:
    if p.text.strip() == 'A.6 - Yield Calculation' and p.style.name == 'List Bullet':
        replace_paragraph_text(p, 'A.6 - Yield Calculation', 'A.6 - Yield & Loss Calculation')
        print("  Updated TOC entry")
        break

# Update table 5 activity name
table5 = doc.tables[5]
activity_cell = table5.rows[1].cells[0]
for para in activity_cell.paragraphs:
    if 'Yield Calculation & Analysis' in para.text:
        replace_paragraph_text(para, 'Yield Calculation & Analysis', 'Yield & Loss Calculation & Analysis')
        print("  Updated table activity name")
        break

# ============================================================
# CHANGE 9: Add conversion efficiency and cost analysis in A.6
# ============================================================
print("Change 9: Add conversion efficiency and cost analysis...")
for p in doc.paragraphs:
    if p.text.startswith('This section describes the methodology for calculating manufacturing yield'):
        replace_paragraph_text(
            p,
            'Yield calculation is critical for measuring process efficiency and identifying areas for improvement.',
            'Yield calculation is critical for measuring process efficiency, overall conversion efficiency, and identifying areas for improvement. The variance analysis enables better production control and cost analysis.'
        )
        print("  Updated A.6 description")
        break

# ============================================================
# Save the updated document
# ============================================================
print(f"\nSaving to: {OUTPUT_PATH}")
doc.save(OUTPUT_PATH)
print("Done! Document saved successfully.")
