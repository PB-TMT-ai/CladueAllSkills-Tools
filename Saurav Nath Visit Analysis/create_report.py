from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================
# TITLE AND HEADER
# ============================================
title = doc.add_heading('Field Visit Performance Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Saurav Nath | Sales Manager | Jharkhand Region')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(68, 84, 106)

# Period
period = doc.add_paragraph()
period.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = period.add_run('Analysis Period: November 2025 - January 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(89, 89, 89)

# Report date
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Report Generated: {datetime.now().strftime("%d %B %Y")}')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()  # Spacing

# ============================================
# EXECUTIVE SUMMARY
# ============================================
doc.add_heading('Executive Summary', level=1)

summary_text = """This report analyzes the field visit performance of Saurav Nath covering the Jharkhand territory (primarily East Singhbum and Saraikela Kharsawan districts) over a 3-month period.

Key Highlights:
- Total of 345 customer visits across 55 working days
- Consistently met daily visit target of 6 visits per day (actual average: 6.3/day)
- December achieved 100% target compliance across all 20 working days
- Strong territory coverage spanning 20-24 pin codes and 7-10 districts

Critical Observation: Visit duration is significantly below benchmark (avg 2 min vs 15 min target) with over 90% of visits lasting less than 5 minutes. This requires immediate attention for quality improvement."""

summary = doc.add_paragraph(summary_text)
summary.paragraph_format.space_after = Pt(12)

# ============================================
# PERFORMANCE DASHBOARD TABLE
# ============================================
doc.add_heading('Performance Dashboard', level=1)

# Create table
table1 = doc.add_table(rows=7, cols=6)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
headers = ['Performance Dimension', 'Target', 'Nov 25', 'Dec 25', 'Jan 26', 'Status']
header_row = table1.rows[0].cells
for i, header in enumerate(headers):
    header_row[i].text = header
    for paragraph in header_row[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
data = [
    ['Visit Productivity (Visits/Day)', '6', '6.1', '6.5', '6.2', 'Good'],
    ['Territory Coverage (Unique Customers)', '50', '46', '63', '59', 'Good'],
    ['Geographic Spread (Pin Codes)', '10+', '20', '24', '21', 'Excellent'],
    ['Visit Quality (Avg Duration min)', '15', '1.9', '2.4', '2.4', 'Poor'],
    ['Days Meeting Target (%)', '80%', '69%', '100%', '95%', 'Good'],
    ['New Customer Acquisition', '-', '46', '24', '6', 'Declining'],
]

for row_idx, row_data in enumerate(data, start=1):
    row = table1.rows[row_idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row[col_idx].text = cell_data
        for paragraph in row[col_idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# ============================================
# TERRITORY COVERAGE TABLE
# ============================================
doc.add_heading('Territory Coverage Analysis', level=1)

table2 = doc.add_table(rows=4, cols=8)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
headers2 = ['Month', 'Total Visits', 'Unique Customers', 'Pin Codes', 'Districts', 'Retailer PB', 'Influencer', 'New Customers']
header_row2 = table2.rows[0].cells
for i, header in enumerate(headers2):
    header_row2[i].text = header
    for paragraph in header_row2[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data
terr_data = [
    ['Nov 25', '97', '46', '20', '9', '54 (56%)', '43 (44%)', '46'],
    ['Dec 25', '131', '63', '24', '10', '80 (61%)', '47 (36%)', '24'],
    ['Jan 26', '117', '59', '21', '7', '57 (49%)', '55 (47%)', '6'],
]

for row_idx, row_data in enumerate(terr_data, start=1):
    row = table2.rows[row_idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row[col_idx].text = cell_data
        for paragraph in row[col_idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# ============================================
# MONTHLY TRENDS TABLE
# ============================================
doc.add_heading('Monthly Performance Trends', level=1)

table3 = doc.add_table(rows=7, cols=5)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['Metric', 'Nov 25', 'Dec 25', 'Jan 26', 'Trend']
header_row3 = table3.rows[0].cells
for i, header in enumerate(headers3):
    header_row3[i].text = header
    for paragraph in header_row3[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

trends_data = [
    ['Total Visits', '97', '131', '117', 'Peaked Dec'],
    ['Avg Visits/Day', '6.1', '6.5', '6.2', 'Stable'],
    ['Working Days', '16', '20', '19', '-'],
    ['Unique Customers', '46', '63', '59', 'Stable'],
    ['Total Field Hours', '2.9', '5.1', '4.5', 'Low'],
    ['Short Visits (<5 min)', '92%', '93%', '90%', 'Critical'],
]

for row_idx, row_data in enumerate(trends_data, start=1):
    row = table3.rows[row_idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row[col_idx].text = cell_data
        for paragraph in row[col_idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# ============================================
# KEY INSIGHTS
# ============================================
doc.add_heading('Key Insights', level=1)

# Strengths
doc.add_heading('Strengths', level=2)
strengths = [
    'Consistently exceeding daily visit target of 6 visits (Average: 6.3 visits/day across 3 months)',
    'December achieved 100% target compliance - all 20 working days met the 6+ visit target',
    'Strong geographic coverage spanning 20-24 unique pin codes across 7-10 districts',
    'Balanced customer mix maintained: ~55% Retailer PB and ~45% Influencer engagement',
    'High customer reach: 76 unique customers contacted over the analysis period',
]
for s in strengths:
    p = doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()  # Spacing

# Areas of Concern
doc.add_heading('Areas of Concern', level=2)
concerns = [
    'Visit duration critically low: Average 2 minutes vs 15-minute benchmark - 87% gap',
    'Over 90% of visits lasting less than 5 minutes indicates superficial engagement',
    'New customer acquisition declining sharply: 46 (Nov) to 24 (Dec) to 6 (Jan) - 87% drop',
    '5 customers receiving 10+ visits in the period - potential over-concentration',
    'January showed territory contraction: 7 districts vs December\'s 10 districts',
    'Total field hours extremely low: Only 2.9 to 5.1 hours/month productive time',
]
for c in concerns:
    p = doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()  # Spacing

# ============================================
# RECOMMENDATIONS
# ============================================
doc.add_heading('Recommendations', level=1)

recommendations = [
    ('Quality Focus', 'Mandate minimum 15-minute meaningful engagement per customer visit. Current check-in/check-out pattern suggests compliance-driven visits rather than business development.'),
    ('New Business Target', 'Set monthly target for new customer acquisition (minimum 15-20 new customers). The declining trend from 46 to 6 new customers requires urgent intervention.'),
    ('Territory Review', 'Investigate January\'s district reduction from 10 to 7. Ensure consistent territory coverage is maintained across all allocated areas.'),
    ('Visit Validation', 'Implement enhanced verification mechanisms - GPS location tracking with photo evidence of customer meetings to ensure quality engagement.'),
    ('Customer Rotation', 'Review the 5 high-frequency accounts (10+ visits). Assess ROI and redirect effort to new prospects if not yielding proportionate business.'),
    ('Time Allocation', 'Restructure daily schedule to achieve 3-4 quality visits with 15+ minute engagement rather than 6+ superficial check-ins.'),
]

for title, desc in recommendations:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()  # Spacing

# ============================================
# FOOTER
# ============================================
doc.add_paragraph('_' * 80)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Confidential - For Internal Use Only')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

# Save document
output_path = 'D:/Saurav Nath Visit Analysis/Saurav_Nath_Field_Visit_Report.docx'
doc.save(output_path)
print(f"Report saved successfully to: {output_path}")
