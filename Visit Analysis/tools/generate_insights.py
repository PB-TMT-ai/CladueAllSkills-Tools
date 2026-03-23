"""
Generate Visit Analysis Insights Document
==========================================
Produces a concise Word document summarizing key findings from the visit data.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


OUTPUT_PATH = "D:/Visit Analysis/Visit_Analysis_Insights.docx"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def main():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Visit Analysis - Key Insights", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Data Period: October - December 2025  |  "
        "Total Visits: 18,815  |  "
        "Field Team: 62 members  |  "
        "Accounts Covered: 6,426"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # --- 1. Visit Volume Trend ---
    doc.add_heading("1. Visit Volume Trend", level=1)
    p = doc.add_paragraph()
    p.add_run("Monthly visits are growing steadily: ").bold = True
    p.add_run(
        "Oct (5,652) -> Nov (6,014, +6.4%) -> Dec (7,149, +18.9%). "
        "December saw the sharpest increase, suggesting either end-of-quarter push "
        "or seasonal demand uptick."
    )

    add_table(doc, ["Month", "Visits", "MoM Change"], [
        ["Oct 2025", "5,652", "-"],
        ["Nov 2025", "6,014", "+6.4%"],
        ["Dec 2025", "7,149", "+18.9%"],
    ])

    doc.add_paragraph("")

    # --- 2. Team Performance ---
    doc.add_heading("2. Team Performance Distribution", level=1)
    p = doc.add_paragraph()
    p.add_run("Top performers average 7-8 visits/day, ").bold = True
    p.add_run(
        "while the team averages 5.3 visits/person/day. "
        "The top 10 people account for a significant share of total visits. "
        "Abhijeet Biswas leads in daily visit rate (7.8/day)."
    )

    add_table(doc,
        ["Rank", "Person", "Total Visits"],
        [
            ["1", "Manoj Yadav", "503"],
            ["2", "Yogeshwar Richhariya", "494"],
            ["3", "Neeraj Sharma", "493"],
            ["4", "Abhijeet Biswas", "481"],
            ["5", "Ekeshwar Sahu", "476"],
        ],
    )

    p2 = doc.add_paragraph()
    p2.add_run("\nBottom performers: ").bold = True
    p2.add_run(
        "Amit Rochlani (4 visits), Pardeep Kumar (5 visits), "
        "Rachanbir Biding (41 visits). These warrant review for coverage gaps "
        "or role-specific reasons."
    )

    doc.add_paragraph("")

    # --- 3. Geographic Coverage ---
    doc.add_heading("3. Geographic Coverage", level=1)
    p = doc.add_paragraph()
    p.add_run("Top 5 states cover 62% of all visits. ").bold = True
    p.add_run(
        "Uttar Pradesh and Haryana dominate, followed by Madhya Pradesh, "
        "Rajasthan, and West Bengal. "
        "21 states were covered across the quarter."
    )

    add_table(doc,
        ["State", "Visits", "% Share"],
        [
            ["Uttar Pradesh", "2,944", "15.6%"],
            ["Haryana", "2,899", "15.4%"],
            ["Madhya Pradesh", "1,971", "10.5%"],
            ["Rajasthan", "1,883", "10.0%"],
            ["West Bengal", "1,880", "10.0%"],
            ["Odisha", "1,228", "6.5%"],
            ["Punjab", "1,054", "5.6%"],
            ["Delhi", "842", "4.5%"],
            ["Uttarakhand", "830", "4.4%"],
            ["Jharkhand", "810", "4.3%"],
        ],
    )

    doc.add_paragraph("")

    # --- 4. Visit Quality Concerns ---
    doc.add_heading("4. Visit Quality Concerns", level=1)
    p = doc.add_paragraph()
    p.add_run("63.7% of visits lasted less than 5 minutes. ").bold = True
    p.add_run(
        "29.2% had zero duration (checkout = checkin). "
        "The median visit duration is just 1 minute, while the mean is 18.7 minutes "
        "(skewed by longer visits). This suggests a large portion of visits may be "
        "quick check-ins or drive-bys rather than substantive meetings."
    )

    add_table(doc,
        ["Duration", "Count", "% of Total"],
        [
            ["0 min (zero)", "5,442", "29.2%"],
            ["< 5 min", "11,893", "63.7%"],
            ["5 - 30 min", "3,330", "17.8%"],
            ["30 - 60 min", "1,629", "8.7%"],
            ["60+ min", "1,809", "9.7%"],
        ],
    )

    doc.add_paragraph("")

    # --- 5. Inter-Visit Gap Analysis ---
    doc.add_heading("5. Inter-Visit Gap Analysis", level=1)
    p = doc.add_paragraph()
    p.add_run("40.8% of same-day gaps exceed 60 minutes. ").bold = True
    p.add_run(
        "The average gap between consecutive visits is 68.9 minutes (median: 43 min). "
        "Only 21.9% of gaps are under 10 minutes (back-to-back visits). "
        "The high proportion of 60+ minute gaps indicates significant travel time or "
        "idle time between visits that may be optimizable through better route planning."
    )

    add_table(doc,
        ["Gap Duration", "Count", "% of Gaps"],
        [
            ["< 10 min", "3,379", "21.9%"],
            ["10 - 30 min", "2,943", "19.0%"],
            ["30 - 60 min", "2,824", "18.3%"],
            ["60+ min", "6,308", "40.8%"],
        ],
    )

    doc.add_paragraph("")

    # --- 6. Meeting Outcomes ---
    doc.add_heading("6. Meeting Outcomes", level=1)
    p = doc.add_paragraph()
    p.add_run("Only 1.7% of visits generated an opportunity. ").bold = True
    p.add_run(
        "51.7% of outcomes were logged as 'Others' (unclassified), "
        "and 33% resulted in 'Got future date'. "
        "Lead qualification rate is 5.1%. "
        "The high 'Others' category suggests outcome tracking discipline needs improvement."
    )

    add_table(doc,
        ["Outcome", "Count", "%"],
        [
            ["Others", "9,723", "51.7%"],
            ["Got future date", "6,211", "33.0%"],
            ["Lead qualified", "966", "5.1%"],
            ["No opportunity", "936", "5.0%"],
            ["Opportunity generated", "315", "1.7%"],
            ["POC unavailable", "281", "1.5%"],
            ["Account created", "156", "0.8%"],
            ["Opportunity lost / Disqualification / Handover", "91", "0.5%"],
        ],
    )

    doc.add_paragraph("")

    # --- 7. Customer Mix ---
    doc.add_heading("7. Customer Type Coverage", level=1)
    p = doc.add_paragraph()
    p.add_run("74% of visits target Private Brand Retailers. ").bold = True
    p.add_run(
        "Influencers account for 18.1% of visits, distributors 1.3%. "
        "Projects, sites, and homes contractors together are under 1%, "
        "suggesting a retail-focused coverage model."
    )

    doc.add_paragraph("")

    # --- 8. Key Recommendations ---
    doc.add_heading("8. Key Recommendations", level=1)

    recommendations = [
        ("Visit Quality Audit",
         "63.7% of visits are under 5 minutes. Establish minimum visit duration "
         "thresholds and investigate zero-duration visits for potential data integrity issues."),
        ("Route Optimization",
         "40.8% of inter-visit gaps exceed 60 minutes. Clustering nearby accounts "
         "and optimizing daily routes could recover 1-2 additional visits per day per person."),
        ("Outcome Tracking Discipline",
         "51.7% of outcomes are 'Others'. Enforce structured outcome selection "
         "to enable meaningful conversion funnel analysis."),
        ("Low-Performer Review",
         "3 team members have fewer than 10 visits over 3 months. "
         "Review for territory assignment, role changes, or support needs."),
        ("High-Gap Individuals",
         "Use the Visit_Gap_Analysis Excel to identify specific people with disproportionately "
         "high 60+ minute gaps. These are the best candidates for route optimization."),
    ]

    for i, (title, desc) in enumerate(recommendations, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title}: ").bold = True
        p.add_run(desc)

    # --- Save ---
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
