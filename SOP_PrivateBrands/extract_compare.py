import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

def extract_tables(filepath, label):
    doc = Document(filepath)
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"  Tables: {len(doc.tables)}, Paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*60}")

    # Print paragraphs (non-empty)
    print("\n--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            print(f"  P{i}: [{p.style.name}] {txt[:200]}")

    # Print tables
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n',' | ')[:150] for cell in row.cells]
            print(f"  Row {j}: {cells}")

extract_tables(r'D:\SOP_PrivateBrands\Documents\JSWOrderLogging_JOTS_modification.docx', 'JOTS MODIFICATION')
extract_tables(r'D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V7.docx', 'V7 FORMATTED')
