#!/usr/bin/env python
"""
JSW Order Logging V16 -> V17 Renumbering Script

Changes ONLY activity number text in table cells (column 0) and heading paragraphs.
Uses run.text modification to preserve ALL formatting, layout, images, and structure.

Changes made:
  - Fix orphaned "i" -> "1.b." (Table 0, Row 2)
  - Fix "8. i." -> "8.i." and "8. ii" -> "8.ii." (Table 3)
  - Add trailing periods to Activity 6 sub-items (Table 8)
  - Renumber Activities 10-16 -> 9-15 (close Activity 9 gap)
  - Fix "11 D"/"11 D"/"11 E" -> "10.d."/"10.e."/"10.f." (Table 12)
  - Update heading paragraphs "Activity N:" for N=10..16
"""

import os
import re
from docx import Document
from lxml import etree

V16_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V16.docx"
V17_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V17.docx"

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WML_NS}

# ============================================================
# RENUMBERING MAP: (table_index, row_index) -> new_text
# ============================================================
CELL_RENUMBER_MAP = {
    # --- Fix formatting (Activities 1-8, numbers unchanged) ---
    (0, 3): "1.b.",       # orphaned "i" -> "1.b."
    (3, 1): "8.i.",       # "8. i." -> "8.i."
    (3, 2): "8.ii.",      # "8. ii" -> "8.ii."
    (8, 1): "6.a.",       # "6.a" -> "6.a."
    (8, 2): "6.b.",       # "6.b" -> "6.b."
    (8, 3): "6.c.",       # "6.c" -> "6.c."
    (8, 4): "6.d.",       # "6.d" -> "6.d."
    (8, 5): "6.e.",       # "6.e" -> "6.e."
    (8, 6): "6.f.",       # "6.f" -> "6.f."

    # --- Renumber Activity 10 -> 9 (Table 11) ---
    (11, 1): "9.",
    (11, 2): "9.a.",
    (11, 3): "9.b.",
    (11, 4): "9.c.",

    # --- Renumber Activity 11 -> 10 (Table 12) + fix D/E format ---
    (12, 1): "10.",
    (12, 2): "10.a.",
    (12, 3): "10.b.",
    (12, 4): "10.c.",
    (12, 5): "10.d.",     # "11 D" -> "10.d."
    (12, 6): "10.e.",     # "11 D" (duplicate) -> "10.e."
    (12, 7): "10.f.",     # "11 E" -> "10.f."

    # --- Renumber Activity 12 -> 11 (Table 13) ---
    (13, 1): "11.",
    (13, 2): "11.a.",
    (13, 3): "11.b.",
    (13, 4): "11.c.",

    # --- Renumber Activity 13 -> 12 (Table 14) ---
    (14, 1): "12.",
    (14, 2): "12.a.",
    (14, 3): "12.b.",
    (14, 4): "12.c.",

    # --- Renumber Activity 14 -> 13 (Table 15) ---
    (15, 1): "13.",
    (15, 2): "13.a.",
    (15, 3): "13.b.",
    (15, 4): "13.c.",

    # --- Renumber Activity 15 -> 14 (Table 16) ---
    (16, 1): "14.",
    (16, 2): "14.a.",
    (16, 3): "14.b.",
    (16, 4): "14.c.",

    # --- Renumber Activity 16 -> 15 (Table 17) ---
    (17, 1): "15.",
    (17, 2): "15.a.",
    (17, 3): "15.b.",
}

# Heading paragraph renumbering: Activity N -> Activity N-1 for N=10..16
HEADING_RENUMBER = {
    "Activity 10": "Activity 9",
    "Activity 11": "Activity 10",
    "Activity 12": "Activity 11",
    "Activity 13": "Activity 12",
    "Activity 14": "Activity 13",
    "Activity 15": "Activity 14",
    "Activity 16": "Activity 15",
}


def replace_cell_text_safe(cell, new_text):
    """Replace activity text in cell while preserving ALL formatting.

    Only modifies run.text (<w:t> content). Never touches <w:rPr>,
    <w:pPr>, <w:tcPr>, images, or any structural XML elements.
    Skips runs containing page breaks (<w:br>).
    """
    para = cell.paragraphs[0]
    runs = para.runs

    if len(runs) == 0:
        return False

    # Find first text-bearing run (skip runs with page breaks)
    first_text_idx = None
    for i, run in enumerate(runs):
        br_elements = run._element.findall(f'{{{WML_NS}}}br')
        if br_elements:
            continue  # Skip page break runs
        # Also skip runs that are purely whitespace before the actual number
        # but only if there are more runs after
        if run.text.strip() == "" and i < len(runs) - 1:
            continue
        first_text_idx = i
        break

    if first_text_idx is None:
        # All runs have breaks or are empty - set on last run
        runs[-1].text = new_text
        return True

    # Set new text on the first meaningful text run
    runs[first_text_idx].text = new_text

    # Clear all subsequent runs' text (but keep the run elements for formatting)
    for i in range(first_text_idx + 1, len(runs)):
        runs[i].text = ""

    return True


def replace_heading_text(para, old_activity, new_activity):
    """Replace 'Activity N' with 'Activity M' in heading paragraph runs.

    Only modifies the text portion, preserving all run formatting.
    """
    for run in para.runs:
        if old_activity in run.text:
            run.text = run.text.replace(old_activity, new_activity)
            return True
    return False


def count_xml_elements(doc_element):
    """Count key XML element types for verification."""
    body = doc_element.element.body
    counts = {}
    for tag_local in ['r', 'p', 'tc', 'tr', 'tbl']:
        tag = f'{{{WML_NS}}}{tag_local}'
        counts[tag_local] = len(body.findall(f'.//{tag}'))
    # Count images (blipFill in drawing namespace)
    a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    counts['blipFill'] = len(body.findall(f'.//{{{a_ns}}}blipFill'))
    return counts


def get_table_layout_xml(doc):
    """Extract table/cell width XML for layout comparison."""
    layouts = []
    for t_idx, table in enumerate(doc.tables):
        tbl_elem = table._tbl
        # Get table width
        tblW = tbl_elem.find(f'.//{{{WML_NS}}}tblW')
        tblW_xml = etree.tostring(tblW).decode() if tblW is not None else "none"

        row_layouts = []
        for r_idx, row in enumerate(table.rows):
            tr_elem = row._tr
            # Get row properties
            trPr = tr_elem.find(f'{{{WML_NS}}}trPr')
            trPr_xml = etree.tostring(trPr).decode() if trPr is not None else "none"

            cell_widths = []
            for c_idx, cell in enumerate(row.cells):
                tc_elem = cell._tc
                tcW = tc_elem.find(f'.//{{{WML_NS}}}tcW')
                tcW_xml = etree.tostring(tcW).decode() if tcW is not None else "none"
                cell_widths.append(tcW_xml)

            row_layouts.append((trPr_xml, cell_widths))
        layouts.append((tblW_xml, row_layouts))
    return layouts


def main():
    print("=" * 70)
    print("JSW ORDER LOGGING V16 -> V17 RENUMBERING")
    print("=" * 70)
    print()

    # ---- Step 1: Load V16 ----
    print("Step 1: Loading V16...")
    if not os.path.exists(V16_PATH):
        print(f"ERROR: V16 not found at {V16_PATH}")
        return False

    doc = Document(V16_PATH)
    print(f"  Tables: {len(doc.tables)}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print()

    # ---- Step 2: Pre-change analysis ----
    print("Step 2: Pre-change analysis (current state of cells to modify)...")
    for (t_idx, r_idx), new_text in sorted(CELL_RENUMBER_MAP.items()):
        table = doc.tables[t_idx]
        cell = table.rows[r_idx].cells[0]
        old_text = cell.text.strip()
        num_runs = len(cell.paragraphs[0].runs)
        print(f"  T{t_idx:2d} R{r_idx}: '{old_text}' ({num_runs} runs) -> '{new_text}'")
    print()

    # Capture pre-change metrics
    pre_element_counts = count_xml_elements(doc)
    pre_layouts = get_table_layout_xml(doc)
    pre_para_count = len(doc.paragraphs)
    pre_table_count = len(doc.tables)
    pre_row_counts = [len(t.rows) for t in doc.tables]

    # Also capture all non-column-0 cell text for comparison
    pre_other_cells = {}
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx in range(1, len(row.cells)):
                pre_other_cells[(t_idx, r_idx, c_idx)] = row.cells[c_idx].text

    # ---- Step 3: Apply table cell renumbering ----
    print("Step 3: Applying table cell renumbering (39 cells)...")
    changes_made = 0
    for (t_idx, r_idx), new_text in sorted(CELL_RENUMBER_MAP.items()):
        table = doc.tables[t_idx]
        cell = table.rows[r_idx].cells[0]
        old_text = cell.text.strip()

        if replace_cell_text_safe(cell, new_text):
            changes_made += 1
            # Verify the change took effect
            actual = cell.text.strip()
            status = "OK" if actual == new_text else f"MISMATCH (got '{actual}')"
            print(f"  [{status}] T{t_idx:2d} R{r_idx}: '{old_text}' -> '{new_text}'")
        else:
            print(f"  [WARN] T{t_idx:2d} R{r_idx}: No runs found, skipped")

    print(f"\n  Total cell changes: {changes_made}")
    print()

    # ---- Step 4: Apply heading paragraph renumbering ----
    print("Step 4: Applying heading paragraph renumbering (7 headings)...")
    heading_changes = 0
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for old_act, new_act in HEADING_RENUMBER.items():
            if text.startswith(old_act + ":") or text.startswith(old_act + " "):
                if replace_heading_text(para, old_act, new_act):
                    heading_changes += 1
                    print(f"  [OK] Para {p_idx}: '{old_act}' -> '{new_act}' in: {text[:70]}")
                break

    print(f"\n  Total heading changes: {heading_changes}")
    print()

    # ---- Step 5: Save as V17 ----
    print("Step 5: Saving as V17...")
    doc.save(V17_PATH)
    print(f"  Saved to: {V17_PATH}")
    print()

    # ---- Step 6: Verification ----
    print("Step 6: Full verification...")
    print("-" * 50)

    v17 = Document(V17_PATH)
    errors = []

    # 6a. Paragraph count
    if len(v17.paragraphs) != pre_para_count:
        errors.append(f"Paragraph count: {pre_para_count} -> {len(v17.paragraphs)}")
    else:
        print(f"  [OK] Paragraph count: {len(v17.paragraphs)}")

    # 6b. Table count
    if len(v17.tables) != pre_table_count:
        errors.append(f"Table count: {pre_table_count} -> {len(v17.tables)}")
    else:
        print(f"  [OK] Table count: {len(v17.tables)}")

    # 6c. Row counts per table
    post_row_counts = [len(t.rows) for t in v17.tables]
    if post_row_counts != pre_row_counts:
        errors.append(f"Row counts changed: {pre_row_counts} -> {post_row_counts}")
    else:
        print(f"  [OK] Row counts per table: all match")

    # 6d. Verify all renumbered cells have correct values
    cell_errors = 0
    for (t_idx, r_idx), expected in sorted(CELL_RENUMBER_MAP.items()):
        actual = v17.tables[t_idx].rows[r_idx].cells[0].text.strip()
        if actual != expected:
            errors.append(f"Cell T{t_idx} R{r_idx}: expected '{expected}', got '{actual}'")
            cell_errors += 1
    print(f"  [{'OK' if cell_errors == 0 else 'FAIL'}] Renumbered cells: {39 - cell_errors}/39 correct")

    # 6e. Verify all non-column-0 cells unchanged
    other_cell_diffs = 0
    for (t_idx, r_idx, c_idx), pre_text in pre_other_cells.items():
        post_text = v17.tables[t_idx].rows[r_idx].cells[c_idx].text
        if post_text != pre_text:
            errors.append(f"Non-activity cell T{t_idx} R{r_idx} C{c_idx} changed!")
            other_cell_diffs += 1
    print(f"  [{'OK' if other_cell_diffs == 0 else 'FAIL'}] Non-activity cells: {len(pre_other_cells) - other_cell_diffs}/{len(pre_other_cells)} unchanged")

    # 6f. XML element counts
    post_element_counts = count_xml_elements(v17)
    element_diffs = []
    for key in pre_element_counts:
        if pre_element_counts[key] != post_element_counts[key]:
            element_diffs.append(f"{key}: {pre_element_counts[key]} -> {post_element_counts[key]}")
    if element_diffs:
        errors.append(f"XML element count changes: {', '.join(element_diffs)}")
        print(f"  [FAIL] XML element counts: {', '.join(element_diffs)}")
    else:
        print(f"  [OK] XML element counts: all match (r={pre_element_counts['r']}, p={pre_element_counts['p']}, tc={pre_element_counts['tc']}, images={pre_element_counts['blipFill']})")

    # 6g. Table/cell layout comparison
    post_layouts = get_table_layout_xml(v17)
    layout_diffs = 0
    for t_idx in range(len(pre_layouts)):
        pre_tblW, pre_rows = pre_layouts[t_idx]
        post_tblW, post_rows = post_layouts[t_idx]
        if pre_tblW != post_tblW:
            errors.append(f"Table {t_idx} width changed")
            layout_diffs += 1
        for r_idx in range(len(pre_rows)):
            pre_trPr, pre_cells = pre_rows[r_idx]
            post_trPr, post_cells = post_rows[r_idx]
            if pre_trPr != post_trPr:
                errors.append(f"Table {t_idx} Row {r_idx} trPr changed")
                layout_diffs += 1
            for c_idx in range(len(pre_cells)):
                if pre_cells[c_idx] != post_cells[c_idx]:
                    errors.append(f"Table {t_idx} Row {r_idx} Cell {c_idx} width changed")
                    layout_diffs += 1
    print(f"  [{'OK' if layout_diffs == 0 else 'FAIL'}] Table/cell layout: {'all match' if layout_diffs == 0 else f'{layout_diffs} differences'}")

    # 6h. Verify heading paragraphs by checking specific paragraph indices
    heading_errors = 0
    # Expected new headings at specific paragraph indices
    expected_headings = {
        74: "Activity 9:",
        80: "Activity 10:",
        84: "Activity 11:",
        87: "Activity 12:",
        90: "Activity 13:",
        93: "Activity 14:",
        109: "Activity 15:",
    }
    for p_idx, expected_prefix in expected_headings.items():
        if p_idx < len(v17.paragraphs):
            text = v17.paragraphs[p_idx].text.strip()
            if text.startswith(expected_prefix):
                pass  # correct
            else:
                errors.append(f"Heading Para {p_idx}: expected '{expected_prefix}...', got '{text[:60]}'")
                heading_errors += 1
        else:
            errors.append(f"Heading Para {p_idx}: paragraph index out of range")
            heading_errors += 1
    # Also verify "Activity 16" no longer exists anywhere in headings
    for p_idx, para in enumerate(v17.paragraphs):
        text = para.text.strip()
        if text.startswith("Activity 16:") or text.startswith("Activity 16 "):
            errors.append(f"Old 'Activity 16' still found at Para {p_idx}: {text[:60]}")
            heading_errors += 1
    print(f"  [{'OK' if heading_errors == 0 else 'FAIL'}] Heading paragraphs: {7 - heading_errors}/7 correct")

    # 6i. Spot-check run formatting on modified cells
    v16_reloaded = Document(V16_PATH)
    format_diffs = 0
    for (t_idx, r_idx) in list(CELL_RENUMBER_MAP.keys())[:10]:  # spot-check first 10
        v16_cell = v16_reloaded.tables[t_idx].rows[r_idx].cells[0]
        v17_cell = v17.tables[t_idx].rows[r_idx].cells[0]
        v16_runs = v16_cell.paragraphs[0].runs
        v17_runs = v17_cell.paragraphs[0].runs

        if len(v16_runs) != len(v17_runs):
            errors.append(f"Run count changed at T{t_idx} R{r_idx}: {len(v16_runs)} -> {len(v17_runs)}")
            format_diffs += 1
            continue

        for run_i in range(len(v16_runs)):
            v16_rPr = v16_runs[run_i]._element.find(f'{{{WML_NS}}}rPr')
            v17_rPr = v17_runs[run_i]._element.find(f'{{{WML_NS}}}rPr')
            v16_rPr_xml = etree.tostring(v16_rPr).decode() if v16_rPr is not None else "none"
            v17_rPr_xml = etree.tostring(v17_rPr).decode() if v17_rPr is not None else "none"
            if v16_rPr_xml != v17_rPr_xml:
                errors.append(f"rPr changed at T{t_idx} R{r_idx} Run{run_i}")
                format_diffs += 1
    print(f"  [{'OK' if format_diffs == 0 else 'FAIL'}] Run formatting: {'preserved' if format_diffs == 0 else f'{format_diffs} differences'}")

    # ---- Summary ----
    print()
    print("=" * 70)
    if errors:
        print(f"VERIFICATION: {len(errors)} ISSUE(S) FOUND")
        for err in errors:
            print(f"  ! {err}")
    else:
        print("VERIFICATION: ALL CHECKS PASSED")
    print("=" * 70)
    print()
    print(f"Source:  {V16_PATH}")
    print(f"Output:  {V17_PATH}")
    print(f"Changes: {changes_made} cell(s) + {heading_changes} heading(s)")
    print()

    return len(errors) == 0


if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)
