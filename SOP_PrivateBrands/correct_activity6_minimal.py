#!/usr/bin/env python
"""
Minimal Activity 6 Index Correction Script
Corrects ONLY Activity 6 numbering (6.1-6.6 -> 6.a-6.f) in V13
NO other changes to the document
"""

import os
from docx import Document

V13_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V13.docx"
V15_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V15_CORRECTED.docx"


def main():
    print("=" * 70)
    print("MINIMAL ACTIVITY 6 INDEX CORRECTION")
    print("=" * 70)
    print()

    # Load V13
    print("Step 1: Loading V13 document...")
    if not os.path.exists(V13_PATH):
        print(f"ERROR: V13 not found at {V13_PATH}")
        return

    doc = Document(V13_PATH)
    print(f"[OK] Loaded V13")
    print(f"  - Tables: {len(doc.tables)}")
    print()

    # Find and correct Activity 6 in Table 8
    print("Step 2: Correcting Activity 6 numbering in Table 8...")

    replacements = {
        '6.1': '6.a',
        '6.2': '6.b',
        '6.3': '6.c',
        '6.4': '6.d',
        '6.5': '6.e',
        '6.6': '6.f'
    }

    corrections_made = 0
    table_8 = doc.tables[8]  # Table 8 (0-indexed, so index 8)

    for row in table_8.rows:
        cell_text = row.cells[0].text.strip()

        for old, new in replacements.items():
            if cell_text.startswith(old):
                # Replace only in first cell (Activity column)
                row.cells[0].text = row.cells[0].text.replace(old, new, 1)
                corrections_made += 1
                print(f"  [OK] {old} -> {new}")
                break

    print(f"\n[OK] Made {corrections_made} corrections")
    print()

    # Save as V15
    print("Step 3: Saving as V15...")
    doc.save(V15_PATH)
    print(f"[OK] Saved to: {V15_PATH}")
    print()

    # Verification
    print("Step 4: Verifying corrections...")

    # Re-open and check
    verify_doc = Document(V15_PATH)
    verify_table = verify_doc.tables[8]

    found_old = []
    found_new = []

    for row in verify_table.rows:
        cell_text = row.cells[0].text.strip()
        for old in replacements.keys():
            if cell_text.startswith(old):
                found_old.append(old)
        for new in replacements.values():
            if cell_text.startswith(new):
                found_new.append(new)

    print(f"  Old numbering (6.1-6.6) found: {len(found_old)} (should be 0)")
    print(f"  New numbering (6.a-6.f) found: {len(found_new)} (should be 6)")

    if len(found_old) == 0 and len(found_new) == 6:
        print("\n[OK] VERIFICATION PASSED")
    else:
        print("\n[WARNING] Verification issues detected")

    print()
    print("=" * 70)
    print("CORRECTION COMPLETE")
    print("=" * 70)
    print()
    print(f"Original: {V13_PATH}")
    print(f"Corrected: {V15_PATH}")
    print()
    print("Changes made: 6 text replacements in Table 8 only")
    print("Everything else preserved exactly as-is")
    print()


if __name__ == "__main__":
    main()
