#!/usr/bin/env python
"""
JSW Order Logging V17 -> V18 Layout Fix & Comment Removal Script

Fixes:
  1. Table widths: Scale 12 overflowing tables to fit within 10,800 twip page width
  2. Comments: Remove all 7 reviewer comments and associated package parts

No content changes. Images, text, formatting preserved. Only widths and comments touched.
"""

import os
from docx import Document
from lxml import etree

V17_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V17.docx"
V18_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V18.docx"

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WML_NS}

TARGET_WIDTH = 10800  # twips = 7.5 inches available content width

# Comment relationship type URIs
COMMENT_REL_TYPES = {
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments',
    'http://schemas.microsoft.com/office/2011/relationships/commentsExtended',
    'http://schemas.microsoft.com/office/2016/09/relationships/commentsIds',
    'http://schemas.microsoft.com/office/2018/08/relationships/commentsExtensible',
}


# ============================================================
# COMMENT REMOVAL
# ============================================================

def remove_comment_body_elements(body):
    """Remove all comment-related XML elements from the document body.

    Removes:
    - <w:commentRangeStart> elements
    - <w:commentRangeEnd> elements
    - <w:r> runs that contain <w:commentReference> (these have no text content)

    Returns dict with removal counts.
    """
    removed = {'rangeStart': 0, 'rangeEnd': 0, 'refRuns': 0}

    # 1. Remove <w:commentRangeStart> elements
    for elem in body.findall(f'.//{{{WML_NS}}}commentRangeStart'):
        elem.getparent().remove(elem)
        removed['rangeStart'] += 1

    # 2. Remove <w:commentRangeEnd> elements
    for elem in body.findall(f'.//{{{WML_NS}}}commentRangeEnd'):
        elem.getparent().remove(elem)
        removed['rangeEnd'] += 1

    # 3. Remove <w:r> runs containing <w:commentReference>
    for ref_elem in body.findall(f'.//{{{WML_NS}}}commentReference'):
        run = ref_elem.getparent()  # <w:r>
        if run is not None and run.getparent() is not None:
            run.getparent().remove(run)
            removed['refRuns'] += 1

    return removed


def remove_comment_relationships(doc):
    """Remove comment-related relationship entries from the document package.

    When doc.save() is called, python-docx traverses relationships to find parts.
    Removed relationships mean their target parts are excluded from the output ZIP.
    """
    rels = doc.part.rels
    comment_rids = []

    for rid, rel in list(rels.items()):
        if rel.reltype in COMMENT_REL_TYPES:
            comment_rids.append((rid, rel.reltype))

    for rid, _ in comment_rids:
        del rels[rid]

    return comment_rids


# ============================================================
# TABLE WIDTH SCALING
# ============================================================

def scale_widths(widths, target):
    """Proportionally scale a list of integer widths to sum to target exactly.

    Rounding error is absorbed by the largest column.
    """
    total = sum(widths)
    if total == 0:
        return widths

    new_widths = [round(w * target / total) for w in widths]

    # Fix rounding error
    rounding_error = sum(new_widths) - target
    if rounding_error != 0:
        largest_idx = widths.index(max(widths))
        new_widths[largest_idx] -= rounding_error

    assert sum(new_widths) == target, f"Scaling failed: {sum(new_widths)} != {target}"
    return new_widths


def get_grid_widths(tbl):
    """Get column widths from <w:tblGrid> element."""
    tblGrid = tbl.find(f'{{{WML_NS}}}tblGrid')
    if tblGrid is None:
        return []
    grid_cols = tblGrid.findall(f'{{{WML_NS}}}gridCol')
    return [int(gc.get(f'{{{WML_NS}}}w', '0')) for gc in grid_cols]


def get_row_cell_widths(tr):
    """Get cell widths from a table row's <w:tcW> elements."""
    tcs = tr.findall(f'{{{WML_NS}}}tc')
    widths = []
    for tc in tcs:
        tcPr = tc.find(f'{{{WML_NS}}}tcPr')
        if tcPr is not None:
            tcW = tcPr.find(f'{{{WML_NS}}}tcW')
            if tcW is not None:
                widths.append(int(tcW.get(f'{{{WML_NS}}}w', '0')))
            else:
                widths.append(0)
        else:
            widths.append(0)
    return widths


def set_row_cell_widths(tr, new_widths):
    """Set cell widths in a table row."""
    tcs = tr.findall(f'{{{WML_NS}}}tc')
    for tc, new_w in zip(tcs, new_widths):
        tcPr = tc.find(f'{{{WML_NS}}}tcPr')
        if tcPr is not None:
            tcW = tcPr.find(f'{{{WML_NS}}}tcW')
            if tcW is not None:
                tcW.set(f'{{{WML_NS}}}w', str(new_w))


def fix_table_widths(doc, target=TARGET_WIDTH):
    """Scale overflowing tables to fit within target width.

    Modifies: <w:gridCol>, <w:tblW>, <w:tcW>, <w:tblInd>
    Preserves: layout type, borders, shading, images, formatting
    """
    fixes = []

    for t_idx, table in enumerate(doc.tables):
        tbl = table._tbl
        tblPr = tbl.find(f'{{{WML_NS}}}tblPr')

        # Get grid widths
        grid_widths = get_grid_widths(tbl)
        grid_total = sum(grid_widths)

        # Skip if already within bounds
        if grid_total <= target:
            continue

        # --- Scale grid columns ---
        tblGrid = tbl.find(f'{{{WML_NS}}}tblGrid')
        grid_cols = tblGrid.findall(f'{{{WML_NS}}}gridCol')
        new_grid_widths = scale_widths(grid_widths, target)

        for gc, new_w in zip(grid_cols, new_grid_widths):
            gc.set(f'{{{WML_NS}}}w', str(new_w))

        # --- Update <w:tblW> ---
        if tblPr is not None:
            tblW = tblPr.find(f'{{{WML_NS}}}tblW')
            if tblW is not None:
                tblW.set(f'{{{WML_NS}}}w', str(target))

            # --- Reset table indent to 0 ---
            tblInd = tblPr.find(f'{{{WML_NS}}}tblInd')
            if tblInd is not None:
                tblInd.set(f'{{{WML_NS}}}w', '0')

        # --- Scale cell widths in every row ---
        rows = tbl.findall(f'{{{WML_NS}}}tr')

        # Check if first data row cell widths match grid widths
        if len(rows) > 0:
            first_row_widths = get_row_cell_widths(rows[0])
            first_row_total = sum(first_row_widths)

            if first_row_widths == grid_widths:
                # Cell widths match grid — use same new widths for all rows
                for tr in rows:
                    set_row_cell_widths(tr, new_grid_widths)
            else:
                # Cell widths differ from grid (e.g. Tables 9, 10)
                # Scale each row's cell widths independently to target
                for tr in rows:
                    row_widths = get_row_cell_widths(tr)
                    row_total = sum(row_widths)
                    if row_total > 0 and row_total != target:
                        new_row_widths = scale_widths(row_widths, target)
                        set_row_cell_widths(tr, new_row_widths)

        fixes.append({
            'table': t_idx,
            'old_total': grid_total,
            'new_total': target,
            'excess': grid_total - target,
        })

    return fixes


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 70)
    print("JSW ORDER LOGGING V17 -> V18 LAYOUT FIX & COMMENT REMOVAL")
    print("=" * 70)
    print()

    # ---- Step 1: Load ----
    print("Step 1: Loading V17...")
    if not os.path.exists(V17_PATH):
        print(f"ERROR: V17 not found at {V17_PATH}")
        return False

    doc = Document(V17_PATH)
    body = doc.element.body

    # Capture pre-change metrics
    pre_para_count = len(doc.paragraphs)
    pre_table_count = len(doc.tables)
    pre_row_counts = [len(t.rows) for t in doc.tables]

    # Count drawings (images)
    drawing_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    pre_drawing_count = len(body.findall(f'.//{{{wp_ns}}}inline')) + len(body.findall(f'.//{{{wp_ns}}}anchor'))

    # Count comment elements
    pre_comment_starts = len(body.findall(f'.//{{{WML_NS}}}commentRangeStart'))
    pre_comment_ends = len(body.findall(f'.//{{{WML_NS}}}commentRangeEnd'))
    pre_comment_refs = len(body.findall(f'.//{{{WML_NS}}}commentReference'))

    # Capture all cell text for verification
    pre_cell_text = {}
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx in range(len(row.cells)):
                pre_cell_text[(t_idx, r_idx, c_idx)] = row.cells[c_idx].text

    print(f"  Tables: {pre_table_count}")
    print(f"  Paragraphs: {pre_para_count}")
    print(f"  Images (drawings): {pre_drawing_count}")
    print(f"  Comment elements: {pre_comment_starts} rangeStart, {pre_comment_ends} rangeEnd, {pre_comment_refs} references")
    print()

    # ---- Step 2: Remove comments ----
    print("Step 2: Removing comments...")

    # 2a. Remove comment elements from body
    removed = remove_comment_body_elements(body)
    print(f"  Body elements removed: {removed['rangeStart']} rangeStart, {removed['rangeEnd']} rangeEnd, {removed['refRuns']} reference runs")

    # 2b. Remove comment relationships from package
    removed_rels = remove_comment_relationships(doc)
    print(f"  Package relationships removed: {len(removed_rels)}")
    for rid, reltype in removed_rels:
        print(f"    {rid}: {reltype.split('/')[-1]}")
    print()

    # ---- Step 3: Fix table widths ----
    print("Step 3: Fixing overflowing table widths...")

    # Pre-analysis
    print(f"  Target width: {TARGET_WIDTH} twips ({TARGET_WIDTH / 1440:.2f} inches)")
    print()

    fixes = fix_table_widths(doc, TARGET_WIDTH)

    for fix in fixes:
        print(f"  [OK] Table {fix['table']:2d}: {fix['old_total']} -> {fix['new_total']} (was {fix['excess']} over)")

    print(f"\n  Tables fixed: {len(fixes)}")
    print()

    # ---- Step 4: Save ----
    print("Step 4: Saving as V18...")
    doc.save(V18_PATH)
    print(f"  Saved to: {V18_PATH}")
    print()

    # ---- Step 5: Verification ----
    print("Step 5: Full verification...")
    print("-" * 50)

    v18 = Document(V18_PATH)
    v18_body = v18.element.body
    errors = []

    # 5a. Paragraph count
    post_para = len(v18.paragraphs)
    if post_para != pre_para_count:
        errors.append(f"Paragraph count: {pre_para_count} -> {post_para}")
    print(f"  [{'OK' if post_para == pre_para_count else 'FAIL'}] Paragraphs: {post_para}")

    # 5b. Table count
    post_tables = len(v18.tables)
    if post_tables != pre_table_count:
        errors.append(f"Table count: {pre_table_count} -> {post_tables}")
    print(f"  [{'OK' if post_tables == pre_table_count else 'FAIL'}] Tables: {post_tables}")

    # 5c. Row counts
    post_row_counts = [len(t.rows) for t in v18.tables]
    if post_row_counts != pre_row_counts:
        errors.append(f"Row counts changed")
    print(f"  [{'OK' if post_row_counts == pre_row_counts else 'FAIL'}] Row counts: all match")

    # 5d. Drawing count (images)
    post_drawing_count = len(v18_body.findall(f'.//{{{wp_ns}}}inline')) + len(v18_body.findall(f'.//{{{wp_ns}}}anchor'))
    if post_drawing_count != pre_drawing_count:
        errors.append(f"Drawing count: {pre_drawing_count} -> {post_drawing_count}")
    print(f"  [{'OK' if post_drawing_count == pre_drawing_count else 'FAIL'}] Images: {post_drawing_count}")

    # 5e. No comment elements remain
    post_comment_starts = len(v18_body.findall(f'.//{{{WML_NS}}}commentRangeStart'))
    post_comment_ends = len(v18_body.findall(f'.//{{{WML_NS}}}commentRangeEnd'))
    post_comment_refs = len(v18_body.findall(f'.//{{{WML_NS}}}commentReference'))
    total_comments = post_comment_starts + post_comment_ends + post_comment_refs

    if total_comments > 0:
        errors.append(f"Comment elements remain: {post_comment_starts}+{post_comment_ends}+{post_comment_refs}")
    print(f"  [{'OK' if total_comments == 0 else 'FAIL'}] Comments removed: 0 elements remain")

    # 5f. No comment relationships remain
    comment_rels_remaining = 0
    for rid, rel in v18.part.rels.items():
        if rel.reltype in COMMENT_REL_TYPES:
            comment_rels_remaining += 1
            errors.append(f"Comment relationship still exists: {rid}")
    print(f"  [{'OK' if comment_rels_remaining == 0 else 'FAIL'}] Comment relationships: 0 remain")

    # 5g. All tables now within bounds
    tables_over = 0
    for t_idx, table in enumerate(v18.tables):
        grid_widths = get_grid_widths(table._tbl)
        grid_total = sum(grid_widths)
        if grid_total > TARGET_WIDTH:
            errors.append(f"Table {t_idx} still overflows: {grid_total} > {TARGET_WIDTH}")
            tables_over += 1
    print(f"  [{'OK' if tables_over == 0 else 'FAIL'}] Table widths: all {post_tables} within {TARGET_WIDTH} twips")

    # 5h. Within-bounds tables unchanged
    v17_reload = Document(V17_PATH)
    unchanged_tables = [0, 3, 4, 5, 6, 16]
    unchanged_ok = True
    for t_idx in unchanged_tables:
        v17_grid = get_grid_widths(v17_reload.tables[t_idx]._tbl)
        v18_grid = get_grid_widths(v18.tables[t_idx]._tbl)
        if v17_grid != v18_grid:
            errors.append(f"Within-bounds table {t_idx} was modified: {v17_grid} -> {v18_grid}")
            unchanged_ok = False
    print(f"  [{'OK' if unchanged_ok else 'FAIL'}] Within-bounds tables {unchanged_tables}: unchanged")

    # 5i. Cell text content unchanged
    text_diffs = 0
    for (t_idx, r_idx, c_idx), pre_text in pre_cell_text.items():
        try:
            post_text = v18.tables[t_idx].rows[r_idx].cells[c_idx].text
            if pre_text != post_text:
                errors.append(f"T{t_idx}R{r_idx}C{c_idx} text changed: '{pre_text[:40]}' -> '{post_text[:40]}'")
                text_diffs += 1
        except (IndexError, KeyError):
            errors.append(f"T{t_idx}R{r_idx}C{c_idx} cell missing")
            text_diffs += 1
    print(f"  [{'OK' if text_diffs == 0 else 'FAIL'}] Cell text: {len(pre_cell_text) - text_diffs}/{len(pre_cell_text)} identical")

    # 5j. Verify comment ZIP parts are gone
    import zipfile
    with zipfile.ZipFile(V18_PATH, 'r') as z:
        comment_files = [n for n in z.namelist() if 'comment' in n.lower()]
    if comment_files:
        errors.append(f"Comment files still in ZIP: {comment_files}")
    print(f"  [{'OK' if not comment_files else 'FAIL'}] ZIP comment files: {len(comment_files)} remain")

    # ---- Summary ----
    print()
    print("=" * 70)
    if errors:
        print(f"VERIFICATION: {len(errors)} ISSUE(S) FOUND")
        for err in errors:
            print(f"  ! {err}")
    else:
        print("VERIFICATION: ALL CHECKS PASSED")
    print("=" * 70)
    print()
    print(f"Source:  {V17_PATH}")
    print(f"Output:  {V18_PATH}")
    print(f"Changes: {len(fixes)} tables scaled, {removed['rangeStart'] + removed['rangeEnd'] + removed['refRuns']} comment elements removed, {len(removed_rels)} comment parts removed")
    print()

    return len(errors) == 0


if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)
