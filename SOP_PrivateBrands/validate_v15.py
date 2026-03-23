#!/usr/bin/env python
"""
V15 QA Validation Script
Verifies V15 document against quality checklist
"""

import os
from docx import Document

V15_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V15.docx"
V13_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V13.docx"
EXEC_SUMMARY_PATH = r"D:\SOP_PrivateBrands\JSWOrderLogging_ExecutiveSummary.docx"


def check_activity_6_numbering(doc):
    """Verify Activity 6 uses 6.a-6.f format, not 6.1-6.6"""
    errors = []
    correct_count = 0

    for table in doc.tables:
        for row in table.rows:
            cell_text = row.cells[0].text.strip()

            # Check for old numbering (should not exist)
            if any(cell_text.startswith(f'6.{i}') for i in range(1, 7)):
                errors.append(f"Found old numbering: {cell_text}")

            # Check for correct numbering
            if any(cell_text.startswith(f'6.{letter}') for letter in 'abcdef'):
                correct_count += 1

    return {
        'passed': len(errors) == 0,
        'correct_count': correct_count,
        'errors': errors
    }


def check_activity_9_gap(doc):
    """Verify Activity 9 gap is documented"""
    activity_9_mentioned = False

    for para in doc.paragraphs:
        text = para.text.lower()
        if 'activity 9' in text and ('intentionally removed' in text or 'not applicable' in text):
            activity_9_mentioned = True
            break

    return {
        'passed': activity_9_mentioned,
        'message': 'Activity 9 gap documented' if activity_9_mentioned else 'Activity 9 gap not documented'
    }


def check_table_count(v13_doc, v15_doc):
    """Verify same number of tables as V13"""
    v13_count = len(v13_doc.tables)
    v15_count = len(v15_doc.tables)

    # V15 has additional tables (exec summary, TOC), so should have MORE tables
    return {
        'passed': v15_count >= v13_count,
        'v13_count': v13_count,
        'v15_count': v15_count,
        'message': f"V13: {v13_count} tables, V15: {v15_count} tables"
    }


def check_5_column_structure(doc):
    """Verify all process tables have 5-column structure"""
    non_5_col_tables = []

    # Skip first few tables (title, exec summary, TOC)
    for i, table in enumerate(doc.tables[3:], start=4):
        if len(table.columns) != 5:
            non_5_col_tables.append(f"Table {i}: {len(table.columns)} columns")

    return {
        'passed': len(non_5_col_tables) == 0,
        'errors': non_5_col_tables
    }


def check_executive_summary(doc):
    """Verify executive summary exists and has required sections"""
    required_sections = [
        'PROCESS OVERVIEW',
        'PHASE BREAKDOWN',
        'TEAM RESPONSIBILITIES',
        'KEY SYSTEMS',
        'SPECIAL WORKFLOWS',
        'VERSION UPDATES'
    ]

    found_sections = []
    doc_text = '\n'.join([para.text for para in doc.paragraphs])

    for section in required_sections:
        if section in doc_text:
            found_sections.append(section)

    return {
        'passed': len(found_sections) == len(required_sections),
        'found': found_sections,
        'missing': [s for s in required_sections if s not in found_sections]
    }


def check_toc_exists(doc):
    """Verify Table of Contents exists"""
    toc_exists = False

    for para in doc.paragraphs:
        if 'TABLE OF CONTENTS' in para.text:
            toc_exists = True
            break

    return {
        'passed': toc_exists,
        'message': 'Table of Contents found' if toc_exists else 'Table of Contents missing'
    }


def main():
    """Run all validation checks"""
    print("=" * 70)
    print("V15 QA VALIDATION")
    print("=" * 70)
    print()

    # Load documents
    print("Loading documents...")
    if not os.path.exists(V15_PATH):
        print(f"ERROR: V15 not found at {V15_PATH}")
        return

    if not os.path.exists(V13_PATH):
        print(f"ERROR: V13 not found at {V13_PATH}")
        return

    v15_doc = Document(V15_PATH)
    v13_doc = Document(V13_PATH)
    print("[OK] Documents loaded")
    print()

    # Run checks
    checks = []

    # 1. Activity 6 numbering
    print("Check 1: Activity 6 Numbering...")
    result = check_activity_6_numbering(v15_doc)
    checks.append(('Activity 6 numbering (6.a-6.f)', result['passed']))
    if result['passed']:
        print(f"  [PASS] Found {result['correct_count']} correctly numbered Activity 6 items")
    else:
        print(f"  [FAIL] Errors: {result['errors']}")
    print()

    # 2. Activity 9 gap
    print("Check 2: Activity 9 Gap Documentation...")
    result = check_activity_9_gap(v15_doc)
    checks.append(('Activity 9 gap documented', result['passed']))
    print(f"  [{'PASS' if result['passed'] else 'FAIL'}] {result['message']}")
    print()

    # 3. Table count
    print("Check 3: Table Count Comparison...")
    result = check_table_count(v13_doc, v15_doc)
    checks.append(('Table count preserved', result['passed']))
    print(f"  [{'PASS' if result['passed'] else 'FAIL'}] {result['message']}")
    print()

    # 4. 5-column structure
    print("Check 4: 5-Column Table Structure...")
    result = check_5_column_structure(v15_doc)
    checks.append(('5-column structure', result['passed']))
    if result['passed']:
        print("  [PASS] All process tables have 5 columns")
    else:
        print(f"  [FAIL] Non-5-column tables: {result['errors']}")
    print()

    # 5. Executive summary
    print("Check 5: Executive Summary Sections...")
    result = check_executive_summary(v15_doc)
    checks.append(('Executive summary complete', result['passed']))
    if result['passed']:
        print(f"  [PASS] All {len(result['found'])} sections present")
    else:
        print(f"  [FAIL] Missing sections: {result['missing']}")
    print()

    # 6. Table of Contents
    print("Check 6: Table of Contents...")
    result = check_toc_exists(v15_doc)
    checks.append(('Table of Contents exists', result['passed']))
    print(f"  [{'PASS' if result['passed'] else 'FAIL'}] {result['message']}")
    print()

    # Summary
    print("=" * 70)
    print("VALIDATION SUMMARY")
    print("=" * 70)
    print()

    passed = sum(1 for _, result in checks if result)
    total = len(checks)

    for check_name, result in checks:
        status = "[PASS]" if result else "[FAIL]"
        print(f"{status} {check_name}")

    print()
    print(f"Overall: {passed}/{total} checks passed")

    if passed == total:
        print()
        print("[OK] V15 validation PASSED - Ready for review")
    else:
        print()
        print("[WARNING] Some checks failed - Review required")

    print()


if __name__ == "__main__":
    main()
