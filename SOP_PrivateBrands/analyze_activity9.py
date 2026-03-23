import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

def get_full_cell_text(cell):
    return '\n'.join(p.text for p in cell.paragraphs)

# Load both docs
jots = Document(r'D:\SOP_PrivateBrands\Documents\JSWOrderLogging_JOTS_modification.docx')
v7 = Document(r'D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V7.docx')

print("=" * 70)
print("V7 - Activity 9 (Table 11, Row 2) - FULL TEXT")
print("=" * 70)
for cell in v7.tables[10].rows[2].cells:
    print(f"\n--- Cell ---")
    print(get_full_cell_text(cell))

print("\n" + "=" * 70)
print("JOTS - Activity 9 (Table 11, Row 2) - FULL TEXT")
print("=" * 70)
for cell in jots.tables[10].rows[2].cells:
    print(f"\n--- Cell ---")
    print(get_full_cell_text(cell))

print("\n" + "=" * 70)
print("V7 - Activity 8 (Table 11, Row 1) - FULL TEXT")
print("=" * 70)
for cell in v7.tables[10].rows[1].cells:
    print(f"\n--- Cell ---")
    print(get_full_cell_text(cell))

print("\n" + "=" * 70)
print("JOTS - Activity 8 (Table 11, Row 1) - FULL TEXT")
print("=" * 70)
for cell in jots.tables[10].rows[1].cells:
    print(f"\n--- Cell ---")
    print(get_full_cell_text(cell))

# Now print JOTS Table 9 fully to see overlap
print("\n" + "=" * 70)
print("JOTS - Table 9 (JOTS Transportation) - FULL TEXT")
print("=" * 70)
for i, row in enumerate(jots.tables[8].rows):
    print(f"\n--- Row {i} ---")
    for j, cell in enumerate(row.cells):
        col_names = ['Activity', 'Steps', 'Team', 'Interface', 'Sign off']
        print(f"  [{col_names[j]}]: {get_full_cell_text(cell)}")

print("\n" + "=" * 70)
print("V7 - Table 9 (JOTS Transportation) - FULL TEXT")
print("=" * 70)
for i, row in enumerate(v7.tables[8].rows):
    print(f"\n--- Row {i} ---")
    for j, cell in enumerate(row.cells):
        col_names = ['Activity', 'Steps', 'Team', 'Interface', 'Sign off']
        print(f"  [{col_names[j]}]: {get_full_cell_text(cell)}")

# Also check Table 10 (Plant Ops) for any DO references
print("\n" + "=" * 70)
print("V7 - Table 10 (Plant Operations) - Row 1 & 2")
print("=" * 70)
for i in [1, 2]:
    print(f"\n--- Row {i} ---")
    for j, cell in enumerate(v7.tables[9].rows[i].cells):
        col_names = ['Activity', 'Steps', 'Team', 'Interface', 'Sign off']
        print(f"  [{col_names[j]}]: {get_full_cell_text(cell)}")
