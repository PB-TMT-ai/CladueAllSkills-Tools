"""
JSW Order Logging V14 Generator
Rebuilds DOCX from V13 extracted content with corrected index
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_borders(table):
    """Add borders to all cells in table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_5col_table(doc, rows_data):
    """Create a 5-column table with header and data rows"""
    table = doc.add_table(rows=len(rows_data), cols=5)
    table.style = 'Table Grid'
    add_table_borders(table)

    # Set column widths
    widths = [Inches(0.8), Inches(3.5), Inches(1.2), Inches(1.8), Inches(0.7)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # Populate cells
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text

            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].bold = True
                    paragraph.runs[0].font.size = Pt(11)
                # Gray background for header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9D9D9')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # Bold activity numbers
                if col_idx == 0 and cell_text.strip():
                    for paragraph in cell.paragraphs:
                        if paragraph.runs:
                            paragraph.runs[0].bold = True

    return table

# Create document
doc = Document()

# Title Page
title = doc.add_heading('JSW ONE TMT', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].bold = True

subtitle1 = doc.add_paragraph('STANDARD OPERATING PROCEDURE')
subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle1.runs[0].font.size = Pt(14)
subtitle1.runs[0].bold = True

subtitle2 = doc.add_paragraph('Order Logging Process')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(13)

doc.add_paragraph()
complete_doc = doc.add_paragraph('Complete End-to-End Process Documentation')
complete_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
complete_doc.runs[0].font.size = Pt(11)

doc.add_page_break()

# Table of Contents
toc_heading = doc.add_heading('Table of Contents', level=1)
toc_heading.runs[0].font.size = Pt(16)
toc_heading.runs[0].bold = True

# PRE-ORDER PHASE
phase1 = doc.add_heading('PRE-ORDER PHASE', level=2)
phase1.runs[0].font.size = Pt(13)
doc.add_paragraph('Section A: FOR Orders (Freight On Road)', style='List Bullet')
doc.add_paragraph('Section A.1: Special Cases - RRP Breach Approvals', style='List Bullet 2')
doc.add_paragraph('Section A.2: Special Cases - Delivery Instruction Changes', style='List Bullet 2')
doc.add_paragraph('Section A.3: Special Cases - Channel Finance Orders', style='List Bullet 2')
doc.add_paragraph('Section B: Ex-works Orders', style='List Bullet')
doc.add_paragraph('Section C: Project / PTR Orders', style='List Bullet')

# ORDER PHASE
phase2 = doc.add_heading('ORDER PHASE', level=2)
phase2.runs[0].font.size = Pt(13)
doc.add_paragraph('Section D - Phase 1: Planning Team Coordination', style='List Bullet')
doc.add_paragraph('Section D - Phase 2: Biz-ops Team Execution', style='List Bullet')
doc.add_paragraph('Section D - Phase 3: JOTS Transportation', style='List Bullet')

# POST-ORDER PHASE
phase3 = doc.add_heading('POST-ORDER PHASE', level=2)
phase3.runs[0].font.size = Pt(13)
doc.add_paragraph('Section D - Phase 4: Plant Operations (Dispatch)', style='List Bullet')
doc.add_paragraph('Section D - Phase 5: Post-Dispatch Activities (GRN & Invoice Posting) [NEW]', style='List Bullet')

doc.add_page_break()

# ==================== PRE-ORDER PHASE ====================
doc.add_heading('PRE-ORDER PHASE', level=1)
intro = doc.add_paragraph('This phase covers order type selection and opportunity creation across all distribution channels: FOR Orders, Ex-works Orders, and Project/PTR Orders.')
intro.runs[0].font.size = Pt(11)

doc.add_page_break()

# SECTION A: FOR Orders
doc.add_heading('SECTION A: FOR Orders (Freight On Road)', level=2)
note_para = doc.add_paragraph('IMPORTANT: Activity 1.f (Freight Approval Request) has been REMOVED from this section as it is not required for retail FOR orders. Order ID is generated automatically upon opportunity submission.')
note_para.runs[0].italic = True
note_para.runs[0].font.size = Pt(10)

table_a = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['1.', "Login to JSW ONE TMT Distributor Portal. Click on the 'Opportunities' tab from the main navigation menu.", 'Sales', 'Distributor Portal', ''],
    ['1.a.', "On the Opportunities page, click on the 'Create new opportunity' button located in the top-right corner.", 'Sales', 'Distributor Portal', ''],
    ['1.b.', "In the 'Create opportunity' wizard (Step 1/4), select 'Order type' from dropdown. Choose 'Retailer' for FOR orders. Click 'Next'.", 'Sales', 'Distributor Portal', ''],
    ['1.c.', "Select specifications: Grade (Fe 500/550), Size (8-32mm), Form. Click 'Next'.", 'Sales', 'Distributor Portal', ''],
    ['1.d.', 'System reflects pricing based on selection. Verify unit price per MT, total pricing, discounts. Confirm accuracy.', 'Sales', 'Distributor Portal', ''],
    ['1.e.', 'Complete required fields: delivery location, expected date, payment terms. Review and submit. Order ID generated automatically upon approval.', 'Sales', 'Distributor Portal', '']
])

doc.add_paragraph()
note2 = doc.add_paragraph('Note: Original activities 1.f, 1.f.i, and 1.f.ii (Freight Approval workflow) have been removed as they are not part of the retail FOR order process. The order proceeds directly to Section D (Order Processing) after opportunity submission.')
note2.runs[0].italic = True
note2.runs[0].font.size = Pt(10)

doc.add_page_break()

# Section A.1: RRP Breach Approval
doc.add_heading('Section A.1: Special Cases - RRP Price Breach Approval', level=2)
intro_a1 = doc.add_paragraph('Order approval is triggered automatically when business rules are violated. Three trigger cases exist: Lower RRP breaches, Upper RRP breaches, and delivery instruction changes.')
intro_a1.runs[0].font.size = Pt(11)

doc.add_paragraph()
doc.add_heading('Trigger Case 1: Breach of Lower RRP', level=3)

table_a1_case1 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['8.a.', 'System detects quoted price is BELOW minimum Recommended Retail Price (Lower RRP) for the SKU. Order submission is automatically blocked.', 'Sales', 'Distributor Portal', ''],
    ['8.a.i.', "Opportunity status automatically changes to 'Awaiting Approval from JSW One'. System sends notification to Regional Sales Manager (RSM) and Category Head.", 'Sales', 'SFDC', ''],
    ['8.a.ii.', "Open opportunity in SFDC. Navigate to 'Approval Request' section. Document price variance reason: competitive pressure, volume discount, or strategic account retention.", 'Sales', 'Distributor Portal', ''],
    ['8.a.iii.', 'Upload supporting documents: competitor quote (if available), customer PO, market analysis. Specify quantity (MT) and expected margin impact.', 'Sales', 'SFDC', ''],
    ['8.a.iv.', "Click 'Submit for Approval' button. System automatically routes to Regional Sales Manager. Timeline: RSM must respond within 24 hours.", 'Sales', 'SFDC', ''],
    ['8.b.', 'RSM receives notification via SFDC and email. Review customer credit history, past order patterns, current account standing, and sales justification.', 'Sales - RSM', 'SFDC', ''],
    ['8.b.i.', 'Check pricing calculation accuracy. Verify margin impact on regional sales targets. Review customer relationship value.', 'Sales - RSM', 'SFDC', ''],
    ['8.b.ii.', 'Select decision: APPROVE (order proceeds to Activity 5.f), REJECT (returns to Sales with feedback), or ESCALATE (forwards to Category Head).', 'Sales - RSM', 'SFDC', ''],
    ['8.c.', "System updates opportunity status to 'Approved'. Order proceeds to Biz-ops team for confirmation (Activity 5.f). Approved price is locked. Audit trail created with approver name and timestamp.", 'Sales', '', ''],
    ['8.c.i.', 'Sales team receives confirmation notification. Order ID generation can now proceed. Continue to standard order processing workflow.', 'Sales', '', ''],
    ['8.d.', "Opportunity status changes to 'Pricing Revision Required'. Sales team receives notification with detailed rejection reason from RSM.", 'Sales', 'SFDC', ''],
    ['8.d.i.', 'Choose action: (a) Revise price to meet minimum RRP, (b) Provide additional justification and resubmit, or (c) Escalate to Category Head with business case.', 'Sales', 'SFDC', ''],
    ['8.d.ii.', 'Document customer communication about price revision. Update opportunity with new approach. Timeline: Revise and resubmit within 48 hours or opportunity auto-closes.', 'Sales', 'SFDC + Customer Email', ''],
    ['8.d.iii.', 'Update opportunity with revised pricing or additional justification. Resubmit for approval (repeat from Activity 8.a.iv).', 'Sales', 'SFDC', '']
])

doc.add_page_break()

# Trigger Case 2: Upper RRP
doc.add_heading('Trigger Case 2: Breach of Upper RRP', level=3)

table_a1_case2 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['8.e.', 'System detects quoted price EXCEEDS maximum Recommended Retail Price (Upper RRP). Order submission blocked. Purpose: Prevent customer overcharging and maintain market competitiveness.', 'Sales', 'SFDC', ''],
    ['8.e.i.', "Opportunity flagged for 'Price Overcharge Review'. Notification sent to Category Head and Pricing Team. Status: 'Awaiting Premium Pricing Approval'.", 'Sales', 'SFDC', ''],
    ['8.e.ii.', "Access opportunity in SFDC. Complete 'Premium Pricing Justification Form' with: special customer requirements, expedited delivery timeline, custom fabrication, or market shortage conditions.", 'Sales', 'SFDC', ''],
    ['8.e.iii.', 'Attach customer PO showing price acceptance. Provide competitor pricing data if premium is market-standard. Confirm customer willingness in writing.', 'Sales', 'SFDC', ''],
    ['8.e.iv.', 'Submit for Category Head review. Timeline: Decision required within 12 hours (premium pricing orders are often time-sensitive).', 'Sales', 'SFDC', ''],
    ['8.f.', 'Verify pricing against current cost structure: raw material costs, logistics constraints, special processing. Validate customer agreement and PO authenticity.', 'Sales - Category Head', 'SFDC + Pricing System', ''],
    ['8.f.i.', 'Assess impact on customer relationship and future order potential. Consult with Pricing Team if needed. Check profitability margin.', 'Sales - Category Head', 'SFDC', ''],
    ['8.f.ii.', 'Select: APPROVE (order proceeds with premium pricing), REVISE (suggest adjusted price within range), or REJECT (require repricing to standard RRP).', 'Sales - Category Head', 'SFDC', ''],
    ['8.g.', "Opportunity status: 'Approved - Premium Pricing'. Approved premium price locked in system. Special flag added for tracking. Finance team notified for revenue recognition.", 'Sales', 'SFDC', ''],
    ['8.g.i.', 'Order proceeds to Biz-ops (Activity 5.f). Premium pricing orders reviewed in monthly Category meetings. Track customer feedback and repeat order rates.', 'Sales', '', ''],
    ['8.h.', "Status: 'Pricing Correction Required'. Sales receives adjusted price recommendation from Category Head. Must communicate revised pricing to customer.", 'Sales', 'SFDC', ''],
    ['8.h.i.', 'Communicate revised pricing professionally. Highlight value proposition and competitive rates. Offer alternatives: volume discounts, extended payment terms.', 'Sales', 'Customer Communication', ''],
    ['8.h.ii.', 'Obtain customer approval for adjusted price. Update opportunity with customer confirmation. Resubmit for standard approval process.', 'Sales', 'SFDC + Email', '']
])

doc.add_page_break()

# Section A.2: Delivery Instruction Changes
doc.add_heading('Section A.2: Special Cases - Delivery Instruction Changes', level=2)
doc.add_heading('Trigger Case 3: Change in Delivery Instruction', level=3)

table_a2 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['8.i.', 'Customer requests modification to delivery location, delivery date, or shipping instructions AFTER initial opportunity approval. Common scenarios: site address change, revised timeline, split delivery.', 'Sales', 'Customer Communication', ''],
    ['8.i.i.', 'If order already in processing, place on hold. System sends notification to Sales, Planning, and Biz-ops teams. Freight cost recalculation triggered automatically.', 'Sales', 'SFDC', ''],
    ['8.i.ii.', "Open opportunity in SFDC. Navigate to 'Change Request' tab. Select change type: Delivery Location, Delivery Date, Delivery Mode, or Split Delivery.", 'Sales', 'SFDC', ''],
    ['8.i.iii.', 'Enter: Original delivery instruction, New delivery instruction, Reason for change (customer relocation, site readiness, urgent need). Upload customer PO amendment if applicable.', 'Sales', 'SFDC', ''],
    ['8.i.iv.', "Upload customer approval document for delivery change. Click 'Initiate Change Request'. System routes to Planning team automatically.", 'Sales', 'SFDC', ''],
    ['8.j.', 'Planning team receives change request notification. Check new delivery location plant proximity. Calculate revised freight cost. Verify inventory at alternate plant if needed.', 'Planning', 'Planning System + SFDC', ''],
    ['8.j.i.', 'Assess operational impact. Calculate: Additional freight charges, Plant shift costs (if applicable), Expedite charges (for urgent delivery). Prepare revised quote with breakdown.', 'Planning', 'Planning System', ''],
    ['8.j.ii.', 'Coordinate with JOTS team for logistics confirmation on new delivery location. Assess delivery date feasibility. Submit findings to Sales team.', 'Planning', 'JOTS Coordination', ''],
    ['8.k.', 'Receive revised quote from Planning. Contact customer with updated pricing: Original freight: Rs. X, Revised freight: Rs. Y, Differential: Rs. (Y-X), Total revised order value.', 'Sales', 'SFDC + Customer Call', ''],
    ['8.k.i.', 'Obtain customer approval via email confirmation, revised PO, or signed change order form. Upload customer approval to SFDC. Update opportunity with revised pricing.', 'Sales', 'SFDC + Email', ''],
    ['8.k.ii.', 'If customer declines revised pricing: Option 1 - Revert to original delivery instruction, Option 2 - Cancel order if change is mandatory. Document decision and reason.', 'Sales', 'SFDC', ''],
    ['8.l.', 'Cost increase < 5%: Auto-approved by system, proceeds to Biz-ops. Cost increase 5-10%: Requires Sales Manager approval (6-hour timeline). Cost increase > 10%: Requires Category Head approval (12-hour timeline).', 'Sales / Sales Manager / Category Head', 'SFDC Approval Workflow', ''],
    ['8.l.i.', 'Sales Manager reviews: Customer relationship value, Order profitability, Strategic account importance. Approves or rejects within 6 hours.', 'Sales Manager', 'SFDC', ''],
    ['8.l.ii.', 'Category Head conducts full business case review. Consider strategic account value. Timeline: 12 hours. Approves, requests revision, or rejects.', 'Category Head', 'SFDC', ''],
    ['8.m.', 'After approval, system updates opportunity with new delivery instruction. Biz-ops team notified to update ERP order. JOTS receives updated delivery details. Planning adjusts freight allocation.', 'Biz-ops', 'SFDC + ERP', ''],
    ['8.m.i.', 'Access order in ERP system. Update: Delivery location address, Delivery date, Freight charges, Shipping instructions. Generate revised order confirmation.', 'Biz-ops', 'ERP', ''],
    ['8.m.ii.', 'Share revised order confirmation with: Customer (email), JOTS team (for vehicle planning update), Plant team (for dispatch scheduling), Sales team (for records).', 'Biz-ops', 'Email + Zoho', ''],
    ['8.m.iii.', "Update Zoho with new delivery order details. Mark change request as 'Completed' in SFDC. All changes logged in audit trail for monthly analysis.", 'Biz-ops', 'Zoho + SFDC', '']
])

doc.add_page_break()

# Section A.3: Channel Finance
doc.add_heading('Section A.3: Special Cases - Channel Finance Orders', level=2)
intro_a3 = doc.add_paragraph('Channel finance enables select customers to procure steel on extended credit terms through financial partners. Currently 3 customers enrolled. Dual approval required: Distributor + Financial Partner.')
intro_a3.runs[0].font.size = Pt(11)

doc.add_page_break()

# SECTION B: Ex-works Orders
doc.add_heading('SECTION B: Ex-works Orders', level=2)

table_b = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['2.', "Login to JSW ONE TMT Distributor Portal. Click 'Opportunities' tab.", 'Sales', 'SFDC + Email', ''],
    ['2.a.', "Click 'Create new opportunity' button.", 'Sales', 'Zoho / Email', ''],
    ['2.b.', "In wizard Step 1/4, select 'Self-stocking' for Ex-works orders. Click 'Next'.", 'Sales', 'SFDC', ''],
    ['2.c.', "Select specifications: Grade, Size, Form as per requirement. Click 'Next'.", 'Sales', 'SFDC', ''],
    ['2.d.', 'System reflects Ex-works pricing (no freight). Verify pricing accuracy.', 'Sales', 'ERP', ''],
    ['2.e.', 'Complete all fields and submit. No freight approval required for Ex-works orders.', 'Sales', 'ERP', ''],
    ['2.f.', 'Upon approval, Order ID generated automatically. Proceed to Section D (Order Processing).', 'Sales', 'ERP', '']
])

doc.add_page_break()

# SECTION C: Project / PTR Orders
doc.add_heading('SECTION C: Project / PTR Orders', level=2)

table_c = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['3.', 'Create customer account in Salesforce if not existing. Enter customer details, site address, business segment.', 'Sales', 'Shared Excel Sheet', ''],
    ['3.a.', 'Search and locate customer account in SFDC system.', 'Sales', 'Zoho', ''],
    ['3.b.', 'Create new opportunity record for project order in SFDC.', 'Sales', 'Zoho', ''],
    ['3.c.', 'Add complete site/delivery address in opportunity record.', 'Sales', 'Zoho', ''],
    ['3.d.', 'Pass opportunity to HO team for entering: SKU-wise price, delivery instructions, price type, avg. selling price/MT, avg. procurement price/MT.', 'Sales', 'Zoho + Internal System', ''],
    ['3.e.', 'HO team adds pricing. Coordinate with Planning for freight calculation. Add freight details. | Coordination is done though excel sheets', 'Biz Ops', 'Excel', ''],
    ['3.f.', 'Submit opportunity with freight details for approval.', 'Biz Ops', 'SFDC', ''],
    ['3.f.i.', 'Upon approval, Order ID generated in SFDC. Proceed to Section D.', 'Sales', 'SFDC', ''],
    ['3.f.ii.', 'Opportunity returned. Revise freight coordination between JOT and Biz Ops. Resubmit (repeat 3.f).', 'Biz Ops', 'SFDC', '']
])

doc.add_page_break()

# ==================== ORDER PHASE ====================
doc.add_heading('ORDER PHASE', level=1)

# Section D - Phase 1
doc.add_heading('Section D - Phase 1: Planning Team Coordination', level=2)

table_d1 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['4.', 'Analyze: previous sales data, PO history, category mix, product distribution, current inventory at plants.', 'Planning', 'Excel / Planning Systems', ''],
    ['4.a.', 'Assess inventory shortage or urgent order requirements across all plant locations.', 'Biz Ops', 'ERP + Excel', ''],
    ['4.b.', 'Suggest optimal plants based on: inventory availability, cost optimization (PO + freight), new distribution locations.', 'Planning', 'Internal Planning Sheet', ''],
    ['4.c.', 'Coordinate with pricing team to finalize price structure.', 'Planning', 'Email / SFDC', ''],
    ['4.d.', 'Coordinate with JOTS team for accurate freight rates and logistics planning.', 'Biz Ops', 'Shared Freight Sheet', ''],
    ['4.e.', 'Coordinate with Product team for timely inventory resolution at selected plant.', 'Biz Ops', 'Email / Meeting', ''],
    ['4.f.', 'Update freight details on Planning sheet for Biz-ops reference.', 'Biz Ops', 'Shared Excel Sheet', '']
])

doc.add_page_break()

# Section D - Phase 2
doc.add_heading('Section D - Phase 2: Biz-ops Team Execution', level=2)

table_d2 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['5.', 'Coordinate with Planning for plant shift orders based on inventory optimization.', 'Biz-ops', 'SFDC + Email', ''],
    ['5.a.', 'Biz-ops coordinates with JOTS for freight arrangement and vehicle planning.', 'Biz-ops', 'Zoho / Email', ''],
    ['5.b.', 'Calculate and enter freight costs in SFDC order record from Planning sheet.', 'Biz-ops', 'SFDC', ''],
    ['5.c.', 'System auto-generates Order ID in SFDC (not created by Biz Ops).', 'Biz-ops', 'SFDC', ''],
    ['5.d.', 'Verify generated order is visible in ERP system.', 'Biz-ops', 'ERP', ''],
    ['5.e.', 'Conduct checks: Inventory Ledger verification, Bill-to address accuracy, Order quantity validation.', 'Biz-ops', 'ERP', ''],
    ['5.f.', 'Confirm order in ERP after all checks passed.', 'Biz-ops', 'ERP', ''],
    ['5.g.', 'Verify key fields post-confirmation: Customer Details, Business Segment & Supply source, SKU & Pricing.', 'Biz-ops', 'ERP', ''],
    ['5.h.', 'Release delivery order quantity in ERP.', 'Biz-ops', 'ERP', ''],
    ['5.i.', 'Complete delivery order release and update system status.', 'Biz-ops', 'ERP', ''],
    ['5.j.', 'Share system-generated delivery order with JOTS for vehicle planning.', 'Biz-ops', 'Zoho + Email', '']
])

doc.add_page_break()

# Section D - Phase 3
doc.add_heading('Section D - Phase 3: JOTS Transportation', level=2)

table_d3 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['6.1', '-Receive the system-generated Delivery Order (DO) from Biz Ops on TMS.', 'JOTS', 'TMS', ''],
    ['6.2', '-Verify the order type to confirm whether it is FOR / Self-Stocking or Project.', 'JOTS', 'TMS', ''],
    ['6.3', '-For Project orders, refer to the freight details updated by VM in the shared freight sheet.', 'JOTS', 'Shared Excel Sheet', ''],
    ['6.4', '-Release the Delivery Order (DO) to the transporter for vehicle alignment. | - Coordinate with the transporter to obtain vehicle details and SIM consent. | -Coordinate with the transporter and CM for vehicle gate-in, and release the Plant DO for gate-in.', 'JOTS', 'TMS', ''],
    ['6.5', '-After loading, perform pre-dispatch checks, including tarpaulin coverage, material condition, seal integrity, etc. | -Upload pre-dispatch documents on TMS (e.g., weighment slip, LR, etc.). | Hand over the invoice to the driver and process the dispatch.', 'JOTS', 'TMS', ''],
    ['6.6', '-Post-dispatch, track the vehicle on TMS and manually follow up in case of any issues. | -Post-reporting, handle weighment discrepancies and any unloading-related issues, if applicable. | -Upload post-unloading documents on TMS, such as POD.', 'JOTS', 'TMS', '']
])

doc.add_page_break()

# ==================== POST-ORDER PHASE ====================
doc.add_heading('POST-ORDER PHASE', level=1)

# Section D - Phase 4
doc.add_heading('Section D - Phase 4: Plant Operations (Dispatch)', level=2)

table_d4 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['7.', 'Receive freight order details generated by Biz-ops. | DO is shared on WhatsApp group for the respective plant .', 'JOTS', 'Freight Tiger TMS', ''],
    ['7.a.', 'Receive system-generated or Zoho-generated DO from JOTS team.', 'JOTS', 'Plant Portal / Zoho', ''],
    ['7.b.', 'Generate: Invoice (TC - Test Certificate), E-Way Bill for interstate transportation.', 'Plant Operations', 'ERP/Zoho', ''],
    ['7.c.', 'Push Invoice Registration Portal (IRP) data and generate E-Way Bill.', 'Plant Operations', 'IRP Portal / E-Way Bill System', ''],
    ['7.d.', 'Complete ERP /Zoho invoice approval process.', 'Plant Operations', 'ERP', ''],
    ['7.e.', 'Conduct final dispatch with Tax Challan, Test Certificate , Eway Bill , Bilty', 'Plant Operations', 'Physical + ERP', '']
])

doc.add_page_break()

# Section D - Phase 5
doc.add_heading('Section D - Phase 5: Post-Dispatch Activities (GRN & Invoice Posting) [NEW]', level=2)
intro_d5 = doc.add_paragraph('This NEW phase covers all post-dispatch operations including Delivery Order system integration, vehicle coordination, GRN entry, invoice generation, ERP approval, and order closure. This phase represents 25 new activities (8-16 with sub-activities) integrated from PB-Ops & Biz-ops procedures.')
intro_d5.runs[0].font.size = Pt(11)

doc.add_paragraph()

table_d5_act8 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['8.', 'DO Release & System Integration: Once Delivery Order (DO) released in ERP, it automatically reflects as "Released" against Order ID. Visible in BQ Union Book and FT TMS System.BQ Union Link - Union BookFT TMS Link - FT TMS SystemNote: JOTS SOP provided separately by JOTS Team.', 'Biz-ops', 'ERP, Union Book, FT TMS System', ''],
    ['', '', '', '', '']
])

doc.add_paragraph()
doc.add_heading('Activity 10: Vehicle Arrival & Weighment Coordination', level=3)

table_d5_act10 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['10.', 'Vehicle Arrival and Weighment: Ensure proper coordination and quality control during vehicle arrival, inspection, loading, and final weighment process.', 'Plant Operations', 'Plant Gate System', ''],
    ['10.a.', 'Vehicle Arrival: Plant Ops monitor vehicle arrival at plant. Upon arrival at gate, ensure vehicle completed two weighments outside plant. Direct vehicle to third weighment inside plant.', 'Plant Operations / JOTS', 'Plant Gate, Weighbridge', ''],
    ['10.b.', 'Loading Advice Verification: Loading advice prepared by CM team. JSW ONE staff verify loading advice matches Delivery Order (DO).', 'Plant Operations', 'Loading Advice System, DO', ''],
    ['10.c.', 'Vehicle Inspection (Post Internal Weighment): Conduct thorough inspection: Confirm length sufficient for loading, Ensure vehicle body intact and dry, Verify vehicle equipped with tarpaulin and ropes. |  | Based on inspection outcome: | • If vehicle passes: proceed with loading | • If vehicle fails: escalate to JOTS |  | Capture two photos: | • Vehicle number and driver in single frame | • Vehicle body image', 'Plant Operations', 'Inspection Checklist, Camera', '']
])

doc.add_paragraph()
doc.add_heading('Activity 11: Loading Process Execution', level=3)

table_d5_act11 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['11.', 'Loading Process: Load materials strictly as per the Loading Advice.', 'Plant Operations', 'Loading Advice, Weighbridge', ''],
    ['11.a.', 'Load materials strictly per Loading Advice. Ensure loading done within specified tolerance limits.', 'Plant Operations', 'Loading Advice', ''],
    ['11.b.', 'Perform weighment for each SKU loaded. Record weights accurately in system.', 'Plant Operations', 'Weighbridge System', ''],
    ['11.c.', 'Post-Loading Steps: After loading, secure material using 5-meter seals.  | Capture photo of fully loaded and sealed vehicle.  | Direct vehicle for final weighment. Submit final weighment slip to CM Dispatch Office.', 'Plant Operations', 'Sealing Equipment, Camera, Weighbridge, CM Dispatch Office', ''],
    ['11 D', 'MTC generation through Zoho after entering material details e.g. Dia, Heat no, Batch no & material quality parameters', 'Plant Operations', 'JTR SGA Test report', ''],
    ['11 D', 'LR Copy Generation', 'JOTS', "Transporter's copy", ''],
    ['11 E', 'Handing over documents to Vehicle driver- Invoice, MTC, E-Waybill & LR Copy', 'Plant Operations', '', '']
])

doc.add_page_break()

doc.add_heading('Activity 12: GRN Entry in Zoho Books (PB Plant Ops)', level=3)

table_d5_act12 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['12.', 'GRN & Invoicing Post-Final Weighment: Following successful completion of final weighment after loading process, invoice generation and finalization must begin.', 'Plant Operations', 'Zoho Books', ''],
    ['12.a.', 'Collect Seller Invoice: PB Plant Ops team must collect seller invoice from CM Dispatch Office.', 'Plant Operations', 'CM Dispatch Office', ''],
    ['12.b.', 'Perform GRN in Zoho Books: Upon receiving seller invoice, perform Goods Receipt (GRN) entry in ZOHO. |  | Path → Zoho books → Purchase → Purchase receives |  | Accurately input mandatory details: | • Seller Invoice Number | • Sales Reference Number | • Payment Terms | • Purchase Order (PO) Number (mentioned in Seller Invoice) | • Material Form', 'Plant Operations', '', ''],
    ['12.c.', 'Validate GRN vs Seller Invoice:  | Ensure GRN details strictly match Seller Invoice and Delivery Order (DO). |  |  In case of any discrepancies in Seller Invoice, discuss with CM and make corrections before proceeding with GRN entry. |  | Once Good receipt done in Zoho, search for sales order associated with order ID.', 'Plant Operations', '', '']
])

doc.add_paragraph()
doc.add_heading('Activity 13: Invoice & E-Way Bill Generation', level=3)

table_d5_act13 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['13.', 'Invoice and E-Way Bill Generation: PB-Plant Ops team responsible for creating Invoice in system and generating E-Way Bill for dispatch. |  | PB Plant Ops search Sales Order in ZOHO Books to create JODL Invoice against Seller Invoice. |  | Path – ZOHO Books → Sales → Sales Orders → SO → Convert to Invoice', 'Plant Operations', 'Zoho Books', ''],
    ['13.a.', 'Create JODL Invoice in Zoho Books: Search Sales Order and convert to invoice in Zoho Books system.', 'Plant Operations', '', ''],
    ['13.b.', 'Fill Mandatory Bills Section Fields: Following fields must be accurately filled by PB-Ops in Bills Section: | • Dispatch from Address | • Bill to and Ship to addresses – must match GST records | • JODL Invoice Date | • Payment Terms | • Sales Order Reference Number | • Business Segment | • Supply Source | • Incoming Payment details | • E-Commerce (if applicable) | • Payment Mode | • Purchase Bill Reference Number | • E-Way Bill Details: Vehicle Number, Motor Vehicle Number, Mode of Transport (Road to be selected) | • SKU and Quantity – must match Seller Invoice', 'Plant Operations', '', ''],
    ['13.c.', 'IRP Push and E-Way Bill Generation: After submitting invoice, mark "Push to IRP." |  | Based on invoice value: | • If invoice value > ₹50,000: Push both Invoice and E-Way Bill details to IRP | • If invoice value < ₹50,000: Push only Invoice details to IRP |  | After pushing to IRP, create E-Way Bill. Enter E-Way Bill details including Vehicle Number and Motor Vehicle Details. Submit E-Way Bill by selecting Push to IRP again for E-Way Bill-specific information. |  | After successful submission: Select "Mark as Sent" in system. Email invoice to Biz-Ops team for further processing.', 'Plant Operations', '', '']
])

doc.add_page_break()

doc.add_heading('Activity 14: ERP Invoice Approval (Biz-Ops)', level=3)

table_d5_act14 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['14.', 'ERP Invoice Approval: After receiving JODL invoice, Biz-Ops approves it in JOPL ERP Supply Panel. |  | JOPL ERP Supply Panel Link: JOPL ERP |  | Enter Invoice Number or Order ID.', 'Biz-ops', 'JOPL ERP Supply Panel', ''],
    ['14.a.', 'Pre-Approval Validation Checks: | • Invoice Value | • Grade and Quantity | • Bill to / Ship to Details mentioned in Invoice should match with Sales order/Union Book/GRN sheet', 'Biz-ops', '', ''],
    ['14.b.', 'Approve in JOPL ERP: After validation checks pass, approve invoice in ERP system.', 'Biz-ops', '', ''],
    ['14.c.', 'Monitor Posting Status: Monitor posting status to ensure it reaches "Success" & "Dispatched" stage. |  | If invoice is not auto-posting, raise SQAD ticket. |  | SQAD Ticket Link: SQAD Ticket', 'Biz-ops', '', '']
])

doc.add_paragraph()
doc.add_heading('Activity 15: Manual Shipment Creation (If Auto Failure)', level=3)

table_d5_act15 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['15.', 'Manual Shipment Creation Conditions: If shipment is not auto-created, proceed with manual shipment creation in ERP.', 'Biz-ops', 'ERP System', ''],
    ['15.a.', 'Conditions for Manual Creation: Determine if manual shipment creation required based on auto-creation failure.', 'Biz-ops', 'ERP System', ''],
    ['15.b.', 'Create Shipment in ERP: Copy Order ID mentioned in invoice. In ERP system: | • Navigate to Order List tab (Link: Jswonemsme/Order-list) | • Search for order using Order ID | • Go to Shipments tab | • Click "Create Shipment" | • Enter SKU-wise quantity exactly as mentioned in invoice | • Ensure Gross Value entered matches value displayed in ERP | • Add Sales Freight, if applicable', 'Biz-ops', '', ''],
    ['15.c.', 'Finalize Shipment: Change Shipment Status to "Ready to Ship". Attach Invoice to shipment record. Click "Update" to save all changes.', 'Biz-ops', '', '']
])

doc.add_page_break()

doc.add_heading('Activity 16: Order Short Closure (Manual & Auto)', level=3)

table_d5_act16 = create_5col_table(doc, [
    ['Activity', 'Steps', 'Team', 'Interface / Utilities used', 'Sign off'],
    ['16.', 'Order Short Closure: Process for closing orders with quantity variance.', 'Biz-ops', 'JOPL ERP', ''],
    ['16.a.', 'Manual Short Closure (<90%): Once invoice uploaded in ERP system, if delivered quantity < 90% of ordered quantity, remaining quantity must be manually short closed after obtaining approval from respective Business Head. |  | Path – JOPL ERP → Order ID → Dispatch order → Details → Request for DO Short Closure → Short close request raise to Biz-ops team → Biz-ops team approves request |  | Post Approval from respective Business holders, select short close SKU & close Qty in system.', 'Biz-ops', '', ''],
    ['16.b.', 'Auto Short Closure (≥95%): If delivered quantity is above 95% or more of ordered quantity, system will automatically short-close pending quantity in ERP system.', 'Biz-ops', 'JOPL ERP (Auto)', '']
])

doc.add_page_break()

# ==================== GLOSSARY ====================
doc.add_heading('Glossary', level=1)

glossary_items = [
    ('ERP', 'JOPL ERP (Enterprise Resource Planning)'),
    ('SKU', 'Stock Keeping Unit'),
    ('DO', 'Dispatch Order / Delivery Order'),
    ('PO', 'Purchase Order'),
    ('SO', 'Sales Order'),
    ('GRN', 'Goods Receipt Note'),
    ('IRP', 'Invoice Registration Portal'),
    ('SQAD Ticket', 'JSW One Tech Ticketing System'),
    ('TDS', 'Tax Deducted at Source'),
    ('TCS', 'Tax Collected at Source'),
    ('JOTS', 'JSW ONE Transportation System'),
    ('CM', 'Commercial Team'),
    ('SFDC', 'Salesforce'),
    ('RSM', 'Regional Sales Manager'),
    ('RRP', 'Recommended Retail Price'),
    ('PB Plant Ops', 'Plant Book Plant Operations'),
    ('TC', 'Tax Challan'),
    ('HO', 'Head Office'),
    ('FOR', 'Freight On Road'),
    ('PTR', 'Project/Trader')
]

for term, definition in glossary_items:
    p = doc.add_paragraph()
    p.add_run(f'{term}').bold = True
    p.add_run(f' – {definition}')
    p.paragraph_format.left_indent = Inches(0.25)

# Save document
output_path = r'D:\SOP_PrivateBrands\JSW_OrderLogging_V14.docx'
doc.save(output_path)

print(f"Document generated successfully: {output_path}")
print(f"\nDocument Statistics:")
print(f"  - Total Sections: 9")
print(f"  - Total Activities: 15 (1-8, 10-16)")
print(f"  - Glossary Terms: {len(glossary_items)}")
print(f"  - Format: 5-column tables throughout")
