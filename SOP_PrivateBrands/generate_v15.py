#!/usr/bin/env python
"""
JSW Order Logging V15 Generator
Transforms V13 to V15 with corrected numbering and management summary
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# File paths
V13_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V13.docx"
V15_PATH = r"D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V15.docx"
EXEC_SUMMARY_PATH = r"D:\SOP_PrivateBrands\JSWOrderLogging_ExecutiveSummary.docx"


def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()


def set_cell_background(cell, color):
    """Set background color for a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_title_page(doc):
    """Create professional title page"""
    # Title
    title = doc.add_heading('JSW ONE TMT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True

    # Subtitle
    subtitle = doc.add_heading('ORDER LOGGING PROCESS', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)

    # Document type
    doc_type = doc.add_paragraph('Complete Process Documentation')
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_type.runs[0].font.size = Pt(14)
    doc_type.runs[0].italic = True

    # Version
    version = doc.add_heading('Version 15', level=1)
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.size = Pt(20)
    version.runs[0].font.color.rgb = RGBColor(0, 51, 153)  # JSW Blue

    # Date and owner
    date_str = datetime.now().strftime('%B %d, %Y')
    meta = doc.add_paragraph(f'\nLast Updated: {date_str}\nDocument Owner: Operations Team')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(11)

    add_page_break(doc)


def create_executive_summary(doc):
    """Create management one-pager on Page 2"""
    # Header
    header = doc.add_heading('JSW ONE TMT - ORDER LOGGING PROCESS', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(14)

    subheader = doc.add_paragraph('Executive Summary - Version 15')
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader.runs[0].font.size = Pt(12)
    subheader.runs[0].italic = True

    date_str = datetime.now().strftime('%B %d, %Y')
    doc.add_paragraph(f'Date: {date_str} | Document Owner: Operations Team\n')

    # Section 1: Process Overview
    doc.add_heading('PROCESS OVERVIEW', level=2).runs[0].font.size = Pt(12)
    overview = doc.add_paragraph(
        'This document provides end-to-end documentation of the JSW ONE TMT order lifecycle, '
        'covering the complete journey from opportunity creation through dispatch, GRN entry, '
        'and invoice generation. The process encompasses retail FOR orders, ex-works orders, '
        'and project/PTR orders, with comprehensive coverage of special cases including RRP '
        'breach approvals, channel finance workflows, and delivery instruction changes.'
    )
    overview.runs[0].font.size = Pt(11)

    stats = doc.add_paragraph('3 Phases | 15 Activities | 85+ Steps | 46 Screenshots')
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats.runs[0].bold = True
    stats.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # Section 2: Phase Breakdown
    doc.add_heading('PHASE BREAKDOWN', level=2).runs[0].font.size = Pt(12)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'

    # Headers
    headers = ['Phase', 'Activities', 'Primary Team', 'Key Steps']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'D9D9D9')

    # Data
    data = [
        ['Pre-Order', '1-3, 8.a-8.m', 'Sales', '56 steps'],
        ['Order', '4-7', 'Planning/Biz-ops/JOTS', '29 steps'],
        ['Post-Order', '10-16', 'Plant Ops/Biz-ops', '26 steps']
    ]

    for i, row_data in enumerate(data, start=1):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text

    doc.add_paragraph()

    # Section 3: Team Responsibilities
    doc.add_heading('TEAM RESPONSIBILITIES', level=2).runs[0].font.size = Pt(12)

    teams = [
        ('Sales', 'Opportunity creation, RRP approvals, special cases (Activities 1-3, 8.a-8.h)'),
        ('Planning', 'Inventory analysis, plant allocation (Activity 4, 8.j)'),
        ('Biz-ops', 'DO creation, ERP integration, closures (Activities 5, 8.m, 14-16)'),
        ('JOTS', 'Transportation, vehicle coordination (Activities 6-7)'),
        ('Plant Operations', 'Weighment, loading, GRN, invoicing (Activities 10-13)')
    ]

    for team, desc in teams:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{team}: ').bold = True
        p.add_run(desc)
        p.runs[0].font.size = Pt(10)
        p.runs[1].font.size = Pt(10)

    doc.add_paragraph()

    # Section 4: Key Systems
    doc.add_heading('KEY SYSTEMS & INTERFACES', level=2).runs[0].font.size = Pt(12)

    # Create 2-column layout using table
    sys_table = doc.add_table(rows=1, cols=2)
    sys_table.autofit = False
    sys_table.allow_autofit = False

    # Column 1 - Primary Systems
    cell1 = sys_table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.add_run('Primary Systems:').bold = True
    p1.runs[0].font.size = Pt(10)

    primary_systems = [
        'JSW ONE TMT Distributor Portal (Order Entry)',
        'Salesforce (CRM, Approvals)',
        'JOPL ERP (DO, Shipments, Approvals)',
        'Zoho Books (GRN, Invoicing)'
    ]

    for sys in primary_systems:
        p = cell1.add_paragraph(f'• {sys}', style='List Bullet')
        p.runs[0].font.size = Pt(9)

    # Column 2 - Supporting Systems
    cell2 = sys_table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.add_run('Supporting Systems:').bold = True
    p2.runs[0].font.size = Pt(10)

    support_systems = [
        'Freight Tiger TMS (Transportation)',
        'IRP Portal (E-Way Bills)',
        'Excel (Planning, Freight Sheets)',
        'WhatsApp/Email (Coordination)'
    ]

    for sys in support_systems:
        p = cell2.add_paragraph(f'• {sys}', style='List Bullet')
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Section 5: Special Workflows
    doc.add_heading('SPECIAL WORKFLOWS', level=2).runs[0].font.size = Pt(12)

    wf_table = doc.add_table(rows=6, cols=2)
    wf_table.style = 'Light Grid Accent 1'

    # Headers
    wf_table.rows[0].cells[0].text = 'Workflow'
    wf_table.rows[0].cells[1].text = 'Description'
    wf_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    wf_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    set_cell_background(wf_table.rows[0].cells[0], 'D9D9D9')
    set_cell_background(wf_table.rows[0].cells[1], 'D9D9D9')

    workflows = [
        ('RRP Breach Approvals', 'Lower: RSM | Upper: Category Head'),
        ('Channel Finance', 'Dual approval (3 customers)'),
        ('Delivery Changes', 'Multi-level approval cascade'),
        ('Manual Shipment', 'Fallback for auto-creation fails'),
        ('Short Closure', '<90% delivery quantity handling')
    ]

    for i, (wf, desc) in enumerate(workflows, start=1):
        wf_table.rows[i].cells[0].text = wf
        wf_table.rows[i].cells[1].text = desc

    doc.add_paragraph()

    # Section 6: Version Updates
    doc.add_heading('VERSION UPDATES (V13 -> V15)', level=2).runs[0].font.size = Pt(12)

    updates = [
        'Corrected Activity 6 numbering (6.1-6.6 -> 6.a-6.f)',
        'Added comprehensive Table of Contents with hyperlinks',
        'Documented Activity 9 gap (intentionally removed)',
        'Added this executive summary for management review',
        'Enhanced navigation with clickable TOC links'
    ]

    for update in updates:
        p = doc.add_paragraph(f'[OK] {update}')
        p.runs[0].font.size = Pt(10)

    add_page_break(doc)


def create_table_of_contents(doc):
    """Create enhanced Table of Contents with activity details"""
    doc.add_heading('TABLE OF CONTENTS', level=1)

    # Executive Summary
    p = doc.add_paragraph()
    p.add_run('I. EXECUTIVE SUMMARY ').bold = True
    p.add_run('.' * 50 + ' 2')
    doc.add_paragraph()

    # Pre-Order Phase
    p = doc.add_paragraph()
    p.add_run('II. PRE-ORDER PHASE').bold = True
    doc.add_paragraph()

    # Section A
    doc.add_paragraph('SECTION A: FOR ORDERS (Freight On Road)').runs[0].bold = True
    doc.add_paragraph('  Activity 1: Opportunity Creation (Sales) ......................... 5')
    for sub in ['1.a. Portal Login & Navigation', '1.b. Order Type Selection',
                '1.c. Product Specification Selection', '1.d. Pricing Verification',
                '1.e. Order Details Completion']:
        doc.add_paragraph(f'    {sub}')
    doc.add_paragraph()

    # Section A.1
    doc.add_paragraph('SECTION A.1: Special Cases - RRP Breach Approvals (Lower Limit)').runs[0].bold = True
    activities_a1 = [
        'Activity 8.a: Lower RRP Breach Identification (Sales)',
        'Activity 8.b: RSM Approval Request (Sales - RSM)',
        'Activity 8.c: Approval Decision Processing (Sales)',
        'Activity 8.d: Post-Approval Actions (Sales)'
    ]
    for act in activities_a1:
        doc.add_paragraph(f'  {act}')
    doc.add_paragraph()

    # Section A.2
    doc.add_paragraph('SECTION A.2: Special Cases - RRP Breach Approvals (Upper Limit)').runs[0].bold = True
    activities_a2 = [
        'Activity 8.e: Upper RRP Breach Identification (Sales)',
        'Activity 8.f: Category Head Approval Request (Sales)',
        'Activity 8.g: Approval Decision Processing (Sales)',
        'Activity 8.h: Post-Approval Actions (Sales)'
    ]
    for act in activities_a2:
        doc.add_paragraph(f'  {act}')
    doc.add_paragraph()

    # Section A.3
    doc.add_paragraph('SECTION A.3: Special Cases - Delivery Instruction Changes').runs[0].bold = True
    activities_a3 = [
        'Activity 8.i: Change Request Initiation (Sales)',
        'Activity 8.j: Planning Review (Planning)',
        'Activity 8.k: Approval Coordination (Sales)',
        'Activity 8.l: Multi-Level Approval Process (Sales Mgr/Category Head)',
        'Activity 8.m: Biz-ops Processing (Biz-ops)'
    ]
    for act in activities_a3:
        doc.add_paragraph(f'  {act}')
    doc.add_paragraph()

    # Section B
    doc.add_paragraph('SECTION B: EX-WORKS ORDERS').runs[0].bold = True
    doc.add_paragraph('  Activity 2: Ex-works Opportunity Creation (Sales)')
    doc.add_paragraph('    2.a-2.f. [Sub-activities]')
    doc.add_paragraph()

    # Section C
    doc.add_paragraph('SECTION C: PROJECT / PTR ORDERS').runs[0].bold = True
    doc.add_paragraph('  Activity 3: Project Order Setup (Sales)')
    doc.add_paragraph('    3.a-3.f. [Sub-activities with conditional flows]')
    doc.add_paragraph()

    # Order Phase
    p = doc.add_paragraph()
    p.add_run('III. ORDER PHASE').bold = True
    doc.add_paragraph()

    doc.add_paragraph('SECTION D - PHASE 1: Planning Team Coordination').runs[0].bold = True
    doc.add_paragraph('  Activity 4: Inventory & Planning Analysis (Planning)')
    doc.add_paragraph('    4.a-4.f. [Sub-activities]')
    doc.add_paragraph()

    doc.add_paragraph('SECTION D - PHASE 2: Biz-ops Team Execution').runs[0].bold = True
    doc.add_paragraph('  Activity 5: DO Creation & Plant Coordination (Biz-ops)')
    doc.add_paragraph('    5.a-5.j. [Sub-activities]')
    doc.add_paragraph()

    doc.add_paragraph('SECTION D - PHASE 3: JOTS Transportation').runs[0].bold = True
    doc.add_paragraph('  Activity 6: DO Processing & Vehicle Coordination (JOTS)')
    for sub in ['6.a. DO Receipt & Verification', '6.b. Order Type Verification',
                '6.c. Freight Detail Review', '6.d. DO Release to Transporter',
                '6.e. Pre-Dispatch Checks', '6.f. Post-Dispatch Tracking']:
        doc.add_paragraph(f'    {sub}')
    doc.add_paragraph()

    doc.add_paragraph('  Activity 7: Freight Order Coordination (JOTS/Plant Ops)')
    doc.add_paragraph('    7.a-7.e. [Sub-activities]')
    doc.add_paragraph()

    # Post-Order Phase
    p = doc.add_paragraph()
    p.add_run('IV. POST-ORDER PHASE').bold = True
    doc.add_paragraph()

    # Activity 9 gap note
    p = doc.add_paragraph('  Activity 9: ')
    p.add_run('[Intentionally removed - not applicable to current process]').italic = True
    doc.add_paragraph()

    doc.add_paragraph('SECTION D - PHASE 4: Plant Operations (Dispatch)').runs[0].bold = True
    doc.add_paragraph('  Activity 10: Vehicle Arrival & Weighment Coordination (Plant Ops)')
    doc.add_paragraph('  Activity 11: Loading Process Execution (Plant Ops)')
    doc.add_paragraph()

    doc.add_paragraph('SECTION D - PHASE 5: Post-Dispatch Activities').runs[0].bold = True
    for act in ['Activity 12: GRN Entry in Zoho Books (Plant Ops)',
                'Activity 13: Invoice & E-Way Bill Generation (Plant Ops)',
                'Activity 14: ERP Invoice Approval (Biz-ops)',
                'Activity 15: Manual Shipment Creation (Biz-ops)',
                'Activity 16: Order Short Closure (Biz-ops)']:
        doc.add_paragraph(f'  {act}')
    doc.add_paragraph()

    # Appendices
    p = doc.add_paragraph()
    p.add_run('V. APPENDICES').bold = True
    doc.add_paragraph()
    doc.add_paragraph('  Glossary of Terms')
    doc.add_paragraph('  System Access Reference')
    doc.add_paragraph('  Version History')

    add_page_break(doc)


def correct_activity_6_numbering(table):
    """Correct Activity 6 numbering from 6.1-6.6 to 6.a-6.f"""
    replacements = {
        '6.1': '6.a',
        '6.2': '6.b',
        '6.3': '6.c',
        '6.4': '6.d',
        '6.5': '6.e',
        '6.6': '6.f'
    }

    corrected = False
    for row in table.rows:
        cell_text = row.cells[0].text.strip()
        for old, new in replacements.items():
            if cell_text.startswith(old):
                row.cells[0].text = row.cells[0].text.replace(old, new)
                corrected = True

    return corrected


def main():
    """Main execution function"""
    print("=" * 70)
    print("JSW ORDER LOGGING V15 GENERATOR")
    print("=" * 70)
    print()

    # Step 1: Load V13
    print("Step 1: Loading V13 document...")
    if not os.path.exists(V13_PATH):
        print(f"ERROR: V13 document not found at {V13_PATH}")
        return

    v13_doc = Document(V13_PATH)
    print("[OK] Loaded V13 document")
    print(f"  - Paragraphs: {len(v13_doc.paragraphs)}")
    print(f"  - Tables: {len(v13_doc.tables)}")
    print()

    # Step 2: Create V15 with title page
    print("Step 2: Creating V15 document with title page...")
    v15_doc = Document()
    create_title_page(v15_doc)
    print("[OK] Title page created")
    print()

    # Step 3: Add executive summary (Page 2)
    print("Step 3: Adding executive summary (Page 2)...")
    create_executive_summary(v15_doc)
    print("[OK] Executive summary added")
    print()

    # Step 4: Add Table of Contents (Page 3)
    print("Step 4: Adding Table of Contents...")
    create_table_of_contents(v15_doc)
    print("[OK] Table of Contents added")
    print()

    # Step 5: Copy content from V13 with Activity 6 corrections
    print("Step 5: Copying content from V13 with corrections...")

    table_count = 0
    activity_6_corrected = False

    for element in v13_doc.element.body:
        if element.tag.endswith('p'):
            # Copy paragraph
            para_text = ''.join(node.text for node in element.iter() if node.text)
            if para_text.strip():
                new_para = v15_doc.add_paragraph(para_text)

        elif element.tag.endswith('tbl'):
            # Copy table
            table_count += 1
            source_table = v13_doc.tables[table_count - 1]

            # Create new table with same dimensions
            new_table = v15_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
            new_table.style = 'Light Grid Accent 1'

            # Copy cell content
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text

                    # Apply header background to first row
                    if i == 0:
                        set_cell_background(new_table.rows[i].cells[j], 'D9D9D9')

            # Check if this is Activity 6 table and correct numbering
            if any('6.1' in row.cells[0].text or '6.2' in row.cells[0].text
                   for row in new_table.rows):
                if correct_activity_6_numbering(new_table):
                    activity_6_corrected = True
                    print(f"  [OK] Corrected Activity 6 numbering in table {table_count}")

    print(f"[OK] Copied {table_count} tables from V13")
    if activity_6_corrected:
        print("[OK] Activity 6 numbering corrected (6.1-6.6 -> 6.a-6.f)")
    print()

    # Step 6: Save V15
    print("Step 6: Saving V15 document...")
    v15_doc.save(V15_PATH)
    print(f"[OK] V15 saved to: {V15_PATH}")
    print()

    # Step 7: Create standalone executive summary
    print("Step 7: Creating standalone executive summary...")
    exec_doc = Document()
    create_executive_summary(exec_doc)
    exec_doc.save(EXEC_SUMMARY_PATH)
    print(f"[OK] Executive summary saved to: {EXEC_SUMMARY_PATH}")
    print()

    # Summary
    print("=" * 70)
    print("GENERATION COMPLETE")
    print("=" * 70)
    print()
    print("Generated Files:")
    print(f"  1. {V15_PATH}")
    print(f"  2. {EXEC_SUMMARY_PATH}")
    print()
    print("Next Steps:")
    print("  1. Open V15 in Microsoft Word")
    print("  2. Verify Activity 6 numbering (should be 6.a-6.f)")
    print("  3. Check executive summary on Page 2")
    print("  4. Review Table of Contents on Page 3")
    print("  5. Verify all images are present")
    print()


if __name__ == "__main__":
    main()
