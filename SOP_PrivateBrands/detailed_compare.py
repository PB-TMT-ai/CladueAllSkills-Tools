import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

def get_full_cell_text(cell):
    """Get complete cell text including all paragraphs"""
    return '\n'.join(p.text for p in cell.paragraphs)

def extract_doc(filepath):
    doc = Document(filepath)
    data = {
        'paragraphs': [],
        'tables': []
    }
    for p in doc.paragraphs:
        data['paragraphs'].append({
            'text': p.text.strip(),
            'style': p.style.name
        })
    for table in doc.tables:
        tbl = []
        for row in table.rows:
            r = []
            for cell in row.cells:
                r.append(get_full_cell_text(cell))
            tbl.append(r)
        data['tables'].append(tbl)
    return data

jots = extract_doc(r'D:\SOP_PrivateBrands\Documents\JSWOrderLogging_JOTS_modification.docx')
v7 = extract_doc(r'D:\SOP_PrivateBrands\Documents\JSWOrderLogging_V7.docx')

# Compare paragraphs
print("=" * 70)
print("PARAGRAPH COMPARISON")
print("=" * 70)
max_p = max(len(jots['paragraphs']), len(v7['paragraphs']))
para_diffs = 0
for i in range(max_p):
    j_text = jots['paragraphs'][i]['text'] if i < len(jots['paragraphs']) else '<MISSING>'
    v_text = v7['paragraphs'][i]['text'] if i < len(v7['paragraphs']) else '<MISSING>'
    j_style = jots['paragraphs'][i]['style'] if i < len(jots['paragraphs']) else '<MISSING>'
    v_style = v7['paragraphs'][i]['style'] if i < len(v7['paragraphs']) else '<MISSING>'
    if j_text != v_text or j_style != v_style:
        para_diffs += 1
        print(f"\n--- Paragraph {i} DIFFERS ---")
        if j_style != v_style:
            print(f"  JOTS style: {j_style}")
            print(f"  V7   style: {v_style}")
        if j_text != v_text:
            print(f"  JOTS: {j_text[:300]}")
            print(f"  V7:   {v_text[:300]}")

if para_diffs == 0:
    print("  No paragraph differences found.")
print(f"\nTotal paragraph differences: {para_diffs}")

# Compare tables
print("\n" + "=" * 70)
print("TABLE COMPARISON")
print("=" * 70)
print(f"JOTS tables: {len(jots['tables'])}, V7 tables: {len(v7['tables'])}")

max_t = max(len(jots['tables']), len(v7['tables']))
for t in range(max_t):
    if t >= len(jots['tables']):
        print(f"\n--- Table {t+1}: ONLY IN V7 ({len(v7['tables'][t])} rows) ---")
        continue
    if t >= len(v7['tables']):
        print(f"\n--- Table {t+1}: ONLY IN JOTS ({len(jots['tables'][t])} rows) ---")
        continue

    jt = jots['tables'][t]
    vt = v7['tables'][t]

    if len(jt) != len(vt):
        print(f"\n--- Table {t+1}: ROW COUNT DIFFERS (JOTS={len(jt)}, V7={len(vt)}) ---")

    max_r = max(len(jt), len(vt))
    table_has_diff = False
    for r in range(max_r):
        if r >= len(jt):
            print(f"  Row {r}: ONLY IN V7: {[c[:80] for c in vt[r]]}")
            table_has_diff = True
            continue
        if r >= len(vt):
            print(f"  Row {r}: ONLY IN JOTS: {[c[:80] for c in jt[r]]}")
            table_has_diff = True
            continue

        # Compare each cell
        max_c = max(len(jt[r]), len(vt[r]))
        row_diffs = []
        for c in range(max_c):
            j_cell = jt[r][c] if c < len(jt[r]) else '<MISSING>'
            v_cell = vt[r][c] if c < len(vt[r]) else '<MISSING>'
            if j_cell != v_cell:
                row_diffs.append(c)

        if row_diffs:
            table_has_diff = True
            # Get activity number for context
            act = jt[r][0][:15] if jt[r] else '?'
            print(f"\n  Table {t+1}, Row {r} (Activity: {act.strip()}) - Cols differ: {row_diffs}")
            for c in row_diffs:
                j_cell = jt[r][c] if c < len(jt[r]) else '<MISSING>'
                v_cell = vt[r][c] if c < len(vt[r]) else '<MISSING>'
                col_names = ['Activity', 'Steps', 'Team', 'Interface/Utilities', 'Sign off']
                col_name = col_names[c] if c < len(col_names) else f'Col{c}'
                print(f"    [{col_name}]")
                print(f"      JOTS: {j_cell[:250]}")
                print(f"      V7:   {v_cell[:250]}")

    if not table_has_diff:
        print(f"  Table {t+1}: IDENTICAL ({len(jt)} rows)")
