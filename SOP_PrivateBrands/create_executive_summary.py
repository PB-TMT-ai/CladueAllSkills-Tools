#!/usr/bin/env python
"""
Create Management Executive Summary (One-Pager)
Standalone document for top management presentation
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\SOP_PrivateBrands\JSWOrderLogging_ExecutiveSummary_v1.docx"


def set_cell_background(cell, color):
    """Set background color for a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_executive_summary():
    """Create one-page executive summary"""
    doc = Document()

    # Set narrow margins for more content
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header
    header = doc.add_heading('JSW ONE TMT - ORDER LOGGING PROCESS', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(14)

    subheader = doc.add_paragraph('Executive Summary')
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader.runs[0].font.size = Pt(12)
    subheader.runs[0].italic = True

    date_str = datetime.now().strftime('%B %d, %Y')
    doc.add_paragraph(f'Date: {date_str} | Document Owner: Operations Team')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Section 1: Process Overview
    doc.add_heading('PROCESS OVERVIEW', level=2).runs[0].font.size = Pt(11)
    overview = doc.add_paragraph(
        'End-to-end documentation of JSW ONE TMT order lifecycle covering opportunity '
        'creation through dispatch, GRN entry, and invoice generation. Encompasses retail '
        'FOR orders, ex-works orders, and project/PTR orders with comprehensive coverage of '
        'special cases including RRP breach approvals, channel finance workflows, and delivery '
        'instruction changes.'
    )
    overview.runs[0].font.size = Pt(9)

    stats = doc.add_paragraph('3 Phases | 15 Activities | 85+ Steps | 46 Screenshots | 5 Teams | 8+ Systems')
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats.runs[0].bold = True
    stats.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Section 2: Phase Breakdown
    doc.add_heading('PHASE BREAKDOWN', level=2).runs[0].font.size = Pt(11)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'

    # Headers
    headers = ['Phase', 'Activities', 'Primary Team', 'Key Steps']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_background(cell, 'D9D9D9')

    # Data
    data = [
        ['Pre-Order', '1-3, 8.a-8.m', 'Sales', '56 steps'],
        ['Order', '4-7', 'Planning/Biz-ops/JOTS', '29 steps'],
        ['Post-Order', '10-16', 'Plant Ops/Biz-ops', '26 steps']
    ]

    for i, row_data in enumerate(data, start=1):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Section 3: Team Responsibilities
    doc.add_heading('TEAM RESPONSIBILITIES', level=2).runs[0].font.size = Pt(11)

    teams = [
        ('Sales', 'Opportunity creation, RRP approvals, special cases (Act 1-3, 8.a-8.h)'),
        ('Planning', 'Inventory analysis, plant allocation (Act 4, 8.j)'),
        ('Biz-ops', 'DO creation, ERP integration, closures (Act 5, 8.m, 14-16)'),
        ('JOTS', 'Transportation, vehicle coordination (Act 6-7)'),
        ('Plant Operations', 'Weighment, loading, GRN, invoicing (Act 10-13)')
    ]

    for team, desc in teams:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{team}: ').bold = True
        p.add_run(desc)
        p.runs[0].font.size = Pt(9)
        p.runs[1].font.size = Pt(9)

    doc.add_paragraph()

    # Section 4: Key Systems
    doc.add_heading('KEY SYSTEMS & INTERFACES', level=2).runs[0].font.size = Pt(11)

    sys_table = doc.add_table(rows=1, cols=2)

    # Column 1 - Primary Systems
    cell1 = sys_table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.add_run('Primary Systems:').bold = True
    p1.runs[0].font.size = Pt(9)

    primary_systems = [
        'JSW ONE TMT Distributor Portal',
        'Salesforce (CRM, Approvals)',
        'JOPL ERP (DO, Shipments)',
        'Zoho Books (GRN, Invoicing)'
    ]

    for sys in primary_systems:
        p = cell1.add_paragraph(f'• {sys}')
        p.runs[0].font.size = Pt(8)

    # Column 2 - Supporting Systems
    cell2 = sys_table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.add_run('Supporting Systems:').bold = True
    p2.runs[0].font.size = Pt(9)

    support_systems = [
        'Freight Tiger TMS',
        'IRP Portal (E-Way Bills)',
        'Excel (Planning, Freight)',
        'WhatsApp/Email'
    ]

    for sys in support_systems:
        p = cell2.add_paragraph(f'• {sys}')
        p.runs[0].font.size = Pt(8)

    doc.add_paragraph()

    # Section 5: Special Workflows
    doc.add_heading('SPECIAL WORKFLOWS', level=2).runs[0].font.size = Pt(11)

    wf_table = doc.add_table(rows=6, cols=2)
    wf_table.style = 'Light Grid Accent 1'

    # Headers
    wf_table.rows[0].cells[0].text = 'Workflow'
    wf_table.rows[0].cells[1].text = 'Description'
    wf_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    wf_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    wf_table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    wf_table.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_background(wf_table.rows[0].cells[0], 'D9D9D9')
    set_cell_background(wf_table.rows[0].cells[1], 'D9D9D9')

    workflows = [
        ('RRP Breach Approvals', 'Lower: RSM | Upper: Category Head'),
        ('Channel Finance', 'Dual approval (3 customers enrolled)'),
        ('Delivery Changes', 'Multi-level approval cascade'),
        ('Manual Shipment', 'Fallback for auto-creation fails'),
        ('Short Closure', '<90% delivery quantity handling')
    ]

    for i, (wf, desc) in enumerate(workflows, start=1):
        wf_table.rows[i].cells[0].text = wf
        wf_table.rows[i].cells[1].text = desc
        wf_table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        wf_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph()

    # Section 6: Document Information
    doc.add_heading('DOCUMENT INFORMATION', level=2).runs[0].font.size = Pt(11)

    info = [
        'Current Version: V15 (February 2026)',
        'Total Activities: 15 (Activity 9 intentionally removed)',
        'Coverage: Pre-Order, Order, Post-Order phases',
        'Documentation: 131 paragraphs, 18 tables, 46 screenshots',
        'Latest Update: Activity 6 index numbering standardized (6.a-6.f)'
    ]

    for item in info:
        p = doc.add_paragraph(f'• {item}')
        p.runs[0].font.size = Pt(9)

    return doc


def main():
    print("=" * 70)
    print("CREATING MANAGEMENT EXECUTIVE SUMMARY")
    print("=" * 70)
    print()

    print("Generating one-page executive summary...")
    doc = create_executive_summary()

    print("Saving document...")
    doc.save(OUTPUT_PATH)

    print(f"[OK] Executive summary saved to:")
    print(f"     {OUTPUT_PATH}")
    print()
    print("Use cases:")
    print("  - Board presentations")
    print("  - Executive briefings")
    print("  - Quarterly reviews")
    print("  - Stakeholder onboarding")
    print()


if __name__ == "__main__":
    main()
