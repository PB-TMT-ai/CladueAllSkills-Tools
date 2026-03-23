"""
Generate a distributor one-pager .docx report from extracted data.
Replicates the 'Basic format' template layout from the Excel file.
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Add parent dir to path for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from extract_distributor_data import extract_all, FY26_MONTHS


# ─── Styling helpers ─────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=8, align='left', color=None):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)

    # Reduce cell spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def format_number(val):
    """Format number as integer with comma separators."""
    if val is None or val == 0:
        return '-'
    return f"{int(round(val)):,}"


def format_pct(val):
    """Format percentage as integer."""
    if val is None or val == 0:
        return '-'
    return f"{val:.0f}%"


def set_table_borders(table, border_color="000000", border_size="4"):
    """Set borders for entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'  <w:insideH w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'  <w:insideV w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_cell_vertical_center(cell):
    """Set vertical alignment to center for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def add_section_header(doc, title):
    """Add a section header paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)


HEADER_BG = "4472C4"  # Blue header
HEADER_COLOR = (255, 255, 255)  # White text
LIGHT_BG = "D9E2F3"  # Light blue for sub-headers


# ─── Section Builders ────────────────────────────────────────

def add_section1(doc, data):
    """Distributor basic details."""
    add_section_header(doc, "Distributor basic details:")
    s = data['section1']

    table = doc.add_table(rows=5, cols=2)
    set_table_borders(table)
    table.autofit = True

    fields = [
        ("Distributor name:", s['distributor_name']),
        ("State:", s['state']),
        ("Districts served:", s['districts_served']),
        ("Retail potential/month (MT):", format_number(s['retail_potential_month'])),
        ("# Very High / High / Medium districts:", s['vh_h_m_districts']),
    ]

    for i, (label, value) in enumerate(fields):
        set_cell_text(table.cell(i, 0), label, bold=True, size=9)
        set_cell_text(table.cell(i, 1), str(value), size=9)
        set_cell_shading(table.cell(i, 0), LIGHT_BG)


def add_section2(doc, data):
    """Sales manpower."""
    add_section_header(doc, "Sales manpower:")
    s = data['section2']

    table = doc.add_table(rows=5, cols=2)
    set_table_borders(table)
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(8)

    # Header
    for j, h in enumerate(["Designation", "Name"]):
        set_cell_text(table.cell(0, j), h, bold=True, size=8, color=HEADER_COLOR)
        set_cell_shading(table.cell(0, j), HEADER_BG)

    rows_data = [
        ("SM:", s['sm']),
        ("TM:", s['tm']),
        ("# DGO:", str(s['dgo'])),
        ("# DSR:", str(s['dsr'])),
    ]
    for i, (a, b) in enumerate(rows_data):
        set_cell_text(table.cell(i + 1, 0), a, bold=True, size=8)
        set_cell_text(table.cell(i + 1, 1), b, size=8)


def add_section3(doc, data):
    """Sales performance."""
    add_section_header(doc, "Sales performance:")
    s = data['section3']

    # Summary row
    table1 = doc.add_table(rows=2, cols=3)
    set_table_borders(table1)
    for j, h in enumerate(["FY Revised MoU", "FY 26 Achieved", "% Achievement"]):
        set_cell_text(table1.cell(0, j), h, bold=True, size=8, color=HEADER_COLOR, align='center')
        set_cell_shading(table1.cell(0, j), HEADER_BG)

    set_cell_text(table1.cell(1, 0), format_number(s['revised_mou']), size=9, align='center')
    set_cell_text(table1.cell(1, 1), format_number(s['achieved_total']), size=9, align='center')
    set_cell_text(table1.cell(1, 2), format_pct(s['pct_achievement']), size=9, align='center')

    # Monthly detail table (11 months, no Mar'26)
    doc.add_paragraph()  # spacer
    months = data['months'][:11]

    table2 = doc.add_table(rows=4, cols=14)
    set_table_borders(table2)

    # Set column widths: label + MoU + TD + 11 months
    table2.columns[0].width = Cm(1.8)
    table2.columns[1].width = Cm(2.0)
    table2.columns[2].width = Cm(2.0)
    for j in range(3, 14):
        table2.columns[j].width = Cm(2.0)

    # Headers
    headers = ["FY 26:", "FY 26 MoU*", "FY 26 TD\n(in MT)"] + months
    for j, h in enumerate(headers):
        set_cell_text(table2.cell(0, j), h, bold=True, size=7, color=HEADER_COLOR, align='center')
        set_cell_shading(table2.cell(0, j), HEADER_BG)

    # Target row
    set_cell_text(table2.cell(1, 0), "Target:", bold=True, size=7)
    set_cell_text(table2.cell(1, 1), format_number(s['target_total']), size=7, align='center')
    set_cell_text(table2.cell(1, 2), format_number(s['target_total']), size=7, align='center')
    for i in range(11):
        set_cell_text(table2.cell(1, 3 + i), format_number(s['monthly_targets'][i]), size=7, align='center')

    # Achieved row
    set_cell_text(table2.cell(2, 0), "Achieved:", bold=True, size=7)
    achieved_td = sum(s['monthly_achieved'][:11])
    set_cell_text(table2.cell(2, 1), format_number(achieved_td), size=7, align='center')
    set_cell_text(table2.cell(2, 2), format_number(s['achieved_total']), size=7, align='center')
    for i in range(11):
        set_cell_text(table2.cell(2, 3 + i), format_number(s['monthly_achieved'][i]), size=7, align='center')

    # % achieved row
    set_cell_text(table2.cell(3, 0), "% achieved:", bold=True, size=7)
    total_pct_mou = (achieved_td / s['target_total'] * 100) if s['target_total'] else 0
    set_cell_text(table2.cell(3, 1), format_pct(total_pct_mou), size=7, align='center')
    td_pct = (s['achieved_total'] / s['target_total'] * 100) if s['target_total'] else 0
    set_cell_text(table2.cell(3, 2), format_pct(td_pct), size=7, align='center')
    for i in range(11):
        set_cell_text(table2.cell(3, 3 + i), format_pct(s['monthly_pct'][i]), size=7, align='center')

    # Footnote
    p = doc.add_paragraph()
    run = p.add_run("*Original MoU target")
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.name = 'Calibri'


def add_section4(doc, data):
    """Sales by location."""
    add_section_header(doc, "Sales by location:")
    s = data['section4']
    months = data['months'][:11]

    table = doc.add_table(rows=7, cols=12)
    set_table_borders(table)

    # Header row with months (no Mar'26)
    set_cell_text(table.cell(0, 0), "", size=7)
    for j, m in enumerate(months):
        set_cell_text(table.cell(0, j + 1), m, bold=True, size=7, color=HEADER_COLOR, align='center')
        set_cell_shading(table.cell(0, j + 1), HEADER_BG)

    # Row specs: (label, is_parent, value_key, pct_key)
    rows_spec = [
        ("Retailer", True, 'retailer_total', None),
        ("  Shop", False, 'retailer_shop', 'retailer_shop_pct'),
        ("  Site", False, 'retailer_site', 'retailer_site_pct'),
        ("Self-stocking", True, 'selfstock_total', None),
        ("  Warehouse", False, 'selfstock_warehouse', 'selfstock_warehouse_pct'),
        ("  Site", False, 'selfstock_site', 'selfstock_site_pct'),
    ]

    for i, (label, is_parent, val_key, pct_key) in enumerate(rows_spec):
        set_cell_text(table.cell(i + 1, 0), label, bold=is_parent, size=7)
        if is_parent:
            set_cell_shading(table.cell(i + 1, 0), LIGHT_BG)
        for j in range(11):
            val = s[val_key][j]
            if is_parent:
                # Parent row: show integer value
                set_cell_text(table.cell(i + 1, j + 1),
                              format_number(val) if val else '-', size=7, align='center')
            else:
                # Child row: show only percentage
                pct = s[pct_key][j] if pct_key else 0
                if val:
                    set_cell_text(table.cell(i + 1, j + 1), f"{pct:.0f}%", size=7, align='center')
                else:
                    set_cell_text(table.cell(i + 1, j + 1), '-', size=7, align='center')


def add_section5(doc, data):
    """Channel performance."""
    add_section_header(doc, "Channel performance:")
    s = data['section5']
    months = data['months'][:11]

    table = doc.add_table(rows=6, cols=13)
    set_table_borders(table)

    # Headers (no Mar'26)
    headers = ["Metric", "FY 26"] + months
    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, bold=True, size=7, color=HEADER_COLOR, align='center')
        set_cell_shading(table.cell(0, j), HEADER_BG)

    # Row definitions: (label, monthly_data, fy26_total, is_pct)
    row_defs = [
        ("Sec. sales:", s['sec_sales'], s['sec_sales_total'], False),
        ("# new dealers:", s['new_dealers'], s['new_dealers_total'], False),
        ("# transacting dealers:", s['transacting'], s['transacting_fy26'], False),
        ("# active dealers:", s['active'], s['active_fy26'], False),
        ("trans./active ratio:", s['trans_active_ratio'], None, True),
    ]

    for i, (label, monthly, total, is_pct) in enumerate(row_defs):
        set_cell_text(table.cell(i + 1, 0), label, bold=True, size=7)
        if total is not None:
            set_cell_text(table.cell(i + 1, 1), format_number(total), size=7, align='center')

        for j in range(11):
            val = monthly[j] if monthly[j] is not None else None
            if val is not None:
                if is_pct:
                    txt = format_pct(val * 100)
                else:
                    txt = format_number(val)
                set_cell_text(table.cell(i + 1, j + 2), txt, size=7, align='center')


def add_section6(doc, data):
    """Performance in key districts."""
    add_section_header(doc, "Performance in key districts (current month):")
    districts = data['section6']

    table = doc.add_table(rows=1 + len(districts), cols=10)
    set_table_borders(table)

    headers = ["District Name", "District\nCategory", "District\nPotential (MT)",
               "Sales vol.\n(MT)", "# billed\ndealers TD", "# transacting\ndealers",
               "# active\ndealers", "SoB (%)", "Reach (%)", "Market\nshare (%)"]

    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, bold=True, size=7, color=HEADER_COLOR, align='center')
        set_cell_shading(table.cell(0, j), HEADER_BG)

    for i, d in enumerate(districts):
        set_cell_text(table.cell(i + 1, 0), d['district'], size=7)
        set_cell_text(table.cell(i + 1, 1), d['category'], size=7, align='center')
        set_cell_text(table.cell(i + 1, 2), format_number(d['potential']), size=7, align='center')
        set_cell_text(table.cell(i + 1, 3), format_number(d['sales_vol']), size=7, align='center')
        set_cell_text(table.cell(i + 1, 4), str(d['billed_td']), size=7, align='center')
        set_cell_text(table.cell(i + 1, 5), str(d['transacting']), size=7, align='center')
        set_cell_text(table.cell(i + 1, 6), str(d['active']), size=7, align='center')
        set_cell_text(table.cell(i + 1, 7), format_pct(d['sob_pct'] * 100), size=7, align='center')
        set_cell_text(table.cell(i + 1, 8), format_pct(d['reach_pct'] * 100), size=7, align='center')
        set_cell_text(table.cell(i + 1, 9), format_pct(d['market_share_pct'] * 100), size=7, align='center')


def add_section7(doc, data):
    """Performance by key dealers (top 10)."""
    dealers = data['section7']
    months = data['months'][:11]

    add_section_header(doc, f"Performance by key dealers (top {len(dealers)} dealers):")

    table = doc.add_table(rows=1 + len(dealers), cols=14)
    set_table_borders(table)

    # Set column widths: wider for name/district, narrower for months
    table.columns[0].width = Cm(5.5)   # Dealer Name
    table.columns[1].width = Cm(3.0)   # District
    table.columns[2].width = Cm(2.5)   # Classification
    for j in range(3, 14):
        table.columns[j].width = Cm(1.5)  # Month cols

    headers = ["Dealer Name", "District", "Classification"] + months
    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, bold=True, size=7, color=HEADER_COLOR, align='center')
        set_cell_shading(table.cell(0, j), HEADER_BG)

    for i, d in enumerate(dealers):
        set_cell_text(table.cell(i + 1, 0), d['name'], size=7)
        set_cell_text(table.cell(i + 1, 1), d['district'], size=7)
        set_cell_text(table.cell(i + 1, 2), d['segmentation'], size=7, align='center')
        for j in range(11):
            val = d['monthly'][j]
            set_cell_text(table.cell(i + 1, 3 + j),
                          format_number(val) if val else '-', size=7, align='center')


# ─── Main Document Builder ───────────────────────────────────

def generate_one_pager(data, output_path):
    """Generate the complete one-pager document."""
    doc = Document()

    # Set narrow margins for landscape-like fit
    for section in doc.sections:
        section.page_width = Cm(29.7)   # A4 landscape width
        section.page_height = Cm(21.0)  # A4 landscape height
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    # Title: JSW ONE TMT - <Name> - Performance <Current-1 Month>
    dist_name = data['section1']['distributor_name']
    # Current-1 month = last month with data (Feb'26 = index 10)
    perf_month = data['months'][10]  # Feb'26
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(f"JSW ONE TMT - {dist_name} - Performance {perf_month}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # Build all sections
    add_section1(doc, data)
    add_section2(doc, data)
    add_section3(doc, data)
    add_section4(doc, data)
    add_section5(doc, data)
    add_section6(doc, data)
    add_section7(doc, data)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


def main():
    dist = sys.argv[1] if len(sys.argv) > 1 else "NIKUNJ UDYOG"
    state = sys.argv[2] if len(sys.argv) > 2 else "HARYANA"

    # Extract data
    data = extract_all(dist, state)

    # Generate output
    safe_name = dist.replace(' ', '_').replace('/', '_')
    safe_state = state.replace(' ', '_')
    output_path = os.path.join(r"D:\Distributor One pager\output", f"{safe_name}_{safe_state}.docx")

    generate_one_pager(data, output_path)


if __name__ == '__main__':
    main()
